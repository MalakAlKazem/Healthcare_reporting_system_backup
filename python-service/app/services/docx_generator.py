"""
DOCX Report Generator for Medication Error Analysis
====================================================
Based on mortality docx_generator.py structure.

Pages:
  1. Metadata table, results table (last 6 history rates), analysis box with chart1 (trend)
  2. Chart 2 (ME count comparison), Chart 3 (error cycle pie)
  3. Chart 4 (detection donut), Chart 5 (shift donut), Chart 6 (staff donut)
  Last. Final result, previous/current action tables, approval table
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

from loguru import logger
import re
from app.medication_error.chart_generator import MedicationErrorCharts
from app.medication_error.medication_error_ai_service import medication_error_ai_service

# =============================================================================
# FONTS
# =============================================================================
FONT = 'Traditional Arabic'
FONT_EN = 'Calibri'
FONT_ANALYSIS = 'Frutiger LT Arabic 45 Light'

# =============================================================================
# LAYOUT CONSTANTS
# =============================================================================
PAGE_WIDTH = Inches(8.27)
PAGE_HEIGHT = Inches(11.69)
MARGIN = Inches(0.7)
USABLE_WIDTH = Inches(6.87)

# =============================================================================
# BORDER PRESETS
# =============================================================================
BORDER_THICK  = {'val': 'single', 'sz': '18', 'color': '000000', 'space': '0'}
BORDER_MEDIUM = {'val': 'single', 'sz': '12', 'color': '000000', 'space': '0'}
BORDER_THIN   = {'val': 'single', 'sz': '8',  'color': '000000', 'space': '0'}
BORDER_FINE   = {'val': 'single', 'sz': '4',  'color': '000000', 'space': '0'}
BORDER_NONE   = {'val': 'none',   'sz': '0',  'color': 'FFFFFF', 'space': '0'}

# =============================================================================
# SHADING COLORS
# =============================================================================
SHADE_HEADER = 'EDEBE3'
SHADE_TABLE  = 'D9E2F3'
SHADE_TOTAL  = 'F2F2F2'


class MedicationErrorDocxGenerator:
    """Generates a medication error analysis DOCX report."""

    def __init__(self):
        self.logo_path = 'app/assets/LOGO.png'

    # =========================================================================
    # PUBLIC API
    # =========================================================================

    async def generate_report(self, stats, history=None, options=None):
        """
        Generate DOCX report with medication error charts.

        Args:
            stats:   Statistics dict from MedicationErrorStatistics.calculate_all_statistics()
            history: List of historical quarters from MedicationErrorHistory
            options: Dict with 'quarter' and 'year'

        Returns:
            dict with 'filePath' and 'fileName'
        """
        if 'statistics' in (stats or {}):
            stats = stats['statistics']
        options  = options or {}
        summary  = stats.get('summary', {})
        quarter  = options.get('quarter', summary.get('quarter', 'الفصل الثالث'))
        year     = str(options.get('year', summary.get('year', '2025')))
        logger.info(f"Generating Medication Error DOCX report for {quarter} {year}...")

        # Build chart history: history entries + current quarter appended
        # (charts 1 & 2 need the current quarter so it shows up on the trend/comparison)
        current_entry = {
            'quarter':      quarter,
            'year':         str(year),
            'error_rate':   summary.get('error_rate', 0),
            'total_errors': summary.get('total_errors', 0),
            'total_doses':  summary.get('total_doses', 0),
        }
        chart_history = list(history or []) + [current_entry]

        # Generate charts
        chart_gen = MedicationErrorCharts(output_dir='storage/charts/medication_error')
        charts = chart_gen.generate_all_charts(stats, chart_history)
        logger.info(f"Generated {len(charts)} charts")

        # Build document
        doc = Document()
        self._setup_section(doc.sections[0])

        self._build_page1(doc, quarter, year, stats, charts, history or [])
        self._build_page2(doc, charts, stats, history=history or [])
        self._build_page3(doc, charts, stats=stats)
        self._build_page4(doc, charts, stats, history=history or [])
        self._build_page5(doc, stats)
        self._build_last_page(doc)

        # Save
        reports_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'storage', 'reports')
        reports_dir = os.path.normpath(reports_dir)
        os.makedirs(reports_dir, exist_ok=True)
        file_name = f'medication-error report {quarter} {year}.docx'
        file_path = os.path.join(reports_dir, file_name)
        doc.save(file_path)
        logger.success(f"DOCX report saved: {file_path}")

        return {'filePath': file_path, 'fileName': file_name}

    # =========================================================================
    # SECTION SETUP
    # =========================================================================

    def _setup_section(self, section):
        section.page_width   = PAGE_WIDTH
        section.page_height  = PAGE_HEIGHT
        section.top_margin   = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin  = MARGIN
        section.right_margin = MARGIN

        # Page border
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '12')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

        self._add_header(section)
        self._add_footer(section)

    def _add_header(self, section):
        header = section.header
        for element in list(header._element):
            header._element.remove(element)

        table = header.add_table(rows=1, cols=3, width=Inches(7))
        self._set_table_cell_margins(table, top=0, bottom=0, left=0, right=0)
        cells = table.rows[0].cells
        cells[0].width = Inches(1.0)
        cells[1].width = Inches(5.0)
        cells[2].width = Inches(1.0)

        # Logo
        if os.path.exists(self.logo_path):
            para = cells[0].paragraphs[0]
            para.add_run().add_picture(self.logo_path, width=Inches(0.5))
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._zero_spacing(para)

        # Title
        para = cells[1].paragraphs[0]
        run = para.add_run("نموذج تحليل البيانات")
        run.font.name = FONT_ANALYSIS
        run.font.size = Pt(16)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(para)
        self._zero_spacing(para)

        # Form number
        para = cells[2].paragraphs[0]
        run = para.add_run('QS-36F-103(4)\nApd 15/12/2019')
        run.font.name = FONT_EN
        run.font.size = Pt(7)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._zero_spacing(para)

        # Underline below header
        line_para = header.add_paragraph()
        line_para.paragraph_format.space_before = Pt(0)
        line_para.paragraph_format.space_after  = Pt(8)
        line_para.paragraph_format.line_spacing = Pt(1)
        line_para.add_run("").font.size = Pt(1)

        pPr   = line_para._element.get_or_add_pPr()
        pBdr  = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_footer(self, section):
        footer = section.footer
        for element in list(footer._element):
            footer._element.remove(element)

        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._element.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

        run.font.size = Pt(10)
        run.font.name = FONT_EN

    # =========================================================================
    # PAGE 1: Metadata + Results table + Analysis box (chart1 trend)
    # =========================================================================

    def _build_page1(self, doc, quarter, year, stats, charts, history):
        self._add_spacer(doc, space_before=6)
        self._add_metadata_table(doc, quarter, year)
        doc.add_paragraph()
        self._add_results_table(doc, stats, history)
        self._add_spacer(doc)
        self._add_analysis_box(doc, quarter, year, stats, charts, history=history)

    def _add_metadata_table(self, doc, quarter, year):
        """
        3-row × 4-col metadata table.
        Changes vs mortality:
          - موضوع التحليل value → 'Medication Error'
          - الإدارة value       → 'التمريض'
          - مصادر البيانات      → 3 medication-error sources
          - الفصل/العام format  → 'quarter/year'
        """
        table = doc.add_table(rows=3, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=0, bottom=0, left=60, right=60)

        col_w = [Inches(1.93), Inches(1.06), Inches(2.71), Inches(1.17)]
        for i, width in enumerate(col_w):
            for row in table.rows:
                row.cells[i].width = width

        # --- Row 0 ---
        self._style_cell_text(table.cell(0, 3), "موضوع التحليل", FONT, 10, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)
        self._style_cell_text(table.cell(0, 2), "Medication Error", FONT_EN, 11, bold=True)
        self._style_cell_text(table.cell(0, 1), "الإدارة", FONT, 10, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(0, 0), "التمريض", FONT, 11, rtl=True)

        # --- Row 1 ---
        self._style_cell_text(table.cell(1, 3), "", FONT, 9, rtl=True)
        self._set_cell_shading(table.cell(1, 3), SHADE_HEADER)
        self._style_cell_text(table.cell(1, 2), "", FONT_EN, 10)
        self._style_cell_text(table.cell(1, 1), "الوحدة الإدارية", FONT, 10, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(1, 0), "....................", FONT, 10, rtl=True)

        # --- Row 2 ---
        self._style_cell_text(table.cell(2, 3),
            "مصادر\nالبيانات\n)تحديد اسماء\nالنماذج(", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(2, 3), SHADE_HEADER)

        sources = (
            "- نموذج حادث مريض\n"
            "- MEDICATION ERROR INVESTIGATION FORM\n"
            "- تدقيق المشرف التمريضي"
        )
        self._style_cell_text(table.cell(2, 2), sources, FONT, 9, rtl=True)

        self._style_cell_text(table.cell(2, 1),
            "الشهر / العام\nأو\nالفصل / العام", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(2, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(2, 0), f"{quarter}/{year}", FONT, 11, rtl=True)

        # Vertical merges (cols 2 & 3: rows 0+1)
        table.cell(0, 2).merge(table.cell(1, 2))
        table.cell(0, 3).merge(table.cell(1, 3))
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)

        # Borders
        row0_tcs = table.rows[0]._tr.findall(qn('w:tc'))
        row1_tcs = table.rows[1]._tr.findall(qn('w:tc'))

        self._set_cell_border(table.cell(0, 0), top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THICK,  right=BORDER_THIN)
        self._set_cell_border(table.cell(1, 0), top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THICK,  right=BORDER_THIN)
        self._set_cell_border(table.cell(2, 0), top=BORDER_THIN,  bottom=BORDER_THICK, left=BORDER_THICK,  right=BORDER_THIN)

        self._set_cell_border(table.cell(0, 1), top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THIN,   right=BORDER_THICK)
        self._set_cell_border(table.cell(1, 1), top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THIN,   right=BORDER_THICK)
        self._set_cell_border(table.cell(2, 1), top=BORDER_THIN,  bottom=BORDER_THICK, left=BORDER_THIN,   right=BORDER_THICK)

        self._set_tc_border(row0_tcs[2], top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THICK, right=BORDER_THIN)
        self._set_tc_border(row1_tcs[2], top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THICK, right=BORDER_THIN)
        self._set_cell_border(table.cell(2, 2), top=BORDER_THIN, bottom=BORDER_THICK, left=BORDER_THICK, right=BORDER_THIN)

        self._set_tc_border(row0_tcs[3], top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THIN, right=BORDER_THICK)
        self._set_tc_border(row1_tcs[3], top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THIN, right=BORDER_THICK)
        self._set_cell_border(table.cell(2, 3), top=BORDER_THIN, bottom=BORDER_THICK, left=BORDER_THIN, right=BORDER_THICK)

    def _add_results_table(self, doc, stats, history):
        """
        4-row × 7-col results table.
        Cols 0-5 → last 6 history quarters (oldest→newest left→right)
        Col 6    → النتيجة الحالية (current error rate)
        """
        COLS = 7
        table = doc.add_table(rows=4, cols=COLS)
        table.autofit = False
        self._set_table_cell_margins(table, top=40, bottom=40, left=60, right=60)

        w_col = Inches(6.87 / COLS)
        for i in range(COLS):
            for row in table.rows:
                row.cells[i].width = w_col

        # Row 0: merged title
        cell0 = table.cell(0, 0)
        for c in range(1, COLS):
            cell0 = cell0.merge(table.cell(0, c))
        self._style_cell_text(table.cell(0, 0), "النتائج", FONT, 11, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)

        # Row 1: instruction (cols 0-5 merged) + current label (col 6)
        instr = table.cell(1, 0)
        for c in range(1, COLS - 1):
            instr = instr.merge(table.cell(1, c))
        self._style_cell_text(
            table.cell(1, 0),
            "في حال كان هناك نتائج سابقة عن الموضوع الذي يتم تحليله يجب ذكره مع تحديد الفترة:\nالشهر / الفصل / العام",
            FONT, 10, rtl=True
        )
        self._set_cell_shading(table.cell(1, 0), SHADE_HEADER)
        self._style_cell_text(table.cell(1, COLS - 1), "النتيجة الحالية", FONT, 11, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, COLS - 1), SHADE_HEADER)

        # Row 2: last 6 history quarter labels
        last6 = history[-(COLS - 1):] if len(history) >= COLS - 1 else history
        while len(last6) < COLS - 1:
            last6 = [{}] + last6

        for i in range(COLS - 1):
            entry = last6[i]
            label = f"{entry.get('quarter', '')} {entry.get('year', '')}" if entry.get('quarter') else ""
            self._style_cell_text(table.cell(2, i), label, FONT, 9, rtl=True)
            self._set_cell_shading(table.cell(2, i), SHADE_HEADER)
        self._style_cell_text(table.cell(2, COLS - 1), "", FONT, 10, rtl=True)
        self._set_cell_shading(table.cell(2, COLS - 1), SHADE_HEADER)

        # Row 3: last 6 error rates + current rate
        for i in range(COLS - 1):
            entry = last6[i]
            val = f"{entry['error_rate']:.4f}%" if entry.get('error_rate') is not None else ""
            self._style_cell_text(table.cell(3, i), val, FONT_EN, 10)

        current_rate = stats.get('summary', {}).get('error_rate', 0)
        try:
            curr = float(current_rate)
        except Exception:
            curr = 0.0
        self._style_cell_text(table.cell(3, COLS - 1), f"{curr:.4f}%", FONT_EN, 11, bold=True)

        # Borders
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, top=BORDER_THIN, bottom=BORDER_THIN,
                                      left=BORDER_THIN, right=BORDER_THIN)
        for c in range(COLS):
            self._set_cell_border(table.cell(0, c), top=BORDER_THICK,
                                  left=BORDER_THIN, right=BORDER_THIN, bottom=BORDER_THIN)
            self._set_cell_border(table.cell(3, c), bottom=BORDER_THICK,
                                  left=BORDER_THIN, right=BORDER_THIN, top=BORDER_THIN)
        for r in range(4):
            self._set_cell_border(table.cell(r, 0),         left=BORDER_THICK,
                                  top=BORDER_THIN, right=BORDER_THIN, bottom=BORDER_THIN)
            self._set_cell_border(table.cell(r, COLS - 1),  right=BORDER_THICK,
                                  top=BORDER_THIN, left=BORDER_THIN,  bottom=BORDER_THIN)

    def _add_analysis_box(self, doc, quarter, year, stats, charts, history=None):
        history = history or []
        """Analysis box: header + analysis text + Chart 1 (trend)."""
        summary      = stats.get('summary', {})
        error_rate   = summary.get('error_rate', 0)
        total_errors = summary.get('total_errors', 0)
        total_doses  = summary.get('total_doses', 0)
        target       = summary.get('target', 0.03)

        # Container box
        box = doc.add_table(rows=1, cols=1)
        box.autofit = False
        para_before = doc.paragraphs[-2]
        para_before.paragraph_format.space_before = Pt(0)
        para_before.paragraph_format.space_after  = Pt(0)
        box.columns[0].width = USABLE_WIDTH
        cell = box.cell(0, 0)
        self._zero_spacing(cell.paragraphs[0])
        self._set_table_cell_margins(box, top=0, bottom=10, left=20, right=20)
        self._set_cell_border(cell, top=BORDER_MEDIUM, bottom=BORDER_MEDIUM,
                              left=BORDER_MEDIUM, right=BORDER_MEDIUM)

        cell.text = ""
        for paragraph in cell.paragraphs:
            paragraph._element.getparent().remove(paragraph._element)

        # Header row
        title_tbl = cell.add_table(rows=1, cols=1)
        title_tbl.autofit = False
        title_tbl.columns[0].width = Inches(9)
        title_cell = title_tbl.cell(0, 0)
        self._set_table_cell_margins(title_tbl, top=0, bottom=0, left=20, right=20)
        self._set_cell_shading(title_cell, SHADE_HEADER)
        self._set_cell_border(title_cell, bottom=BORDER_MEDIUM)
        self._style_cell_text(title_cell, "*التحليل", FONT, 11, bold=True, rtl=True)

        for p_after in cell.paragraphs:
            p_after._element.getparent().remove(p_after._element)

        # ── prev-quarter data from history ─────────────────────────────
        _prev      = history[-1] if history else {}
        _prev_q    = (f"{_prev.get('quarter','')} {_prev.get('year','')}").strip() or "الفصل السابق"
        _prev_err  = _prev.get('total_errors', 0)
        _prev_rate = _prev.get('error_rate', 0)

        # ── cause counts ───────────────────────────────────────────────
        _causes   = stats.get('error_causes', {})
        _c_sorted = sorted(_causes.get('counts', {}).items(), key=lambda x: x[1], reverse=True)
        _c_total  = sum(v for _, v in _c_sorted) or 1
        _top_c    = _c_sorted[0][0] if _c_sorted else ""
        _top_c_p  = round(_c_sorted[0][1] / _c_total * 100) if _c_sorted else 0
        _sec_c    = _c_sorted[1][0] if len(_c_sorted) > 1 else ""
        _sec_c_p  = round(_c_sorted[1][1] / _c_total * 100) if len(_c_sorted) > 1 else 0

        # ── AI analysis (6 \n-separated lines) ────────────────────────
        analysis_text = medication_error_ai_service.analyze_summary(
            quarter=quarter, year=year,
            total_errors=total_errors, total_doses=total_doses,
            error_rate=error_rate, target=target,
            prev_quarter=_prev_q, prev_errors=_prev_err, prev_rate=_prev_rate,
            ncc_merp=stats.get('ncc_merp', {}).get('counts', {}),
            error_cycle=stats.get('error_cycle', {}).get('counts', {}),
            detected_by=stats.get('detected_by', {}).get('counts', {}),
            staff_involved=stats.get('staff_involved', {}).get('counts', {}),
            cause_counts=stats.get('error_causes', {}).get('counts', {}),
            top_cause=_top_c, top_cause_pct=_top_c_p,
            second_cause=_sec_c, second_cause_pct=_sec_c_p,
        )

        # ── Render each line with BiDi-safe split runs ─────────────────
        _rate_str = f"{error_rate:.2f}%"
        for _li, _line in enumerate(analysis_text.split('\n')):
            _line = _line.strip()
            if _line:
                self._add_bidi_para(cell, _line, FONT_ANALYSIS, 11,
                                    bold_values=[_rate_str] if _li == 0 else None)

        # Chart title — bold, RTL
        p_ct = cell.add_paragraph()
        self._set_rtl(p_ct)
        self._zero_spacing(p_ct)
        r_ct = p_ct.add_run("مقارنة نتيجة المؤشر مع الفصول السابقة على الشكل التالي :")
        r_ct.font.name = FONT_ANALYSIS
        r_ct.font.size = Pt(11)
        r_ct.bold = True

        # Chart 1: trend
        if 'chart1_trend' in charts:
            pimg = cell.add_paragraph()
            pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pimg.add_run().add_picture(charts['chart1_trend'], width=Inches(6.8))

        # ── Post-chart1 AI text ───────────────────────────────────────
        _rates4    = [h.get('error_rate', 0) for h in history[-4:]] if len(history) >= 4 else []
        _trend_dir = 'improving' if (len(_rates4) >= 2 and _rates4[-1] < _rates4[0]) else 'worsening'
        _post_text = medication_error_ai_service.analyze_trend_post(
            quarter=quarter, year=year,
            total_errors=total_errors, error_rate=error_rate,
            prev_quarter=_prev_q, prev_errors=_prev_err, prev_rate=_prev_rate,
            trend_direction=_trend_dir,
        )
        self._add_bidi_para(cell, _post_text, FONT_ANALYSIS, 11)

        for p in cell.paragraphs:
            self._zero_spacing(p)

    # =========================================================================
    # PAGE 2: Chart 2 + NCC MERP + WHO Process Stages
    # =========================================================================

    def _build_page2(self, doc, charts, stats, history=None):
        history = history or []
        _prev      = history[-1] if history else {}
        _prev_q    = (f"{_prev.get('quarter','')} {_prev.get('year','')}").strip() or "الفصل السابق"
        _prev_err  = _prev.get('total_errors', 0)
        doc.add_page_break()
        summary      = stats.get('summary', {})
        quarter      = summary.get('quarter', 'الفصل الثالث')
        year         = str(summary.get('year', '2025'))
        total_errors = summary.get('total_errors', 0)

        # ── Chart 2: ME count comparison ─────────────────────────────────
        self._add_rtl_para(doc, 'مقارنة عدد ME خلال الفصول',
                           11, bold=True, font=FONT_ANALYSIS)
        if 'chart2_comparison' in charts:
            logger.info("Embedding Chart 2: ME count comparison...")
            self._add_centered_chart(doc, charts['chart2_comparison'], width=6.5)
            logger.info("Chart 2 embedded")

        # ── Post-chart2: AI comparison ───────────────────────────────
        _cmp = medication_error_ai_service.analyze_comparison(
            quarter=quarter, year=year, total_errors=total_errors,
            prev_quarter=_prev_q, prev_errors=_prev_err,
            has_prescribing_method_change=True,
        )
        for _cl in _cmp.split('\n'):
            _cl = _cl.strip()
            if _cl:
                self._add_bidi_para(doc, _cl, FONT_ANALYSIS, 11)

        self._add_spacer(doc, space_before=6)

        # ── NCC MERP section ──────────────────────────────────────────────
        self._add_rtl_para(doc, 'توزيع ME بحسب تصنيف NCC MERP*',
                           11, bold=True, font=FONT_ANALYSIS)

        self._add_bidi_para(doc,
            f"بحسب المجلس الوطني للابلاغ عن اخطاء الادوية و الوقاية منها (NCC MERP) "
            f"و التي صنفت اخطاء الادوية بناءً على نتائجها بدءاً من عدم وجود خطأ وصولاً "
            f"الى الضرر المميت و تشمل 6 تصنيفات و بناءً عليه فقد توزع عدد ME في {quarter} "
            f"بحسب هذه التصنيفات ضمن الجدول ادناه :",
            FONT_ANALYSIS, 11)

        self._add_ncc_merp_table(doc, stats)

        # NCC MERP footnote title (static)
        self._add_spacer(doc, space_before=4)
        self._add_bidi_para(doc,
            "NCC MERP* : (National Coordination Council for Medication Error Reporting and Prevention)",
            FONT_ANALYSIS, 10)

        _ncc_ai = medication_error_ai_service.analyze_ncc_merp(
            quarter=quarter,
            total_errors=stats.get('summary',{}).get('total_errors', 0),
            ncc_counts=stats.get('ncc_merp',{}).get('counts', {}),
            ncc_pcts=stats.get('ncc_merp',{}).get('percentages', {}),
        )
        self._add_bidi_para(doc, _ncc_ai, FONT_ANALYSIS, 11)

        self._add_spacer(doc, space_before=6)

        # ── WHO Process Stages section ────────────────────────────────────
        self._add_rtl_para(doc, 'توزيع ME بحسب Process Stages',
                           11, bold=True, font=FONT_ANALYSIS)

        self._add_bidi_para(doc,
            "تشير الابحاث الصادرة عن منظمة الصحة العالمية ان معظم اخطاء الادوية تحدث خلال "
            "مرحلة اعطاء الدواء Administration و تتوزع الاخطاء بحسب الجدول الاتي :",
            FONT_ANALYSIS, 11)

        self._add_who_stages_table(doc)

        # WHO footnote text (static)
        self._add_bidi_para(doc,
            f"هذا الجدول هو نموذج و ليس Benchmark و لكن لاعطاء فكرة عن توزع الاخطاء "
            f"بحسب Stages. اما حول توزيع الاخطاء في {quarter} من العام {year} فهي على الشكل التالي :",
            FONT_ANALYSIS, 11)

    # =========================================================================
    # NCC MERP TABLE
    # =========================================================================

    def _add_ncc_merp_table(self, doc, stats=None):
        """
        NCC MERP classification table — 7 rows × 5 cols (RTL column order).
        Columns (left→right in storage, right→left visually in RTL):
          0  NCC MERP category   (rightmost — first read in RTL)
          1  definition
          2  NCC MERP Index      (rows B/C/D merged)
          3  عدد ME              ← dynamic from stats['ncc_merp']
          4  نسبة ME لكل تصنيف   ← dynamic from stats['ncc_merp']
        """
        # Fixed template: category key, display definition, NCC MERP index label
        NCC_TEMPLATE = [
            ('Category A', 'Circumstance',       'No error'),
            ('Category B', 'No Reach-Near miss', 'Error, no Harm'),
            ('Category C', 'No Harm',            ''),
            ('Category D', 'Monitoring',         ''),
            ('Category E', 'Temporary Harm',     'Error, Harm'),
            ('Category F', 'Harm/death',         'Error, death'),
        ]

        # Pull dynamic counts & percentages from stats
        ncc_stats   = (stats or {}).get('ncc_merp', {})
        ncc_counts  = ncc_stats.get('counts', {})
        ncc_pcts    = ncc_stats.get('percentages', {})

        # Build final data rows: (category, definition, index, count, pct)
        NCC_DATA = [
            (cat, defn, idx,
             ncc_counts.get(cat, 0),
             ncc_pcts.get(cat, 0))
            for cat, defn, idx in NCC_TEMPLATE
        ]

        NCOLS = 5
        NROWS = len(NCC_DATA) + 1   # 1 header + 6 data = 7

        table = doc.add_table(rows=NROWS, cols=NCOLS)
        table.autofit = False
        self._set_table_rtl(table)
        self._set_table_cell_margins(table, top=30, bottom=30, left=40, right=40)

        # category | definition | Index | عدد ME | نسبة ME
        col_widths = [Inches(1.1), Inches(1.5), Inches(1.6), Inches(1.2), Inches(1.47)]
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

        # Header row: category | definition | Index | عدد ME | نسبة ME
        col_headers = [
            ('NCC MERP\ncategory',                   FONT_EN, False),
            ('definition',                           FONT_EN, False),
            ('NCC MERP\nIndex',                      FONT_EN, False),
            ('عدد ME',                               FONT,    True),
            ('نسبة ME لكل تصنيف\nمن اجمالي ME',     FONT,    True),
        ]
        for j, (txt, fnt, rtl) in enumerate(col_headers):
            hc = table.cell(0, j)
            self._style_cell_text(hc, txt, fnt, 9, bold=True, rtl=rtl, align='center')
            self._set_cell_shading(hc, SHADE_HEADER)

        # Merge the NCC MERP Index cells for rows B, C, D (ri = 2, 3, 4) — col 2 stays middle
        table.cell(2, 2).merge(table.cell(4, 2))
        SKIP_IDX_COL = {3, 4}   # absorbed into merged cell

        # Data rows — column mapping: cat→0, defn→1, idx→2, count→3, pct→4
        for i, (cat, defn, idx, count, pct) in enumerate(NCC_DATA):
            ri = i + 1
            self._style_cell_text(table.cell(ri, 0), cat,             FONT_EN, 9, align='center')
            self._style_cell_text(table.cell(ri, 1), defn,            FONT_EN, 9, align='center')
            if ri not in SKIP_IDX_COL:
                self._style_cell_text(table.cell(ri, 2), idx,         FONT_EN, 9, align='center')
            self._style_cell_text(table.cell(ri, 3), str(count),      FONT_EN, 9, align='center')
            self._style_cell_text(table.cell(ri, 4), f'{pct}%',      FONT_EN, 9, align='center')

        # Borders: thin inside, thick outer
        self._apply_borders_to_table(table, BORDER_THIN)
        for c in range(NCOLS):
            self._set_cell_border(table.cell(0,       c), top=BORDER_THICK)
            self._set_cell_border(table.cell(NROWS-1, c), bottom=BORDER_THICK)
        for r in range(NROWS):
            self._set_cell_border(table.cell(r, 0),       left=BORDER_THICK)
            self._set_cell_border(table.cell(r, NCOLS-1), right=BORDER_THICK)

    # =========================================================================
    # WHO PROCESS STAGES TABLE
    # =========================================================================

    def _add_who_stages_table(self, doc):
        """Static WHO Process Stages reference table (2 cols × 6 rows, RTL).
        Col 0 (rightmost in RTL) = نسبة الاخطاء بحسب WHO
        Col 1 (leftmost in RTL)  = Stage of process
        """
        WHO_DATA = [
            ('Prescribing',    '21.3%'),
            ('Transcription',  '1.4%'),
            ('Dispensing',     '15.9%'),
            ('Administration', '54.4%'),
            ('Monitoring',     '7%'),
        ]

        NROWS = len(WHO_DATA) + 1
        NCOLS = 2

        table = doc.add_table(rows=NROWS, cols=NCOLS)
        table.autofit = False
        self._set_table_rtl(table)
        self._set_table_cell_margins(table, top=30, bottom=30, left=60, right=60)

        # col 0 = نسبة (narrow), col 1 = Stage (wider) — RTL makes col 0 appear on right
        col_widths = [Inches(2.5), Inches(4.37)]
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

        # Header — reversed: col 0 = نسبة, col 1 = Stage
        self._style_cell_text(table.cell(0, 0), 'نسبة الاخطاء بحسب  WHO',
                              FONT, 9, bold=True, rtl=True, align='center')
        self._style_cell_text(table.cell(0, 1), 'Stage of process',
                              FONT_EN, 9, bold=True, align='center')
        for j in range(NCOLS):
            self._set_cell_shading(table.cell(0, j), SHADE_HEADER)

        # Data rows — col 0 = pct, col 1 = stage
        for i, (stage, pct) in enumerate(WHO_DATA):
            ri = i + 1
            self._style_cell_text(table.cell(ri, 0), pct,   FONT_EN, 9, align='center')
            self._style_cell_text(table.cell(ri, 1), stage, FONT_EN, 9, align='center')

        # Borders: thin inside, thick outer
        self._apply_borders_to_table(table, BORDER_THIN)
        for c in range(NCOLS):
            self._set_cell_border(table.cell(0,       c), top=BORDER_THICK)
            self._set_cell_border(table.cell(NROWS-1, c), bottom=BORDER_THICK)
        for r in range(NROWS):
            self._set_cell_border(table.cell(r, 0),       left=BORDER_THICK)
            self._set_cell_border(table.cell(r, NCOLS-1), right=BORDER_THICK)

    # =========================================================================
    # PAGE 3: Chart 3 (cycle) + Chart 4 (detection) + Chart 5 (shift)
    # =========================================================================

    def _build_page3(self, doc, charts, stats=None):
        stats   = stats or {}
        quarter = stats.get('summary', {}).get('quarter', 'الفصل الثالث')
        doc.add_page_break()

        # ── Chart 3: Error cycle pie ──────────────────────────────────────
        if 'chart3_cycle_pie' in charts:
            logger.info("Embedding Chart 3: Error cycle pie...")
            self._add_centered_chart(doc, charts['chart3_cycle_pie'], width=3.3)
            logger.info("Chart 3 embedded")

        _cyc = medication_error_ai_service.analyze_error_cycle(
            quarter=quarter,
            cycle_counts=stats.get('error_cycle',{}).get('counts',{}),
            cycle_pcts=stats.get('error_cycle',{}).get('percentages',{}))
        self._add_bidi_para(doc, _cyc, FONT_ANALYSIS, 10)

        # ── Chart 4: Detection donut ──────────────────────────────────────
        self._add_rtl_para(doc, 'توزيع ME بحسب طريقة اكتشافها',
                           10, bold=True, font=FONT_ANALYSIS)

        p4_intro = doc.add_paragraph()
        self._set_rtl(p4_intro)
        self._zero_spacing(p4_intro)
        r4_intro = p4_intro.add_run(
            "تم اكتشاف الاخطاء من قبل 3 مصادر و هي الصيدلي , HN و RN و كانت على الشكل التالي :"
        )
        r4_intro.font.name = FONT_ANALYSIS
        r4_intro.font.size = Pt(10)

        if 'chart4_detection_donut' in charts:
            logger.info("Embedding Chart 4: Detection donut...")
            self._add_centered_chart(doc, charts['chart4_detection_donut'], width=3.3)
            logger.info("Chart 4 embedded")

        _det = medication_error_ai_service.analyze_detection(
            quarter=quarter,
            detected_counts=stats.get('detected_by',{}).get('counts',{}),
            detected_pcts=stats.get('detected_by',{}).get('percentages',{}))
        self._add_bidi_para(doc, _det, FONT_ANALYSIS, 10)

        # ── Chart 5: Shift donut ──────────────────────────────────────────
        self._add_rtl_para(doc, 'توزيع ME بحسب الدوام',
                           10, bold=True, font=FONT_ANALYSIS)

        if 'chart5_shift_donut' in charts:
            logger.info("Embedding Chart 5: Shift donut...")
            self._add_centered_chart(doc, charts['chart5_shift_donut'], width=3.3)
            logger.info("Chart 5 embedded")

        _sh = medication_error_ai_service.analyze_shift(
            quarter=quarter,
            shift_counts=stats.get('duty_shift',{}).get('counts',{}),
            shift_pcts=stats.get('duty_shift',{}).get('percentages',{}))
        self._add_bidi_para(doc, _sh, FONT_ANALYSIS, 10)

    # =========================================================================
    # PAGE 4: Chart 6 (staff) + Chart 7 (causes) + cross-tab table
    # =========================================================================

    def _build_page4(self, doc, charts, stats, history=None):
        history = history or []
        quarter = stats.get('summary', {}).get('quarter', 'الفصل الثالث')
        doc.add_page_break()

        # ── Chart 6: Staff involved donut ─────────────────────────────────
        self._add_rtl_para(doc, 'توزيع ME بحسب مسبب الخطأ',
                           11, bold=True, font=FONT_ANALYSIS)

        _st = medication_error_ai_service.analyze_staff(
            quarter=quarter,
            staff_counts=stats.get('staff_involved',{}).get('counts',{}),
            staff_pcts=stats.get('staff_involved',{}).get('percentages',{}))
        self._add_bidi_para(doc, _st, FONT_ANALYSIS, 11)

        if 'chart6_staff_donut' in charts:
            logger.info("Embedding Chart 6: Staff donut...")
            self._add_centered_chart(doc, charts['chart6_staff_donut'], width=3.5)
            logger.info("Chart 6 embedded")

        # ── Chart 7: Causes bars ──────────────────────────────────────────
        self._add_rtl_para(doc, 'توزيع ME بحسب السبب (Cause of error)',
                           11, bold=True, font=FONT_ANALYSIS)

        if 'chart7_causes_bars' in charts:
            logger.info("Embedding Chart 7: Causes bars...")
            self._add_centered_chart(doc, charts['chart7_causes_bars'], width=6.0)
            logger.info("Chart 7 embedded")

        _ca = medication_error_ai_service.analyze_causes(
            quarter=quarter,
            cause_counts=stats.get('error_causes',{}).get('counts',{}),
            cause_pcts=stats.get('error_causes',{}).get('percentages',{}))
        self._add_bidi_para(doc, _ca, FONT_ANALYSIS, 11)

        self._add_spacer(doc, space_before=4)

        # ── Cross-tab: Cause of Error × Stage of Process ─────────────────
        self._add_rtl_para(
            doc,
            'توزيع ME بحسب السبب (Cause of error) و Stage of process',
            11, bold=True, font=FONT_ANALYSIS
        )

        self._add_cause_stage_table(doc, stats)

    # -------------------------------------------------------------------------
    def _add_cause_stage_table(self, doc, stats=None):
        """
        Dynamic cross-tabulation: Cause of Error (rows) × Stage of Process (cols).
        Data sourced from stats['cause_stage_matrix'] computed by statistics.py.
        RTL layout — cause column appears on the right (col 0 = rightmost).
        Columns: Cause | <stage 1> | <stage 2> | ... | Total
        """
        cs          = (stats or {}).get('cause_stage_matrix', {})
        causes      = cs.get('causes', [])
        stages      = cs.get('stages', [])
        matrix      = cs.get('matrix', {})
        row_totals  = cs.get('row_totals', {})
        col_totals  = cs.get('col_totals', {})
        grand_total = cs.get('grand_total', 0)

        if not causes or not stages:
            logger.warning("cause_stage_matrix is empty — skipping table")
            return

        NCOLS = len(stages) + 2   # cause col + stage cols + Total col
        NROWS = len(causes) + 2   # header + data rows + total row

        table = doc.add_table(rows=NROWS, cols=NCOLS)
        table.autofit = False
        self._set_table_rtl(table)
        self._set_table_cell_margins(table, top=25, bottom=25, left=35, right=35)

        # ── Column widths ──
        # Cause column fixed at 2.3"; remaining width split among stage cols + Total
        cause_w     = Inches(2.3)
        remaining_w = USABLE_WIDTH - cause_w
        other_w     = remaining_w / (len(stages) + 1)  # stage cols + Total
        col_widths  = [cause_w] + [other_w] * (len(stages) + 1)
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

        # ── Header row ──
        self._style_cell_text(table.cell(0, 0), 'Cause of Error',
                              FONT_EN, 8, bold=True, align='center')
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)
        for si, stage in enumerate(stages):
            hc = table.cell(0, si + 1)
            self._style_cell_text(hc, stage, FONT_EN, 8, bold=True, align='center')
            self._set_cell_shading(hc, SHADE_HEADER)
        self._style_cell_text(table.cell(0, NCOLS - 1), 'Total',
                              FONT_EN, 8, bold=True, align='center')
        self._set_cell_shading(table.cell(0, NCOLS - 1), SHADE_HEADER)

        # ── Data rows ──
        for i, cause in enumerate(causes):
            ri = i + 1
            self._style_cell_text(table.cell(ri, 0), cause,
                                  FONT_EN, 8, align='center')
            for si, stage in enumerate(stages):
                count = matrix.get(cause, {}).get(stage, 0)
                cell_val = str(count) if count else ''
                self._style_cell_text(table.cell(ri, si + 1), cell_val,
                                      FONT_EN, 8, align='center')
            self._style_cell_text(table.cell(ri, NCOLS - 1), str(row_totals.get(cause, 0)),
                                  FONT_EN, 8, bold=True, align='center')

        # ── Total row ──
        tr = NROWS - 1
        self._style_cell_text(table.cell(tr, 0), 'Total',
                              FONT_EN, 8, bold=True, align='center')
        for si, stage in enumerate(stages):
            self._style_cell_text(table.cell(tr, si + 1), str(col_totals.get(stage, 0)),
                                  FONT_EN, 8, bold=True, align='center')
        self._style_cell_text(table.cell(tr, NCOLS - 1), str(grand_total),
                              FONT_EN, 8, bold=True, align='center')
        for j in range(NCOLS):
            self._set_cell_shading(table.cell(tr, j), SHADE_TOTAL)

        # ── Borders: thin inside, thick outer ──
        self._apply_borders_to_table(table, BORDER_THIN)
        for c in range(NCOLS):
            self._set_cell_border(table.cell(0,       c), top=BORDER_THICK)
            self._set_cell_border(table.cell(NROWS-1, c), bottom=BORDER_THICK)
        for r in range(NROWS):
            self._set_cell_border(table.cell(r, 0),       left=BORDER_THICK)
            self._set_cell_border(table.cell(r, NCOLS-1), right=BORDER_THICK)

    # =========================================================================
    # PAGE 5: Type of Error × Stage of Process cross-tab
    # =========================================================================

    def _build_page5(self, doc, stats):
        doc.add_page_break()

        self._add_rtl_para(
            doc,
            'توزيع ME بحسب نوع الخطأ Type of error و Stage of process',
            11, bold=True, font=FONT_ANALYSIS
        )

        self._add_type_stage_table(doc, stats)

        self._add_spacer(doc, space_before=4)

        self._add_rtl_para(
            doc,
            'توزيع ME بحسب الاقسام للتصنيف',
            11, bold=True, font=FONT_ANALYSIS
        )

        self._add_departments_table(doc, stats)

    # -------------------------------------------------------------------------
    def _add_type_stage_table(self, doc, stats=None):
        """
        Dynamic cross-tabulation: Type of Error (rows) × Stage of Process (cols).
        Data sourced from stats['type_stage_matrix'] computed by statistics.py.
        RTL layout — type column appears on the right (col 0 = rightmost).
        Columns: Type of Error | <stage 1> | <stage 2> | ... | Total
        """
        ts          = (stats or {}).get('type_stage_matrix', {})
        types       = ts.get('types', [])
        stages      = ts.get('stages', [])
        matrix      = ts.get('matrix', {})
        row_totals  = ts.get('row_totals', {})
        col_totals  = ts.get('col_totals', {})
        grand_total = ts.get('grand_total', 0)

        if not types or not stages:
            logger.warning("type_stage_matrix is empty — skipping table")
            return

        NCOLS = len(stages) + 2   # type col + stage cols + Total col
        NROWS = len(types) + 2    # header + data rows + total row

        table = doc.add_table(rows=NROWS, cols=NCOLS)
        table.autofit = False
        self._set_table_rtl(table)
        self._set_table_cell_margins(table, top=15, bottom=15, left=30, right=30)

        # ── Column widths ──
        type_w      = Inches(2.3)
        remaining_w = USABLE_WIDTH - type_w
        other_w     = remaining_w / (len(stages) + 1)   # stage cols + Total
        col_widths  = [type_w] + [other_w] * (len(stages) + 1)
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

        # ── Header row ──
        self._style_cell_text(table.cell(0, 0), 'Type of Error',
                              FONT_EN, 8, bold=True, align='center')
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)
        for si, stage in enumerate(stages):
            hc = table.cell(0, si + 1)
            self._style_cell_text(hc, stage, FONT_EN, 8, bold=True, align='center')
            self._set_cell_shading(hc, SHADE_HEADER)
        self._style_cell_text(table.cell(0, NCOLS - 1), 'Total',
                              FONT_EN, 8, bold=True, align='center')
        self._set_cell_shading(table.cell(0, NCOLS - 1), SHADE_HEADER)

        # ── Data rows ──
        for i, etype in enumerate(types):
            ri = i + 1
            self._style_cell_text(table.cell(ri, 0), etype,
                                  FONT_EN, 8, align='center')
            for si, stage in enumerate(stages):
                count = matrix.get(etype, {}).get(stage, 0)
                cell_val = str(count) if count else ''
                self._style_cell_text(table.cell(ri, si + 1), cell_val,
                                      FONT_EN, 8, align='center')
            self._style_cell_text(table.cell(ri, NCOLS - 1), str(row_totals.get(etype, 0)),
                                  FONT_EN, 8, bold=True, align='center')

        # ── Total row ──
        tr = NROWS - 1
        self._style_cell_text(table.cell(tr, 0), 'Total',
                              FONT_EN, 8, bold=True, align='center')
        for si, stage in enumerate(stages):
            self._style_cell_text(table.cell(tr, si + 1), str(col_totals.get(stage, 0)),
                                  FONT_EN, 8, bold=True, align='center')
        self._style_cell_text(table.cell(tr, NCOLS - 1), str(grand_total),
                              FONT_EN, 8, bold=True, align='center')
        for j in range(NCOLS):
            self._set_cell_shading(table.cell(tr, j), SHADE_TOTAL)

        # ── Borders: thin inside, thick outer ──
        self._apply_borders_to_table(table, BORDER_THIN)
        for c in range(NCOLS):
            self._set_cell_border(table.cell(0,       c), top=BORDER_THICK)
            self._set_cell_border(table.cell(NROWS-1, c), bottom=BORDER_THICK)
        for r in range(NROWS):
            self._set_cell_border(table.cell(r, 0),       left=BORDER_THICK)
            self._set_cell_border(table.cell(r, NCOLS-1), right=BORDER_THICK)

    # -------------------------------------------------------------------------
    def _add_departments_table(self, doc, stats=None):
        """
        Nursing-unit distribution — 4-col split layout to halve vertical height.
        RTL table, 4 columns:
          Col 0 (rightmost): unit  — first half
          Col 1:             count — first half
          Col 2:             unit  — second half
          Col 3 (leftmost):  count — second half
        Last row: merged Grand Total.
        """
        dept   = (stats or {}).get('departments_all', {})
        units  = dept.get('units', [])
        counts = dept.get('counts', {})
        total  = dept.get('total', 0)

        if not units:
            logger.warning("departments_all is empty — skipping table")
            return

        # Split units into two halves
        mid         = (len(units) + 1) // 2   # ceiling half → right column
        right_units = units[:mid]              # first half goes in right cols (0,1)
        left_units  = units[mid:]              # second half goes in left cols (2,3)

        NCOLS = 4
        NROWS = mid + 2   # header + right-half rows + grand-total row

        table = doc.add_table(rows=NROWS, cols=NCOLS)
        table.autofit = False
        self._set_table_rtl(table)
        self._set_table_cell_margins(table, top=15, bottom=15, left=25, right=25)

        # Column widths: each "pair" gets half the usable width
        half_w  = USABLE_WIDTH / 2
        label_w = half_w * 0.67   # unit name column
        cnt_w   = half_w * 0.33   # count column
        for row in table.rows:
            row.cells[0].width = label_w
            row.cells[1].width = cnt_w
            row.cells[2].width = label_w
            row.cells[3].width = cnt_w

        # ── Header row ──
        for j, txt in enumerate(['Row Labels', 'Sum', 'Row Labels', 'Sum']):
            self._style_cell_text(table.cell(0, j), txt,
                                  FONT_EN, 7, bold=True, align='center')
            self._set_cell_shading(table.cell(0, j), SHADE_HEADER)

        # ── Data rows ──
        for i, unit in enumerate(right_units):
            ri = i + 1
            self._style_cell_text(table.cell(ri, 0), unit,
                                  FONT_EN, 7, align='center')
            self._style_cell_text(table.cell(ri, 1), str(counts.get(unit, 0)),
                                  FONT_EN, 7, align='center')

        for i, unit in enumerate(left_units):
            ri = i + 1
            self._style_cell_text(table.cell(ri, 2), unit,
                                  FONT_EN, 7, align='center')
            self._style_cell_text(table.cell(ri, 3), str(counts.get(unit, 0)),
                                  FONT_EN, 7, align='center')
        # Fill empty cells in left half if right > left
        for i in range(len(left_units), len(right_units)):
            ri = i + 1
            self._style_cell_text(table.cell(ri, 2), '', FONT_EN, 7, align='center')
            self._style_cell_text(table.cell(ri, 3), '', FONT_EN, 7, align='center')

        # ── Grand total row — merged across all 4 cols ──
        tr = NROWS - 1
        table.cell(tr, 0).merge(table.cell(tr, 3))
        self._style_cell_text(table.cell(tr, 0),
                              f'Grand Total    {total}',
                              FONT_EN, 7, bold=True, align='center')
        self._set_cell_shading(table.cell(tr, 0), SHADE_TOTAL)

        # ── Borders ──
        self._apply_borders_to_table(table, BORDER_THIN)
        for c in range(NCOLS):
            self._set_cell_border(table.cell(0,       c), top=BORDER_THICK)
            self._set_cell_border(table.cell(NROWS-1, c), bottom=BORDER_THICK)
        for r in range(NROWS):
            self._set_cell_border(table.cell(r, 0),       left=BORDER_THICK)
            self._set_cell_border(table.cell(r, NCOLS-1), right=BORDER_THICK)
        # Divider between the two halves
        for r in range(NROWS):
            self._set_cell_border(table.cell(r, 1), right=BORDER_MEDIUM)
            self._set_cell_border(table.cell(r, 2), left=BORDER_MEDIUM)

    # =========================================================================
    # LAST PAGE: Final result, action tables, approval
    # =========================================================================

    def _build_last_page(self, doc):
        doc.add_page_break()
        self._add_spacer(doc, space_before=6)

        self._add_titled_row(doc, "النتيجة النهائية", shade=SHADE_HEADER, border=BORDER_MEDIUM)
        self._add_titled_row(doc, "نتيجة المؤشر مشجعة جدا (0.01%) فهي لم تتخطى الـ target المحدد, 0.03% كذلك لم يظهر أي حالة adverse event أو sentinel.", shade=None, border=BORDER_FINE)

        self._add_spacer(doc, space_before=6)
        self._add_checkbox_question_box(doc)
        self._add_spacer(doc, space_before=6)
        self._add_previous_actions_table(doc)
        self._add_spacer(doc, space_before=6)
        self._add_current_actions_table(doc)

        p_title = doc.add_paragraph()
        self._set_rtl(p_title)
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after  = Pt(4)
        p_title.paragraph_format.line_spacing = 1
        run_t = p_title.add_run('يجب ان تكون الاجراءات الواجب اتخاذها محددة وقابلة للقياس')
        run_t.font.name  = FONT_ANALYSIS
        run_t.font.size  = Pt(11)
        run_t.bold       = True
        run_t.underline  = True

        self._add_approval_table(doc)

        p_note = doc.add_paragraph()
        self._set_rtl(p_note)
        p_note.paragraph_format.space_before = Pt(4)
        p_note.paragraph_format.space_after  = Pt(0)
        p_note.paragraph_format.line_spacing = 1
        run_note = p_note.add_run(
            'يجب توقيع جميع المعنيين بمتابعة الاجراءات المتخذة او من ينوب عنهم في خانة الموافقة '
            'اضافة الى مدراء الادارات المعنيين'
        )
        run_note.font.name = FONT_ANALYSIS
        run_note.font.size = Pt(11)
        run_note.bold      = True

    # =========================================================================
    # LAST PAGE COMPONENTS (identical to mortality)
    # =========================================================================

    def _add_checkbox_question_box(self, doc):
        box_tbl = doc.add_table(rows=1, cols=1)
        box_tbl.autofit = False
        box_tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(box_tbl, top=30, bottom=30, left=60, right=60)

        cell = box_tbl.cell(0, 0)
        self._set_cell_border(cell, top=BORDER_FINE, bottom=BORDER_FINE,
                              left=BORDER_FINE, right=BORDER_FINE)

        p = cell.paragraphs[0]
        self._set_rtl(p)
        self._zero_spacing(p)
        self._add_run(p, 'هل تم اتخاذ اجراءات تحسينية سابقا حول موضوع التحليل :      ', FONT_ANALYSIS, 11)
        self._add_run(p, 'نعم', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐', None, 11)
        self._add_run(p, '         ', None, 11)
        self._add_run(p, 'كلا  ', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐ ', None, 11)

        p2 = cell.add_paragraph()
        self._set_rtl(p2)
        self._zero_spacing(p2)
        self._add_run(p2,
            'في حال نعم , يجب تعبئة جدول " الاجراءات المتخذة سابقا" و في حال كلا , الانتقال الى جدول "الاجراءات المتخذة الحالية"',
            FONT_ANALYSIS, 11)

    def _add_previous_actions_table(self, doc):
        table = doc.add_table(rows=3, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 4))
        self._style_cell_text(table.cell(0, 0), 'الاجراءات المتخذة  السابقة',
                              FONT, 10, bold=True, rtl=True, align='center')

        headers = ['الرقم', 'الاجراءات المقترحة سابقا', 'التاريخ المتوقع لتنفيذ الاجراء',
                   'تم التنفيذ )نعم/كلا(', 'في حال كلا ) الاجراءات التصحيحية الجديدة(']
        for j, h in enumerate(headers):
            self._style_cell_text(table.cell(1, j), h, FONT, 11, bold=True, rtl=True, align='center')

        for j in range(4):
            table.cell(2, j).text = ''

        last_cell = table.cell(2, 4)
        last_cell.text = ''
        p1 = last_cell.paragraphs[0]
        self._add_run(p1, ' ☐  الغاء الاجراء', FONT, 11)
        self._zero_spacing(p1)
        self._set_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = last_cell.add_paragraph()
        self._add_run(p2, ' ☐ اضافة الاجراء ضمن الاجراءات المتخذة الحالية.)الجدول ادناه(', FONT, 11)
        self._zero_spacing(p2)
        self._set_rtl(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(2):
            self._shade_row_cells(table.rows[i].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(1.8), Inches(1.37), Inches(1.1), Inches(2.1)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_current_actions_table(self, doc):
        table = doc.add_table(rows=6, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 3))
        self._style_cell_text(table.cell(0, 0), 'الاجراءات المتخذة  الحالية',
                              FONT, 10, bold=True, rtl=True, align='center')

        headers = ['الرقم', 'الاجراءات الواجب اتخاذها لحل الموضوع',
                   'الشخص المعني بالمتابعة', 'التاريخ المتوقع للتنفيذ']
        for j, h in enumerate(headers):
            self._style_cell_text(table.cell(1, j), h, FONT, 11, bold=True, rtl=True, align='center')

        for i in range(2, 6):
            for j in range(4):
                table.cell(i, j).text = ''

        for i in range(2):
            self._shade_row_cells(table.rows[i].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(3.07), Inches(1.8), Inches(1.5)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_approval_table(self, doc):
        table = doc.add_table(rows=4, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).text = ''
        self._style_cell_text(table.cell(0, 1), 'اعداد', FONT, 11, bold=True, rtl=True, align='center')
        table.cell(0, 2).merge(table.cell(0, 4))
        self._style_cell_text(table.cell(0, 2), 'موافقة', FONT, 11, bold=True, rtl=True, align='center')

        for i, label in enumerate(['الاسم', 'التوقيع', 'التاريخ']):
            self._style_cell_text(table.cell(i + 1, 0), label, FONT, 11, bold=True, rtl=True, align='center')
            for j in range(1, 5):
                table.cell(i + 1, j).text = ''

        self._shade_row_cells(table.rows[0].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.8), Inches(1.5), Inches(1.52), Inches(1.52), Inches(1.53)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    # =========================================================================
    # REUSABLE HELPERS (identical to mortality)
    # =========================================================================

    def _add_spacer(self, doc, space_before=0):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(space_before)
        sp.paragraph_format.space_after  = Pt(0)
        sp.paragraph_format.line_spacing = Pt(1)
        sp.add_run("").font.size = Pt(1)

    def _add_centered_chart(self, doc, chart_bytes, width=6.5):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._zero_spacing(para)
        para.add_run().add_picture(chart_bytes, width=Inches(width))


    def _add_bidi_para(self, container, text, font_name, size,
                       bold=False, bold_values=None):
        """
        Render mixed Arabic/English with correct BiDi bracket/% direction.

        ROOT CAUSE of flipped () and misplaced %:
          In a w:bidi paragraph with ONE run containing both Arabic and LTR
          text, the Unicode BiDi algorithm mirrors bracket characters and may
          reorder adjacent LTR tokens.  Solution: split the text at
          Arabic/LTR boundaries and give each segment its own run with an
          explicit direction marker at the XML level.

          LTR run (English, digits, %,  (), commas …):
            - Remove any inherited <w:rtl/> from rPr
            - Set w:rFonts ascii/hAnsi/cs = FONT_EN  (prevents Arabic shaping)

          Arabic run:
            - Append <w:rtl/> to rPr
            - Set w:rFonts to the Arabic font
        """
        # Matches any LTR token: starts with an ASCII letter or digit,
        # may contain letters, digits, spaces, and common punctuation.
        LTR_PAT = re.compile(r'[A-Za-z0-9][A-Za-z0-9 ().,;:%&/_\-+<>=*]*')

        # Build [(chunk, is_ltr)] list
        parts, cursor = [], 0
        for m in LTR_PAT.finditer(text):
            if m.start() > cursor:
                parts.append((text[cursor:m.start()], False))  # Arabic
            parts.append((m.group(), True))                    # LTR
            cursor = m.end()
        if cursor < len(text):
            parts.append((text[cursor:], False))
        if not parts:
            parts = [(text, False)]

        para = container.add_paragraph()
        self._set_rtl(para)
        self._zero_spacing(para)

        for chunk, is_ltr in parts:
            if not chunk:
                continue
            # Thin spaces around LTR tokens so Arabic words don't merge
            display = (' ' + chunk.strip() + ' ') if is_ltr else chunk
            run = para.add_run(display)
            run.font.size = Pt(size)

            _bold = bold or (bool(bold_values) and any(bv in chunk for bv in bold_values))
            if _bold:
                run.bold = True

            rPr = run._element.get_or_add_rPr()

            # Build/update w:rFonts element
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)

            if is_ltr:
                # ── LTR run ──────────────────────────────────────────────
                run.font.name = FONT_EN
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), FONT_EN)
                # Remove any w:rtl so Word treats this run as LTR
                for el in rPr.findall(qn('w:rtl')):
                    rPr.remove(el)
            else:
                # ── RTL (Arabic) run ─────────────────────────────────────
                run.font.name = font_name
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), font_name)
                # Ensure w:rtl is present
                if not rPr.findall(qn('w:rtl')):
                    rPr.append(OxmlElement('w:rtl'))


    def _add_analysis_text(self, doc, text):
        para = doc.add_paragraph()
        run  = para.add_run(text)
        run.font.name = FONT_ANALYSIS
        run.font.size = Pt(11)
        self._set_rtl(para)
        self._zero_spacing(para)

    def _add_titled_row(self, doc, text, shade=None, border=BORDER_FINE):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(tbl,
            top=0 if shade else 20, bottom=0 if shade else 20, left=20, right=20)
        cell = tbl.cell(0, 0)
        if shade:
            self._set_cell_shading(cell, shade)
        self._set_cell_border(cell, top=border, bottom=border, left=border, right=border)
        self._style_cell_text(cell, text, FONT, 11, bold=True, rtl=True, align='center')

    def _add_rtl_para(self, doc, text, size, bold=False, font=None):
        para = doc.add_paragraph()
        run  = para.add_run(text)
        run.font.name = font or FONT
        run.font.size = Pt(size)
        run.bold      = bold
        self._set_rtl(para)
        self._zero_spacing(para)

    def _add_run(self, para, text, font_name, size, bold=False):
        run = para.add_run(text)
        if font_name:
            run.font.name = font_name
        run.font.size = Pt(size)
        run.bold      = bold
        return run

    def _zero_spacing(self, para):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing = 1

    def _set_rtl(self, para):
        para.paragraph_format.right_to_left = True
        pPr  = para._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def _set_table_rtl(self, table):
        bidi = OxmlElement('w:bidiVisual')
        bidi.set(qn('w:val'), '1')
        table._tbl.tblPr.append(bidi)

    def _shade_row_cells(self, cells, color):
        for cell in cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            shading.set(qn('w:val'),  'clear')
            cell._element.get_or_add_tcPr().append(shading)

    def _apply_borders_to_table(self, table, border_style):
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, top=border_style, bottom=border_style,
                                      left=border_style, right=border_style)

    # =========================================================================
    # LOW-LEVEL XML HELPERS (identical to mortality)
    # =========================================================================

    def _set_table_cell_margins(self, table, top=50, bottom=50, left=80, right=80):
        tblCellMar = OxmlElement('w:tblCellMar')
        for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'),    str(val))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)
        table._tbl.tblPr.append(tblCellMar)

    def _set_cell_shading(self, cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tcPr.append(shd)

    def _set_cell_border(self, cell, **kwargs):
        tc       = cell._element
        tcPr     = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                tag     = f'w:{edge}'
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                for k, v in kwargs[edge].items():
                    element.set(qn('w:' + k), str(v))

    def _set_tc_border(self, tc, **kwargs):
        tcPr      = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                tag     = f'w:{edge}'
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                for k, v in kwargs[edge].items():
                    element.set(qn('w:' + k), str(v))

    def _style_cell_text(self, cell, text, font_name, font_size,
                         bold=False, rtl=False, align='center'):
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(font_size)
        r.bold      = bold
        self._zero_spacing(p)

        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if rtl:
            self._set_rtl(p)


# =============================================================================
# SINGLETON
# =============================================================================
medication_error_docx_generator = MedicationErrorDocxGenerator()

# Alias for routes.py
matplotlib_docx_generator = medication_error_docx_generator