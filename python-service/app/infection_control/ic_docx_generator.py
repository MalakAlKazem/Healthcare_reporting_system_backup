"""
Shared DOCX Report Generator for CLABSI and CAUTI infection control modules.

Report structure (mirrors VAP report):
  Page 1   : Header metadata table · Results table (last 6 quarters) ·
             Analysis box · Floor comparison table
  Pages 2+ : Per-floor sections (dynamic — only floors with data):
               • Trend chart + static analysis
               • Germs chart + static analysis
               • Cases table + static analysis
  Last page: Final result · Checkbox question · Previous/current action tables ·
             Approval table

Config keys (passed at construction):
    indicator         : 'CLABSI' or 'CAUTI'
    indicator_ar      : Arabic abbreviation shown in headings
    days_key          : 'catheter_days' or 'urinary_catheter_days'
    days_label_ar     : Arabic column header for catheter days
    insertion_col     : Excel column name for insertion date
    insertion_col_ar  : Arabic label for that column
    metadata_topic_ar : Arabic analysis topic string for the metadata table
    file_prefix       : 'CLABSI' or 'CAUTI'
"""

import io
import os
import re
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from loguru import logger
except ImportError:
    import logging
    logger = logging.getLogger(__name__)

from app.infection_control.ic_chart_generator import InfectionControlChartGenerator
from app.infection_control.ic_statistics import _match_floor

try:
    from app.infection_control.ic_ai_service import ic_ai_service
except Exception:
    ic_ai_service = None  # graceful degradation — fallbacks used

# ── Fonts ─────────────────────────────────────────────────────────────────────
FONT          = 'Traditional Arabic'
FONT_EN       = 'Calibri'
FONT_ANALYSIS = 'Frutiger LT Arabic 45 Light'

# ── Layout ────────────────────────────────────────────────────────────────────
PAGE_WIDTH   = Inches(8.27)
PAGE_HEIGHT  = Inches(11.69)
MARGIN       = Inches(0.7)
USABLE_WIDTH = Inches(6.87)

# ── Borders ───────────────────────────────────────────────────────────────────
BORDER_THICK  = {'val': 'single', 'sz': '18', 'color': '000000', 'space': '0'}
BORDER_MEDIUM = {'val': 'single', 'sz': '12', 'color': '000000', 'space': '0'}
BORDER_THIN   = {'val': 'single', 'sz': '8',  'color': '000000', 'space': '0'}
BORDER_FINE   = {'val': 'single', 'sz': '4',  'color': '000000', 'space': '0'}
BORDER_NONE   = {'val': 'none',   'sz': '0',  'color': 'FFFFFF', 'space': '0'}

# ── Shading ───────────────────────────────────────────────────────────────────
SHADE_HEADER = 'EDEBE3'
SHADE_TABLE  = 'D9E2F3'
SHADE_TOTAL  = 'F2F2F2'

# ── Quarter names ─────────────────────────────────────────────────────────────
QUARTER_AR = {
    1: 'الفصل الأول',  2: 'الفصل الثاني',
    3: 'الفصل الثالث', 4: 'الفصل الرابع',
    # with hamza
    'الفصل الأول':  'الفصل الأول',  'الفصل الثاني':  'الفصل الثاني',
    'الفصل الثالث': 'الفصل الثالث', 'الفصل الرابع':  'الفصل الرابع',
    # without hamza (stored in history JSON)
    'الفصل الاول':  'الفصل الأول',
}

# ── Reports dir ───────────────────────────────────────────────────────────────
_THIS_DIR    = os.path.dirname(os.path.abspath(__file__))
_SERVICE_DIR = os.path.normpath(os.path.join(_THIS_DIR, '..', '..'))
_REPORTS_DIR = os.path.join(_SERVICE_DIR, 'storage', 'reports')
_LOGO_PATH   = os.path.join(_SERVICE_DIR, 'assets', 'LOGO.png')


def _q(q) -> str:
    """Convert quarter (int or Arabic string) to Arabic full string."""
    return QUARTER_AR.get(q, str(q))


class InfectionControlDocxGenerator:
    """
    Full-featured Arabic RTL DOCX generator for CLABSI and CAUTI.

    Mirrors VAP report structure but:
      - Floors are dynamic (only floors present in data get per-floor pages)
      - Configurable indicator name, days column, Arabic labels
      - Static analysis text placeholders (AI integration point ready)

    Usage:
        gen    = InfectionControlDocxGenerator(config=CLABSI_CONFIG)
        result = gen.generate_report(history=all_records, targets=CLABSI_TARGETS)
    """

    # ── Type configs ──────────────────────────────────────────────────────────
    CONFIGS = {
        'clabsi': {
            'indicator':         'CLABSI',
            'indicator_ar':      'CLABSI',
            'days_key':          'catheter_days',
            'days_label_ar':     'أيام القسطرة',
            'insertion_col':     'Date of insertion Central line',
            'insertion_col_ar':  'تاريخ إدخال القسطرة المركزية',
            'metadata_topic_ar': 'تحليل بيانات مؤشر ال CLABSI',
            'file_prefix':       'CLABSI',
        },
        'cauti': {
            'indicator':         'CAUTI',
            'indicator_ar':      'CAUTI',
            'days_key':          'urinary_catheter_days',
            'days_label_ar':     'أيام القسطرة البولية',
            'insertion_col':     'Date of foley insertion',
            'insertion_col_ar':  'تاريخ إدخال قسطرة البول',
            'metadata_topic_ar': 'تحليل بيانات مؤشر الCAUTI',
            'file_prefix':       'CAUTI',
        },
        'vap': {
            'indicator':         'VAP',
            'indicator_ar':      'VAP',
            'days_key':          'ventilator_days',
            'days_label_ar':     'أيام التنفس الاصطناعي',
            'insertion_col':     'Date of intubation',
            'insertion_col_ar':  'تاريخ التنبيب',
            'metadata_topic_ar': 'تحليل بيانات مؤشر ال VAP',
            'file_prefix':       'VAP',
        },
    }

    def __init__(self, infection_type: str):
        if infection_type not in self.CONFIGS:
            raise ValueError(f"Unknown type '{infection_type}'. Use 'clabsi' or 'cauti'.")
        self.cfg        = self.CONFIGS[infection_type]
        self.logo_path  = _LOGO_PATH

    # =========================================================================
    # PUBLIC API
    # =========================================================================

    def generate_report(self, history: list, targets: dict, current: dict = None) -> dict:
        """
        Generate the full Arabic DOCX infection control report.

        Args:
            history : All history records for this type (list, oldest first).
                      Each entry has: year, quarter, summary.
                      No cases or germs_distribution stored here.
            targets : {dept: target_rate} dict
            current : Separate dict with cases and germs_distribution for the
                      latest quarter (loaded from *_current.json).

        Returns:
            {'filePath': str, 'fileName': str}
        """
        cfg = self.cfg
        if not history:
            raise ValueError("No history data — cannot generate report.")

        latest  = history[-1]
        quarter = _q(latest.get('quarter', ''))
        year    = str(latest.get('year', datetime.now().year))

        logger.info(f"[{cfg['indicator']}] Starting report generation — {quarter} {year}")

        # Generate charts
        logger.info(f"[{cfg['indicator']}] Generating charts…")
        chart_gen = InfectionControlChartGenerator(cfg['indicator'], cfg['days_key'])
        charts    = chart_gen.generate_all_charts(history, targets)
        logger.info(f"[{cfg['indicator']}] {len(charts)} charts ready")

        # Build document
        doc = Document()
        self._setup_section(doc.sections[0], quarter, year)

        # Normalise history for tables
        norm_history = [self._normalise(e, targets) for e in history]
        norm_current = norm_history[-1]

        # Inject raw cases from the separate current file (germs come from history entry)
        if current:
            norm_current['cases'] = current.get('cases', [])

        # Only include floors that have data in at least one history quarter
        # (avoids empty placeholder floors bloating the comparison table)
        floors_with_any_data = {
            f for e in norm_history
            for f, d in e.get('floors', {}).items()
            if (d.get('cases') or 0) > 0
        }
        floors = [f for f in norm_current['floors'].keys() if f in floors_with_any_data]
        active_floors = [f for f in floors if norm_current.get('floors', {}).get(f, {}).get('cases', 0) > 0]
        logger.info(f"[{cfg['indicator']}] Active floors with cases: {active_floors}")

        logger.info(f"[{cfg['indicator']}] Building page 1 (summary)…")
        self._build_page1(doc, quarter, year, norm_current, norm_history, floors, targets)

        logger.info(f"[{cfg['indicator']}] Building floor pages ({len(active_floors)} floors)…")
        self._build_floor_pages(doc, quarter, year, norm_current, norm_history, floors, charts)

        logger.info(f"[{cfg['indicator']}] Building last page…")
        self._build_last_page(doc)

        # Save
        os.makedirs(_REPORTS_DIR, exist_ok=True)
        file_name = f"{cfg['file_prefix']} report {quarter} {year}.docx"
        file_path = os.path.join(_REPORTS_DIR, file_name)
        try:
            doc.save(file_path)
        except PermissionError:
            ts        = datetime.now().strftime('%H%M%S')
            file_name = f"{cfg['file_prefix']} report {quarter} {year}_{ts}.docx"
            file_path = os.path.join(_REPORTS_DIR, file_name)
            doc.save(file_path)
            logger.warning(f"[{cfg['indicator']}] File locked — saved as {file_name}")

        logger.info(f"[{cfg['indicator']}] Report saved: {file_path}")

        # Release LLM memory after report is done
        if ic_ai_service:
            ic_ai_service.unload()
            logger.info(f"[{cfg['indicator']}] AI model unloaded from memory")

        return {'filePath': file_path, 'fileName': file_name}

    # =========================================================================
    # DATA NORMALISATION
    # =========================================================================

    def _normalise(self, entry: dict, targets: dict) -> dict:
        """Convert raw history record to internal format for the generator.

        All indicators (CLABSI, CAUTI, VAP) use the same 'summary' key format:
          entry['summary'][floor] = {cases, <days_key>, rate}
          entry['germs_distribution'][floor] = {germ: count}
          entry['cases'] = [{floor, age, germs, risk_factors, ...}]
        """
        cfg     = self.cfg
        quarter = _q(entry.get('quarter', ''))
        year    = str(entry.get('year', ''))

        # ── CLABSI / CAUTI / VAP — all use 'summary' key ─────────────────────
        summary = entry.get('summary', {})
        total_cases = sum(v.get('cases') or 0 for v in summary.values())
        total_days  = sum(v.get(cfg['days_key']) or 0 for v in summary.values())
        overall_rate = round((total_cases / total_days) * 1000, 2) if total_days > 0 else 0.0

        floors = {}
        for dept, data in summary.items():
            floors[dept] = {
                'cases':  data.get('cases') or 0,
                'days':   data.get(cfg['days_key']) or 0,
                'rate':   data.get('rate') or 0.0,
                'target': targets.get(dept, 0.0),
            }

        return {
            'quarter':            quarter,
            'year':               year,
            'total_cases':        total_cases,
            'total_days':         total_days,
            'overall_rate':       overall_rate,
            'floors':             floors,
            'germs_distribution': entry.get('germs_distribution', {}),
            'cases':              entry.get('cases', []),
        }

    # =========================================================================
    # SECTION SETUP (header, footer, border)
    # =========================================================================

    def _setup_section(self, section, quarter, year):
        section.page_width    = PAGE_WIDTH
        section.page_height   = PAGE_HEIGHT
        section.top_margin    = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin   = MARGIN
        section.right_margin  = MARGIN

        sectPr    = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '12')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

        self._add_header(section)
        self._add_footer(section)

    def _add_header(self, section):
        header = section.header
        for el in list(header._element):
            header._element.remove(el)

        table = header.add_table(rows=1, cols=3, width=Inches(7))
        self._set_table_cell_margins(table, top=0, bottom=0, left=0, right=0)
        cells = table.rows[0].cells
        cells[0].width = Inches(1.0)
        cells[1].width = Inches(5.0)
        cells[2].width = Inches(1.0)

        if os.path.exists(self.logo_path):
            para = cells[0].paragraphs[0]
            para.add_run().add_picture(self.logo_path, width=Inches(0.5))
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._zero_spacing(para)

        para = cells[1].paragraphs[0]
        run  = para.add_run('نموذج تحليل البيانات')
        run.font.name = FONT_ANALYSIS
        run.font.size = Pt(16)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(para)
        self._zero_spacing(para)

        para = cells[2].paragraphs[0]
        run  = para.add_run(f'QS-36F-103(4)\n{datetime.now().strftime("%d/%m/%Y")}')
        run.font.name = FONT_EN
        run.font.size = Pt(7)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._zero_spacing(para)

        line_para = header.add_paragraph()
        line_para.paragraph_format.space_before = Pt(0)
        line_para.paragraph_format.space_after  = Pt(8)
        line_para.paragraph_format.line_spacing = Pt(1)
        line_para.add_run('').font.size = Pt(1)
        pPr    = line_para._element.get_or_add_pPr()
        pBdr   = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_footer(self, section):
        footer = section.footer
        for el in list(footer._element):
            footer._element.remove(el)

        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run()

        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fc1)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        run._element.append(instr)
        fc2 = OxmlElement('w:fldChar')
        fc2.set(qn('w:fldCharType'), 'end')
        run._element.append(fc2)
        run.font.size = Pt(10)
        run.font.name = FONT_EN

    # =========================================================================
    # PAGE 1
    # =========================================================================

    def _build_page1(self, doc, quarter, year, current, norm_history, floors, targets):
        self._add_metadata_table(doc, quarter, year)
        self._add_mini_spacer(doc)
        self._add_results_table(doc, current, norm_history, floors)
        self._add_mini_spacer(doc)
        self._add_analysis_box(doc, quarter, year, current)
        self._add_mini_spacer(doc)
        self._add_floor_comparison_table(doc, quarter, year, current, norm_history, floors, targets)
        self._add_page_break(doc)

    def _add_mini_spacer(self, doc):
        """Minimal paragraph spacer between page-1 tables."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'),  '0')
        spacing.set(qn('w:line'),   '120')   # ~6pt line height
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)

    # ── Metadata table ────────────────────────────────────────────────────────

    def _add_metadata_table(self, doc, quarter, year):
        cfg   = self.cfg
        table = doc.add_table(rows=3, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=0, bottom=0, left=30, right=30)

        col_w = [Inches(1.93), Inches(1.06), Inches(2.71), Inches(1.17)]
        for i, w in enumerate(col_w):
            for row in table.rows:
                row.cells[i].width = w

        # Row heights for metadata table — generous minimum so content is never cut
        for ri in range(3):
            table.rows[ri].height      = Inches(0.35)
            table.rows[ri].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        # Row 0
        self._style_cell(table.cell(0, 3), 'موضوع التحليل', FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)
        self._style_cell(table.cell(0, 2), cfg['metadata_topic_ar'], FONT_EN, 8, bold=True, rtl=True)
        self._style_cell(table.cell(0, 1), 'الإدارة', FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 1), SHADE_HEADER)
        self._style_cell(table.cell(0, 0), 'الجودة والسلامة', FONT, 10, rtl=True)

        # Row 1
        self._style_cell(table.cell(1, 3), '', FONT, 9, rtl=True)
        self._set_cell_shading(table.cell(1, 3), SHADE_HEADER)
        self._style_cell(table.cell(1, 2), '', FONT_EN, 9)
        self._style_cell(table.cell(1, 1), 'الوحدة الإدارية', FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, 1), SHADE_HEADER)
        self._style_cell(table.cell(1, 0), 'ضبط العدوى', FONT, 9, rtl=True)

        # Row 2 — condensed to single lines to minimise row height
        self._style_cell(
            table.cell(2, 3),
            'مصادر البيانات( تحديد اسماء النماذج)',
            FONT, 9, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(2, 3), SHADE_HEADER)
        self._style_cell(
            table.cell(2, 2),
            'Health care associated infection )HCAI( sheet',
            FONT_EN, 8, rtl=True
        )
        self._style_cell(
            table.cell(2, 1),
            'الشهر / الفصل / العام',
            FONT, 9, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(2, 1), SHADE_HEADER)
        self._style_cell(table.cell(2, 0), f'{quarter} / {year}', FONT, 10, rtl=True)

        # Vertical merges for cols 2 and 3 rows 0-1
        table.cell(0, 2).merge(table.cell(1, 2))
        table.cell(0, 3).merge(table.cell(1, 3))
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)

        # Borders — all sides thick
        for r in range(3):
            for c in range(4):
                self._set_cell_border(table.cell(r, c),
                                      top=BORDER_THICK, bottom=BORDER_THICK,
                                      left=BORDER_THICK, right=BORDER_THICK)

        # Cols 2-3 in row 1 are vMerge continuation cells.
        # table.cell(1, 2/3) resolves to the anchor (row 0) so the loop above
        # never touches the raw w:tc elements in row 1 — set their borders directly.
        row1_tcs = table.rows[1]._tr.findall(qn('w:tc'))
        for tc in row1_tcs[2:4]:
            self._set_tc_border(tc,
                                top=BORDER_THICK, bottom=BORDER_THICK,
                                left=BORDER_THICK, right=BORDER_THICK)

    # ── Results table (last 6 quarters) ──────────────────────────────────────

    def _add_results_table(self, doc, current, norm_history, floors):
        COLS   = 6
        N_HIST = COLS - 1

        last5 = norm_history[:-1][-(N_HIST):]  # exclude current quarter
        while len(last5) < N_HIST:
            last5 = [{}] + last5

        table = doc.add_table(rows=4, cols=COLS)
        table.autofit = False
        self._set_table_cell_margins(table, top=10, bottom=10, left=30, right=30)

        w_col = Inches(6.87 / COLS)
        for i in range(COLS):
            for row in table.rows:
                row.cells[i].width = w_col

        # Compact row heights for results table
        table.rows[0].height      = Inches(0.20)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[1].height      = Inches(0.28)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[2].height      = Inches(0.20)
        table.rows[2].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[3].height      = Inches(0.35)
        table.rows[3].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        # Row 0: merged title
        c0 = table.cell(0, 0)
        for c in range(1, COLS):
            c0 = c0.merge(table.cell(0, c))
        self._style_cell(table.cell(0, 0), 'النتائج', FONT, 11, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)

        # Row 1: instruction + current label
        instr = table.cell(1, 0)
        for c in range(1, N_HIST):
            instr = instr.merge(table.cell(1, c))
        self._style_cell(
            table.cell(1, 0),
            'في حال كان هناك نتائج سابقة عن الموضوع الذي يتم تحليله يجب ذكره مع تحديد الفترة: الشهر / الفصل / العام',
            FONT, 8, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(1, 0), SHADE_HEADER)
        self._style_cell(table.cell(1, N_HIST), 'النتيجة الحالية', FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, N_HIST), SHADE_TABLE)

        # Row 2: quarter labels
        for i, entry in enumerate(last5):
            lbl = f"{entry.get('quarter', '')} {entry.get('year', '')}".strip() if entry else ''
            self._style_cell(table.cell(2, i), lbl, FONT, 9, bold=True, rtl=True)
            self._set_cell_shading(table.cell(2, i), SHADE_HEADER)
        self._style_cell(
            table.cell(2, N_HIST),
            f"{current.get('quarter', '')} {current.get('year', '')}",
            FONT, 10, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(2, N_HIST), SHADE_TABLE)

        # Row 3: rates per floor per quarter
        self._fill_rates_cell(table.cell(3, N_HIST), current.get('floors', {}))
        for i, entry in enumerate(last5):
            if entry and entry.get('floors'):
                self._fill_rates_cell(table.cell(3, i), entry['floors'])

        # Borders — all sides thick
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell,
                    top=BORDER_THICK, bottom=BORDER_THICK,
                    left=BORDER_THICK, right=BORDER_THICK)

    def _fill_rates_cell(self, cell, floors: dict):
        """Fill a results-table cell with floor=rate pairs."""
        pairs = []
        for dept, fd in floors.items():
            try:
                rate = float(fd.get('rate', 0.0) if fd else 0.0)
            except (TypeError, ValueError):
                rate = 0.0
            rs = f'{rate:.2f}'.rstrip('0').rstrip('.') if rate else '0'
            pairs.append(f'{dept} = {rs}‰')

        cell.paragraphs[0].clear()
        for i in range(0, len(pairs), 3):
            chunk = '  \n  '.join(pairs[i:i + 3])
            p     = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(chunk)
            r.font.name = FONT_EN
            r.font.size = Pt(7)
            r.bold      = True

    # ── Analysis box ──────────────────────────────────────────────────────────

    def _add_analysis_box(self, doc, quarter, year, current):
        cfg = self.cfg
        box = doc.add_table(rows=1, cols=1)
        box.autofit = False
        box.columns[0].width = USABLE_WIDTH
        cell = box.cell(0, 0)
        self._zero_spacing(cell.paragraphs[0])
        self._set_table_cell_margins(box, top=0, bottom=0, left=20, right=20)
        self._set_cell_border(cell, top=BORDER_MEDIUM, left=BORDER_THICK, right=BORDER_THICK)

        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)

        # Header row inside box
        title_tbl  = cell.add_table(rows=1, cols=1)
        title_tbl.autofit = False
        title_tbl.columns[0].width = USABLE_WIDTH
        title_cell = title_tbl.cell(0, 0)
        self._set_table_cell_margins(title_tbl, top=0, bottom=0, left=20, right=20)
        self._set_cell_shading(title_cell, SHADE_HEADER)
        self._set_cell_border(title_cell, bottom=BORDER_MEDIUM)
        self._style_cell(title_cell, '*التحليل', FONT, 11, bold=True, rtl=True)

        for p in cell.paragraphs:
            self._zero_spacing(p)

    # ── Floor comparison table ────────────────────────────────────────────────

    def _add_floor_comparison_table(self, doc, quarter, year, current, norm_history, floors, targets):
        cfg = self.cfg

        recent_hist = norm_history[:-1][-4:]  # exclude current quarter, last 4 previous
        all_qy = [(e.get('quarter', ''), e.get('year', '')) for e in recent_hist]
        all_qy.append((quarter, year))

        desc       = self._describe_quarters(all_qy)
        title_text = f"جدول يظهر مقارنة نتائج ال {cfg['indicator_ar']} بين {desc} بحسب الأقسام:"

        tp = doc.add_paragraph()
        self._set_rtl(tp)
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after  = Pt(0)
        tr = tp.add_run(title_text)
        tr.font.name = FONT
        tr.font.size = Pt(10)
        tr.bold      = True
        tr.underline = True
        self._set_run_rtl(tr)

        N_FL       = len(floors)
        TOTAL_COLS = 1 + N_FL * 3
        n_data     = len(recent_hist) + 1
        TOTAL_ROWS = 2 + n_data

        table = doc.add_table(rows=TOTAL_ROWS, cols=TOTAL_COLS)
        table.autofit = False
        self._set_table_cell_margins(table, top=8, bottom=8, left=20, right=20)

        # Set full-width table
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(USABLE_WIDTH / 914400 * 1440)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # Column widths — computed dynamically so total == USABLE_WIDTH
        USABLE_IN  = 6.87          # inches
        Q_IN       = 0.75          # quarter label column
        floor_in   = (USABLE_IN - Q_IN) / N_FL   # per-floor budget (3 sub-cols)
        # Sub-column split ratio 5:4:5
        rate_in    = floor_in * 5 / 14
        cases_in   = floor_in * 4 / 14
        days_in    = floor_in * 5 / 14
        q_w        = Inches(Q_IN)
        rate_w     = Inches(rate_in)
        cases_w    = Inches(cases_in)
        days_w     = Inches(days_in)

        for row in table.rows:
            row.cells[0].width = q_w
            for fi in range(N_FL):
                b = 1 + fi * 3
                row.cells[b].width     = rate_w
                row.cells[b + 1].width = cases_w
                row.cells[b + 2].width = days_w

        # Compact row heights — keep sub-header tall for vertical text
        table.rows[0].height      = Inches(0.28)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for ri in range(2, TOTAL_ROWS):
            table.rows[ri].height      = Inches(0.22)
            table.rows[ri].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        # Row 0: quarter header + floor headers
        table.cell(0, 0).merge(table.cell(1, 0))
        self._style_cell(table.cell(0, 0), 'الفصل\nوالعام', FONT, 9, bold=True, rtl=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)

        for fi, floor in enumerate(floors):
            b      = 1 + fi * 3
            target = targets.get(floor, 0.0)
            t_str  = str(int(target)) if target == int(target) else str(target)
            hdr    = f'{floor}  (Target={t_str}‰)'
            table.cell(0, b).merge(table.cell(0, b + 2))
            self._style_cell(table.cell(0, b), hdr, FONT_EN, 8, bold=True)
            self._set_cell_shading(table.cell(0, b), SHADE_HEADER)

        # Row 1: sub-column labels — tall row so vertical text is fully visible
        table.rows[1].height      = Inches(1.1)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        days_short = cfg['days_label_ar']
        for fi in range(N_FL):
            b = 1 + fi * 3
            self._style_cell(table.cell(1, b),     'النسبة',      FONT, 8, bold=True, rtl=True)
            self._style_cell(table.cell(1, b + 1), 'عدد الحالات', FONT, 8, bold=True, rtl=True)
            self._style_cell(table.cell(1, b + 2), days_short,    FONT, 8, bold=True, rtl=True)
            self._set_cell_text_direction(table.cell(1, b))
            self._set_cell_text_direction(table.cell(1, b + 1))
            self._set_cell_text_direction(table.cell(1, b + 2))
            self._set_cell_shading(table.cell(1, b),     SHADE_HEADER)
            self._set_cell_shading(table.cell(1, b + 1), SHADE_HEADER)
            self._set_cell_shading(table.cell(1, b + 2), SHADE_HEADER)

        # History data rows
        for r_i, entry in enumerate(recent_hist):
            ri      = 2 + r_i
            q_label = f"{entry.get('quarter', '')} {entry.get('year', '')}".strip()
            self._style_cell(table.cell(ri, 0), q_label, FONT, 8, rtl=True,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(table.cell(ri, 0), SHADE_HEADER)

            fl = entry.get('floors', {})
            for fi, floor in enumerate(floors):
                b   = 1 + fi * 3
                fd  = fl.get(floor, {})
                rate = fd.get('rate', 0.0)
                cas  = fd.get('cases', 0)
                dys  = fd.get('days', 0)
                rs   = f'{rate:.2f}‰' if rate else '0‰'
                self._style_cell(table.cell(ri, b),     rs,       FONT_EN, 8)
                self._style_cell(table.cell(ri, b + 1), str(cas), FONT_EN, 8)
                self._style_cell(table.cell(ri, b + 2), str(dys), FONT_EN, 8)

        # Current quarter row (highlighted)
        ci        = 2 + len(recent_hist)
        cur_label = f'{quarter} {year}'
        self._style_cell(table.cell(ci, 0), cur_label, FONT, 8, bold=True, rtl=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(table.cell(ci, 0), SHADE_TABLE)

        fl_cur = current.get('floors', {})
        for fi, floor in enumerate(floors):
            b   = 1 + fi * 3
            fd  = fl_cur.get(floor, {})
            rate = fd.get('rate', 0.0)
            cas  = fd.get('cases', 0)
            dys  = fd.get('days', 0)
            rs   = f'{rate:.2f}‰' if rate else '0‰'
            self._style_cell(table.cell(ci, b),     rs,       FONT_EN, 8, bold=True)
            self._style_cell(table.cell(ci, b + 1), str(cas), FONT_EN, 8, bold=True)
            self._style_cell(table.cell(ci, b + 2), str(dys), FONT_EN, 8, bold=True)
            self._set_cell_shading(table.cell(ci, b),     SHADE_TABLE)
            self._set_cell_shading(table.cell(ci, b + 1), SHADE_TABLE)
            self._set_cell_shading(table.cell(ci, b + 2), SHADE_TABLE)

        # Borders
        for r in range(TOTAL_ROWS):
            for c in range(TOTAL_COLS):
                if r == 1 and c == 0:
                    continue  # continuation cell — handled below via raw w:tc
                top    = BORDER_THICK if r == 0              else BORDER_FINE
                bottom = BORDER_THICK if r == TOTAL_ROWS - 1 else BORDER_FINE
                left   = BORDER_THICK if c == 0              else BORDER_FINE
                right  = BORDER_THICK if c == TOTAL_COLS - 1 else BORDER_FINE
                self._set_cell_border(table.cell(r, c),
                                      top=top, bottom=bottom, left=left, right=right)

        # Col 0 rows 0-1 are a vertical merge (الفصل والعام).
        # table.cell(1, 0) resolves to the row-0 anchor so the loop above never
        # touches the raw w:tc continuation element in row 1 — set it directly.
        row1_tc0 = table.rows[1]._tr.findall(qn('w:tc'))[0]
        self._set_tc_border(row1_tc0,
                            top=BORDER_FINE,
                            bottom=BORDER_THICK if TOTAL_ROWS == 2 else BORDER_FINE,
                            left=BORDER_THICK,
                            right=BORDER_FINE)

        # AI analysis for Table 1
        cfg        = self.cfg
        floor_stats = {
            fl: {
                'rate':   current.get('floors', {}).get(fl, {}).get('rate', 0.0),
                'cases':  current.get('floors', {}).get(fl, {}).get('cases', 0),
                'days':   current.get('floors', {}).get(fl, {}).get('days', 0),
                'target': targets.get(fl, 0.0),
            }
            for fl in floors
        }
        if ic_ai_service:
            analysis_text = ic_ai_service.analyze_floor_comparison_table(
                indicator=cfg['indicator'],
                desc=desc,
                floor_stats=floor_stats,
                n_floors=len(floors),
                quarter=quarter,
                year=str(year),
            )
        else:
            # Static fallback when ai_service is unavailable
            above_floors = {fl: d for fl, d in floor_stats.items() if d.get('rate', 0.0) > d.get('target', 999.0)}
            below_floors = {fl: d for fl, d in floor_stats.items() if d.get('rate', 0.0) <= d.get('target', 999.0)}
            curr_label   = f"{quarter} من العام {year}"
            analysis_text = f"يظهر الجدول رقم 1 مقارنة نتائج ال {cfg['indicator_ar']} خلال {desc} بحسب الأقسام، وفيما يخص {curr_label}"
            if above_floors:
                above_parts = [f"{fl} بـ {round(d['rate'] - d['target'], 2)}‰" for fl, d in above_floors.items()]
                analysis_text += f" فقد تجاوزت أقسام {' و'.join(above_parts)} الهدف"
                if below_floors:
                    analysis_text += f"، بينما حققت {' و'.join(below_floors.keys())} أداءً ضمن الهدف"
            else:
                analysis_text += f" فقد حققت جميع الأقسام أداءً ضمن الهدف"
            analysis_text += "."
        self._add_analysis_paragraph(doc, analysis_text)

    # =========================================================================
    # PER-FLOOR PAGES (dynamic — one page per floor with data)
    # =========================================================================

    def _build_floor_pages(self, doc, quarter, year, current, norm_history, floors, charts):
        cfg          = self.cfg
        indicator_ar = cfg['indicator_ar']

        # Section title — current quarter first, then up to 4 previous quarters
        prev_qy   = [(e.get('quarter', ''), e.get('year', '')) for e in norm_history[:-1][-4:]]
        prev_desc = self._describe_quarters(prev_qy)
        curr_desc = f'{quarter} من العام {year}'
        all_desc  = f'{curr_desc} و{prev_desc}' if prev_desc else curr_desc

        # Two-quarter desc (current + previous) — [-2:-1] is the entry before current
        prev_list = norm_history[-2:-1] if len(norm_history) >= 2 else []
        two_qy    = [(e.get('quarter', ''), e.get('year', '')) for e in prev_list]
        two_qy.append((quarter, year))
        two_desc  = self._describe_quarters(two_qy) if len(two_qy) > 1 else f'{quarter} {year}'

        self._add_section_title(
            doc, f'نتيجة ال {indicator_ar} خلال {all_desc} بحسب الأقسام:'
        )

        chart_num  = 1
        table_num  = 2

        # Only render pages for floors that actually have cases
        active_floors = [f for f in floors if current.get('floors', {}).get(f, {}).get('cases', 0) > 0]

        # Find the floor with the highest rate above its target (for notes)
        above_target_floors = [
            f for f in active_floors
            if current.get('floors', {}).get(f, {}).get('rate', 0.0)
            > current.get('floors', {}).get(f, {}).get('target', 999.0)
        ]

        # Previous-quarter germs for comparison in germs analysis
        prev_entry         = norm_history[-2] if len(norm_history) >= 2 else {}
        prev_germs_all     = prev_entry.get('germs_distribution', {})
        prev_quarter_label = prev_entry.get('quarter', '')
        prev_year_label    = str(prev_entry.get('year', ''))

        for fi, floor in enumerate(active_floors):
            floor_data = current.get('floors', {}).get(floor, {})
            rate       = floor_data.get('rate', 0.0)
            cases      = floor_data.get('cases', 0)
            target     = floor_data.get('target', 0.0)
            status     = 'ABOVE target' if rate > target else 'within target'
            logger.info(
                f"[{cfg['indicator']}] Floor {fi+1}/{len(active_floors)}: {floor} — "
                f"rate={rate:.2f}‰  cases={cases}  target={target}‰  ({status})"
            )

            # Floor that pushes the overall rate above target (for trend note)
            other_above = next(
                (f for f in above_target_floors if f != floor), None
            ) if above_target_floors else None

            # Per-floor history rates (previous quarters only)
            history_rates = []
            for e in norm_history[:-1][-4:]:
                fd = e.get('floors', {}).get(floor, {})
                if fd:
                    history_rates.append({
                        'quarter': e.get('quarter', ''),
                        'year':    str(e.get('year', '')),
                        'rate':    fd.get('rate', 0.0),
                    })

            # ── Trend chart ──────────────────────────────────────────────────
            self._add_sub_title(
                doc,
                f'مقارنة نسبة ال {indicator_ar} بين {all_desc} قسم ال {floor}:'
            )
            trend_key = f'floor_trend_{floor}'
            logger.info(f"[{cfg['indicator']}]   Chart {chart_num}: {floor} trend")
            self._add_chart(doc, charts.get(trend_key), str(chart_num))

            floor_cases = [c for c in current.get('cases', []) if _match_floor(c.get('floor'), floor)]

            if ic_ai_service:
                logger.info(f"[{cfg['indicator']}]   AI: trend analysis for {floor}…")
                trend_analysis = ic_ai_service.analyze_floor_trend(
                    indicator=cfg['indicator'],
                    quarter=quarter,
                    year=year,
                    floor=floor,
                    floor_data=floor_data,
                    history_rates=history_rates,
                    chart_num=chart_num,
                    cases=floor_cases,
                    all_floors_above=other_above,
                )
            else:
                rate      = floor_data.get('rate', 0.0)
                target    = floor_data.get('target', 0.0)
                status_ar = 'فوق المستهدف' if rate > target else 'ضمن المستهدف'
                trend_analysis = (
                    f"إن الرسم البياني رقم {chart_num} يظهر نسبة ال {indicator_ar} "
                    f"في قسم ال {floor} خلال {quarter} من العام {year}. "
                    f"بلغت النسبة {rate:.2f}‰ مقارنةً بالمستهدف {target}‰ — النتيجة {status_ar}."
                )
            self._add_analysis_paragraph(doc, trend_analysis)
            chart_num += 1

            # ── Germs chart ───────────────────────────────────────────────────
            self._add_sub_title(
                doc,
                f'مقارنة نسبة ال {indicator_ar} بين {two_desc} في قسم ال {floor} مع ال Germs المسببة:'
            )
            germs_key = f'floor_germs_{floor}'
            logger.info(f"[{cfg['indicator']}]   Chart {chart_num}: {floor} germs")
            self._add_chart(doc, charts.get(germs_key), str(chart_num))

            floor_germs = current.get('germs_distribution', {}).get(floor, {})
            prev_germs  = prev_germs_all.get(floor, {})

            if ic_ai_service:
                logger.info(f"[{cfg['indicator']}]   AI: germs analysis for {floor} ({len(floor_germs)} germs)…")
                germs_analysis = ic_ai_service.analyze_floor_germs(
                    indicator=cfg['indicator'],
                    quarter=quarter,
                    year=year,
                    floor=floor,
                    current_germs=floor_germs,
                    prev_germs=prev_germs,
                    prev_quarter_label=prev_quarter_label,
                    prev_year_label=prev_year_label,
                    chart_num=chart_num,
                    total_cases=floor_data.get('cases', 0),
                )
            else:
                top_germ = max(floor_germs, key=floor_germs.get) if floor_germs else None
                germs_analysis = (
                    f"إن الرسم البياني رقم {chart_num} يظهر ال Germs في قسم ال {floor} "
                    f"خلال {two_desc}."
                )
                if top_germ:
                    germs_analysis += (
                        f" أكثر الجراثيم انتشاراً هي {top_germ} بعدد {floor_germs[top_germ]} حالة."
                    )
            self._add_analysis_paragraph(doc, germs_analysis)
            chart_num += 1

            # ── Cases table ───────────────────────────────────────────────────
            if floor_cases:
                logger.info(f"[{cfg['indicator']}]   Table {table_num}: {floor} cases ({len(floor_cases)} rows)")
                self._add_page_break(doc)
                self._add_sub_title(
                    doc,
                    f'جدول مفصل عن حالات ال {indicator_ar} في قسم ال {floor} خلال {quarter} من العام {year}:'
                )
                self._add_cases_table(doc, floor_cases)

            # Cases analysis — only when table was rendered
            if floor_cases:
                if ic_ai_service:
                    logger.info(f"[{cfg['indicator']}]   AI: cases analysis for {floor}…")
                    cases_analysis = ic_ai_service.analyze_floor_cases(
                        indicator=cfg['indicator'],
                        quarter=quarter,
                        year=year,
                        floor=floor,
                        cases=floor_cases,
                        table_num=table_num,
                    )
                else:
                    cases_analysis = (
                        f"إن الجدول رقم {table_num} يظهر تفاصيل حالات ال {indicator_ar} "
                        f"في قسم ال {floor} خلال {quarter} من العام {year}. "
                        f"بلغ عدد الحالات {len(floor_cases)} حالة."
                    )
                if cases_analysis:
                    self._add_analysis_paragraph(doc, cases_analysis)
                table_num += 1

            # Page break between floors (skip on last floor)
            if fi < len(active_floors) - 1:
                self._add_page_break(doc)

    def _add_section_title(self, doc, text):
        para = doc.add_paragraph()
        self._set_rtl(para)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(3)
        run = para.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(14)
        run.bold      = True
        run.underline = True
        self._set_run_rtl(run)

    def _add_sub_title(self, doc, text):
        para = doc.add_paragraph()
        self._set_rtl(para)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(13)
        run.bold      = True
        run.underline = True
        self._set_run_rtl(run)

    def _add_chart(self, doc, chart, label):
        if chart is None:
            para = doc.add_paragraph(f'[ الرسم البياني {label} ]')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            return
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run()
        try:
            if isinstance(chart, io.BytesIO):
                chart.seek(0)
                run.add_picture(chart, width=Inches(5.5))
            else:
                run.add_picture(str(chart), width=Inches(5.5))
        except Exception as exc:
            logger.warning(f'Could not embed chart {label}: {exc}')
            para.clear()
            para.add_run(f'[ الرسم البياني {label} – خطأ في التحميل ]')

    # Risk factor column names (snake_case — matches normalized processor output)
    _RF_COLS = [
        "diabetic", "hypertension", "dyslipidemia", "heart_disease",
        "kidney_disease", "copd", "smoker", "obesity",
        "cardiac_congenital_malformation", "advanced_age",
        "length_of_stay", "duration_of_catheter",
        "cancer", "compromised_immune_system", "respiratory_pb",
    ]
    _RF_LABELS = {
        "diabetic":                        "Diabetic",
        "hypertension":                    "Hypertension",
        "dyslipidemia":                    "Dyslipidemia",
        "heart_disease":                   "Heart Disease",
        "kidney_disease":                  "Kidney Disease",
        "copd":                            "COPD",
        "smoker":                          "Smoker",
        "obesity":                         "Obesity",
        "cardiac_congenital_malformation": "Cardiac Congenital Malformation",
        "advanced_age":                    "Advanced Age",
        "length_of_stay":                  "Length of Stay",
        "duration_of_catheter":            "Duration of Catheter",
        "cancer":                          "Cancer",
        "compromised_immune_system":       "Compromised Immune System",
        "respiratory_pb":                  "Respiratory Problem",
    }

    def _get_risk_factors(self, case: dict) -> str:
        """Collect active risk factors.

        Tries boolean fields first (new processor format: diabetic=True/False, ...).
        Falls back to the pre-formatted 'risk_factors' string if no boolean fields
        are present (legacy current.json saved before processor refactor).
        """
        factors = []
        for col in self._RF_COLS:
            v = case.get(col)
            if v is True or str(v).strip().lower() in ('yes', '1', 'true'):
                factors.append(self._RF_LABELS.get(col, col))
        if factors:
            return ', '.join(factors)
        # Fallback: pre-formatted string stored by old processor
        legacy = (case.get('risk_factors') or '').strip()
        return legacy if legacy else '—'

    def _add_cases_table(self, doc, cases):
        cfg        = self.cfg
        is_vap     = cfg['indicator'] == 'VAP'

        if is_vap:
            HEADERS_AR = [
                'عوامل الخطر',
                'العمر',
                'الجرثومة',
                'تاريخ الإصابة',
                'تاريخ التنبيب',
                'تاريخ الدخول',
                'التشخيص',
                'رقم الحالة',
            ]
            col_widths = [
                Inches(1.50), Inches(0.55), Inches(1.00),
                Inches(0.65), Inches(0.65), Inches(0.65),
                Inches(1.00), Inches(0.67),
            ]
        else:
            HEADERS_AR = [
                'عوامل الخطر',
                'العمر',
                'الجرثومة',
                'تاريخ الإصابة',
                cfg['insertion_col_ar'],
                'تاريخ الدخول',
                'التشخيص',
                'رقم الحالة',
            ]
            col_widths = [
                Inches(1.50), Inches(0.55), Inches(1.00),
                Inches(0.65), Inches(0.65), Inches(0.65),
                Inches(1.00), Inches(0.67),
            ]

        n_data = max(len(cases), 1)
        n_cols = len(HEADERS_AR)
        table  = doc.add_table(rows=1 + n_data, cols=n_cols)
        table.autofit = False
        self._set_table_cell_margins(table, top=30, bottom=30, left=40, right=40)
        self._set_table_bidi(table)

        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

        # Header row
        for i, hdr in enumerate(HEADERS_AR):
            self._style_cell(table.cell(0, i), hdr, FONT, 10, bold=True, rtl=True)
            self._set_cell_shading(table.cell(0, i), SHADE_HEADER)
            self._set_cell_border(table.cell(0, i),
                top=BORDER_MEDIUM, bottom=BORDER_MEDIUM,
                left=BORDER_FINE, right=BORDER_FINE)

        def _v(case, *keys):
            for k in keys:
                v = case.get(k)
                if v is not None:
                    return str(v)
            return '—'

        def _cell_val(case, col_idx):
            if is_vap:
                if   col_idx == 0: return self._get_risk_factors(case)
                elif col_idx == 1: return _v(case, 'age_display', 'age')
                elif col_idx == 2: return _v(case, 'germs')
                elif col_idx == 3: return _v(case, 'date_of_infection')
                elif col_idx == 4: return _v(case, 'date_of_intubation')
                elif col_idx == 5: return _v(case, 'date_of_admission')
                elif col_idx == 6: return _v(case, 'diagnosis')
                elif col_idx == 7: return _v(case, 'case_number', 'nb_of_cases')
            else:
                if   col_idx == 0: return self._get_risk_factors(case)
                elif col_idx == 1: return _v(case, 'age_display', 'Age/Year', 'Age')
                elif col_idx == 2: return _v(case, 'germs', 'Germs')
                elif col_idx == 3: return _v(case, 'date_of_infection', 'Date of infection')
                elif col_idx == 4: return _v(case, 'date_of_insertion')
                elif col_idx == 5: return _v(case, 'date_of_admission', 'Date of admission')
                elif col_idx == 6: return _v(case, 'diagnosis', 'Diagnosis')
                elif col_idx == 7: return _v(case, 'case_number', 'nb_of_cases', 'Nb of cases')
            return '—'

        # Data rows
        if cases:
            for r_i, case in enumerate(cases):
                for c_i in range(n_cols):
                    val = _cell_val(case, c_i)
                    self._style_cell(table.cell(1 + r_i, c_i), val, FONT_EN, 8)
                    self._set_cell_border(table.cell(1 + r_i, c_i),
                        top=BORDER_FINE, bottom=BORDER_FINE,
                        left=BORDER_FINE, right=BORDER_FINE)
        else:
            for c_i in range(n_cols):
                self._style_cell(table.cell(1, c_i), '', FONT_EN, 8)
                self._set_cell_border(table.cell(1, c_i),
                    top=BORDER_FINE, bottom=BORDER_FINE,
                    left=BORDER_FINE, right=BORDER_FINE)

    def _add_analysis_paragraph(self, doc, text):
        """Analysis paragraph – RTL, BiDi-safe for mixed Arabic/Latin text."""
        for line in (text or '').split('\n'):
            line = (line or '\u200f').strip()
            if not line:
                line = '\u200f'
            self._add_bidi_para(doc, line, FONT_ANALYSIS, 11,
                                space_before=2, space_after=2)

    def _add_bidi_para(self, container, text, font_name, size,
                       bold=False, bold_values=None, space_before=2, space_after=2):
        """
        Render mixed Arabic/English text with correct BiDi bracket and % placement.

        Splits text at Arabic/LTR boundaries and assigns each segment its own run
        with an explicit direction marker so Word renders brackets and % correctly.
        """
        LTR_PAT = re.compile(r'[A-Za-z0-9][A-Za-z0-9 ().,;:%&/_\-+<>=*‰]*')

        parts, cursor = [], 0
        for m in LTR_PAT.finditer(text):
            if m.start() > cursor:
                parts.append((text[cursor:m.start()], False))  # Arabic
            parts.append((m.group(), True))                    # LTR
            cursor = m.end()
        if cursor < len(text):
            parts.append((text[cursor:], False))
        if not parts:
            parts = [(text, False)]

        para = container.add_paragraph()
        self._set_rtl(para)
        self._zero_spacing(para)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)

        for chunk, is_ltr in parts:
            if not chunk:
                continue
            display = (' ' + chunk.strip() + ' ') if is_ltr else chunk
            run = para.add_run(display)
            run.font.size = Pt(size)

            _bold = bold or (bool(bold_values) and any(bv in chunk for bv in bold_values))
            if _bold:
                run.bold = True

            rPr    = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)

            if is_ltr:
                run.font.name = FONT_EN
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), FONT_EN)
                for el in rPr.findall(qn('w:rtl')):
                    rPr.remove(el)
            else:
                run.font.name = font_name
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), font_name)
                if not rPr.findall(qn('w:rtl')):
                    rPr.append(OxmlElement('w:rtl'))

    # =========================================================================
    # LAST PAGE (identical structure to VAP)
    # =========================================================================

    def _build_last_page(self, doc):
        self._add_page_break(doc)
        self._add_spacer(doc, space_before=3)

        self._add_titled_row(doc, 'النتيجة النهائية', shade=SHADE_HEADER, border=BORDER_MEDIUM)
        self._add_titled_row(doc, 'Favorable',         shade=None,        border=BORDER_FINE)

        self._add_spacer(doc, space_before=3)
        self._add_checkbox_question_box(doc)
        self._add_spacer(doc, space_before=3)
        self._add_previous_actions_table(doc)
        self._add_spacer(doc, space_before=3)
        self._add_current_actions_table(doc)

        p_title = doc.add_paragraph()
        self._set_rtl(p_title)
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after  = Pt(4)
        p_title.paragraph_format.line_spacing = 1
        run_t = p_title.add_run('يجب ان تكون الاجراءات الواجب اتخاذها محددة وقابلة للقياس')
        run_t.font.name  = FONT_ANALYSIS
        run_t.font.size  = Pt(11)
        run_t.bold       = True
        run_t.underline  = True

        self._add_approval_table(doc)

        p_note = doc.add_paragraph()
        self._set_rtl(p_note)
        p_note.paragraph_format.space_before = Pt(4)
        p_note.paragraph_format.space_after  = Pt(0)
        p_note.paragraph_format.line_spacing = 1
        run_note = p_note.add_run(
            'يجب توقيع جميع المعنيين بمتابعة الاجراءات المتخذة او من ينوب عنهم في خانة الموافقة '
            'اضافة الى مدراء الادارات المعنيين'
        )
        run_note.font.name = FONT_ANALYSIS
        run_note.font.size = Pt(11)
        run_note.bold      = True

    def _add_titled_row(self, doc, text, shade=None, border=None):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(tbl, top=20, bottom=20, left=60, right=60)
        cell = tbl.cell(0, 0)
        self._style_cell(cell, text, FONT, 11, bold=True, rtl=True)
        if shade:
            self._set_cell_shading(cell, shade)
        if border:
            self._set_cell_border(cell, top=border, bottom=border, left=border, right=border)

    def _add_checkbox_question_box(self, doc):
        box_tbl = doc.add_table(rows=1, cols=1)
        box_tbl.autofit = False
        box_tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(box_tbl, top=30, bottom=30, left=60, right=60)
        cell = box_tbl.cell(0, 0)
        self._set_cell_border(cell, top=BORDER_FINE, bottom=BORDER_FINE,
                              left=BORDER_FINE, right=BORDER_FINE)

        p = cell.paragraphs[0]
        self._set_rtl(p)
        self._zero_spacing(p)
        self._add_run(p, 'هل تم اتخاذ اجراءات تحسينية سابقا حول موضوع التحليل :      ', FONT_ANALYSIS, 11)
        self._add_run(p, 'نعم', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐', None, 11)
        self._add_run(p, '         ', None, 11)
        self._add_run(p, 'كلا  ', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐ ', None, 11)

        p2 = cell.add_paragraph()
        self._set_rtl(p2)
        self._zero_spacing(p2)
        self._add_run(p2,
            'في حال نعم , يجب تعبئة جدول " الاجراءات المتخذة سابقا" و في حال كلا , الانتقال الى جدول "الاجراءات المتخذة الحالية"',
            FONT_ANALYSIS, 11)

    def _add_previous_actions_table(self, doc):
        table = doc.add_table(rows=3, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 4))
        self._style_cell(table.cell(0, 0), 'الاجراءات المتخذة  السابقة', FONT, 10, bold=True, rtl=True)

        headers = ['الرقم', 'الاجراءات المقترحة سابقا', 'التاريخ المتوقع لتنفيذ الاجراء',
                   'تم التنفيذ (نعم/كلا)', 'في حال كلا (الاجراءات التصحيحية الجديدة)']
        for j, h in enumerate(headers):
            self._style_cell(table.cell(1, j), h, FONT, 11, bold=True, rtl=True)

        for j in range(4):
            table.cell(2, j).text = ''

        last_cell = table.cell(2, 4)
        last_cell.text = ''
        p1 = last_cell.paragraphs[0]
        self._add_run(p1, ' ☐  الغاء الاجراء', FONT, 11)
        self._zero_spacing(p1)
        self._set_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = last_cell.add_paragraph()
        self._add_run(p2, ' ☐ اضافة الاجراء ضمن الاجراءات المتخذة الحالية. )الجدول ادناه(', FONT, 11)
        self._zero_spacing(p2)
        self._set_rtl(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(2):
            for cell in table.rows[i].cells:
                self._set_cell_shading(cell, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(1.8), Inches(1.37), Inches(1.1), Inches(2.1)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_current_actions_table(self, doc):
        table = doc.add_table(rows=6, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 3))
        self._style_cell(table.cell(0, 0), 'الاجراءات المتخذة  الحالية', FONT, 10, bold=True, rtl=True)

        headers = ['الرقم', 'الاجراءات الواجب اتخاذها لحل الموضوع',
                   'الشخص المعني بالمتابعة', 'التاريخ المتوقع للتنفيذ']
        for j, h in enumerate(headers):
            self._style_cell(table.cell(1, j), h, FONT, 11, bold=True, rtl=True)

        for i in range(2, 6):
            for j in range(4):
                table.cell(i, j).text = ''

        for i in range(2):
            for cell in table.rows[i].cells:
                self._set_cell_shading(cell, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(3.07), Inches(1.8), Inches(1.5)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_approval_table(self, doc):
        table = doc.add_table(rows=4, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).text = ''
        self._style_cell_text(table.cell(0, 1), 'اعداد', FONT, 11, bold=True, rtl=True)
        table.cell(0, 2).merge(table.cell(0, 4))
        self._style_cell_text(table.cell(0, 2), 'موافقة', FONT, 11, bold=True, rtl=True)

        for i, label in enumerate(['الاسم', 'التوقيع', 'التاريخ']):
            self._style_cell_text(table.cell(i + 1, 0), label, FONT, 11, bold=True, rtl=True)
            for j in range(1, 5):
                table.cell(i + 1, j).text = ''

        self._shade_row_cells(table.rows[0].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.8), Inches(1.5), Inches(1.52), Inches(1.52), Inches(1.53)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    # =========================================================================
    # UTILITIES (shared helpers)
    # =========================================================================

    def _describe_quarters(self, quarter_year_list):
        ORDINAL = {
            'الفصل الأول':  'الأول',
            'الفصل الثاني': 'الثاني',
            'الفصل الثالث': 'الثالث',
            'الفصل الرابع': 'الرابع',
        }
        year_quarters: dict = {}
        for q, y in quarter_year_list:
            if q and y:
                year_quarters.setdefault(str(y), []).append(str(q))

        parts = []
        for y in sorted(year_quarters.keys()):
            qs    = year_quarters[y]
            n     = len(qs)
            names = [ORDINAL.get(q, q) for q in qs]
            if n == 1:
                part = f'الفصل {names[0]} من العام {y}'
            elif n == 2:
                part = f'الفصلين {names[0]} و{names[1]} من العام {y}'
            elif n == 3:
                if qs == ['الفصل الأول', 'الفصل الثاني', 'الفصل الثالث']:
                    part = f'الفصول الثلاثة الأولى من العام {y}'
                elif qs == ['الفصل الثاني', 'الفصل الثالث', 'الفصل الرابع']:
                    part = f'الفصول الثلاثة الأخيرة من العام {y}'
                else:
                    part = f'الفصول {" و".join(names)} من العام {y}'
            elif n == 4:
                part = f'فصول العام {y} الأربعة'
            else:
                part = f'فصول العام {y}'
            parts.append(part)

        return ' و'.join(parts) if parts else ''

    def _style_cell(self, cell, text, font_name, font_size, bold=False,
                    rtl=False, alignment=None):
        """Alias kept for callers; delegates to _style_cell_text."""
        al = alignment or WD_ALIGN_PARAGRAPH.CENTER
        self._style_cell_text(cell, text, font_name, font_size,
                              bold=bold, rtl=rtl, alignment=al)

    def _style_cell_text(self, cell, text, font_name, font_size,
                         bold=False, italic=False, rtl=False,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.clear()
        para.alignment = alignment
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        if rtl:
            self._set_rtl(para)
        run = para.add_run(str(text) if text is not None else '')
        run.bold       = bold
        run.italic     = italic
        run.font.name  = font_name
        run.font.size  = Pt(font_size)
        if rtl:
            self._set_run_rtl(run)

    def _set_cell_shading(self, cell, fill_hex):
        tc_pr = cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn('w:shd')):
            tc_pr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),  'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tc_pr.append(shd)

    def _set_cell_border(self, cell, top=None, bottom=None, left=None, right=None):
        tc_pr = cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn('w:tcBorders')):
            tc_pr.remove(old)
        tc_borders = OxmlElement('w:tcBorders')
        for side, attrs in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            if attrs:
                el = OxmlElement(f'w:{side}')
                for k, v in attrs.items():
                    el.set(qn(f'w:{k}'), v)
                tc_borders.append(el)
        tc_pr.append(tc_borders)

    def _set_cell_text_direction(self, cell, direction: str = 'btLr'):
        """Rotate cell text vertically. direction: 'btLr' (upward) or 'tbRl' (downward)."""
        tc_pr = cell._tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn('w:textDirection')):
            tc_pr.remove(old)
        td = OxmlElement('w:textDirection')
        td.set(qn('w:val'), direction)
        tc_pr.append(td)

    def _set_tc_border(self, tc, top=None, bottom=None, left=None, right=None):
        """Same as _set_cell_border but takes a raw w:tc element (for merged cells)."""
        tc_pr = tc.get_or_add_tcPr()
        for old in tc_pr.findall(qn('w:tcBorders')):
            tc_pr.remove(old)
        tc_borders = OxmlElement('w:tcBorders')
        for side, attrs in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            if attrs:
                el = OxmlElement(f'w:{side}')
                for k, v in attrs.items():
                    el.set(qn(f'w:{k}'), v)
                tc_borders.append(el)
        tc_pr.append(tc_borders)

    def _set_table_cell_margins(self, table, top=0, bottom=0, left=0, right=0):
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tbl_pr)
        cell_mar = OxmlElement('w:tblCellMar')
        for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'),    str(val))
            el.set(qn('w:type'), 'dxa')
            cell_mar.append(el)
        for old in tbl_pr.findall(qn('w:tblCellMar')):
            tbl_pr.remove(old)
        tbl_pr.append(cell_mar)

    def _set_rtl(self, para):
        para.paragraph_format.right_to_left = True
        pPr = para._element.get_or_add_pPr()
        if pPr.find(qn('w:bidi')) is None:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

    def _set_run_rtl(self, run):
        rPr = run._element.get_or_add_rPr()
        # 1. RTL direction marker
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        # 2. Mirror font name to CS slot so Arabic glyphs use correct shaper
        try:
            cs_font = run.font.name
            if cs_font:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:cs'), cs_font)
        except Exception:
            pass
        # 3. CS bold so bold Arabic glyphs actually render bold
        try:
            if run.bold:
                rPr.append(OxmlElement('w:bCs'))
        except Exception:
            pass

    def _set_table_bidi(self, table):
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tbl_pr)
        bidi = OxmlElement('w:bidiVisual')
        bidi.set(qn('w:val'), '1')
        tbl_pr.append(bidi)

    def _set_table_rtl(self, table):
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tbl_pr)
        bidi = OxmlElement('w:bidiVisual')
        bidi.set(qn('w:val'), '1')
        tbl_pr.append(bidi)

    def _apply_borders_to_table(self, table, border):
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, top=border, bottom=border,
                                      left=border, right=border)

    def _add_run(self, para, text, font_name, font_size, bold=False):
        run = para.add_run(str(text))
        if font_name:
            run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold      = bold
        return run

    def _zero_spacing(self, para):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing = 1

    def _add_spacer(self, doc, space_before=3):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(0)

    def _add_page_break(self, doc):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing = Pt(1)
        run  = para.add_run()
        br   = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._element.append(br)

    def _shade_row_cells(self, cells, fill):
        for cell in cells:
            self._set_cell_shading(cell, fill)


