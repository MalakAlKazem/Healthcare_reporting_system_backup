"""
VAP (Ventilator-Associated Pneumonia) DOCX Report Generator

Report structure
────────────────
Page 1    : Metadata table · Results table (6 quarters) · Analysis box ·
            Floor comparison table (19 cols, up to 6 quarters)
Page 2    : Section II – ICU results
              1-1  ICU trend chart   + analysis
              1-2  ICU germs chart   + analysis
Page 3    : 1-3  ICU cases table    + analysis
            2-1  CCU trend chart    + analysis
Page 4    : 3-2  CCU germs chart    + analysis
            2-3  CCU cases table    + analysis
Page 5    : 3-1  ICN trend chart    + analysis
            3-2  ICN germs chart    + analysis
Page 6    : 3-3  ICN cases table    + analysis
Last page : Final result · Action tables · Approval table

Analysis texts are static placeholders; see _add_ai_analysis() for the
integration point where AI-generated text will replace them.
"""

import io
import os
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# ── Logging shim (works whether loguru is installed or not) ──────────────────
try:
    from loguru import logger
except ImportError:
    import logging
    logger = logging.getLogger(__name__)
    logging.basicConfig(level=logging.INFO)

# =============================================================================
# FONTS
# =============================================================================
FONT          = 'Traditional Arabic'
FONT_EN       = 'Calibri'
FONT_ANALYSIS = 'Frutiger LT Arabic 45 Light'

# =============================================================================
# LAYOUT CONSTANTS
# =============================================================================
PAGE_WIDTH   = Inches(8.27)
PAGE_HEIGHT  = Inches(11.69)
MARGIN       = Inches(0.7)
USABLE_WIDTH = Inches(6.87)

# =============================================================================
# BORDER PRESETS
# =============================================================================
BORDER_THICK  = {'val': 'single', 'sz': '18', 'color': '000000', 'space': '0'}
BORDER_MEDIUM = {'val': 'single', 'sz': '12', 'color': '000000', 'space': '0'}
BORDER_THIN   = {'val': 'single', 'sz': '8',  'color': '000000', 'space': '0'}
BORDER_FINE   = {'val': 'single', 'sz': '4',  'color': '000000', 'space': '0'}
BORDER_NONE   = {'val': 'none',   'sz': '0',  'color': 'FFFFFF', 'space': '0'}

# =============================================================================
# SHADING COLORS
# =============================================================================
SHADE_HEADER = 'EDEBE3'
SHADE_TABLE  = 'D9E2F3'
SHADE_TOTAL  = 'F2F2F2'

STANDARD_FLOORS = ['ICU', 'CCU', 'CSU', 'Ped', 'ICN', 'ITU', 'Neonatal']
FLOOR_TARGETS   = {
    'ICU': 25.0, 'CCU': 15.0, 'CSU': 9.5,
    'Ped': 5.5,  'ICN': 10.0, 'ITU': 25.0, 'Neonatal': 0.0,
}


class VAPDocxGenerator:
    """Generates a VAP (Ventilator-Associated Pneumonia) analysis DOCX report."""

    def __init__(self):
        self.logo_path = 'app/assets/LOGO.png'

    # =========================================================================
    # PUBLIC API
    # =========================================================================

    def generate_report(self, stats, history=None, chart_paths=None, options=None):
        """
        Generate DOCX report for VAP infection control.

        Args:
            stats       : Statistics dict from VAPStatistics.calculate_all_statistics()
            history     : List of historical quarters from VAPHistory.get_all()
            chart_paths : Dict of chart_name → BytesIO or file path (optional)
            options     : Dict with 'quarter' and 'year'

        Returns:
            dict with 'filePath' and 'fileName'
        """
        chart_paths = chart_paths or {}
        options     = options or {}
        summary     = stats.get('summary', {})
        quarter     = options.get('quarter', summary.get('quarter', 'الفصل الرابع'))
        year        = str(options.get('year', summary.get('year', '2025')))
        logger.info(f"Generating VAP DOCX report for {quarter} {year}...")

        history = list(history or [])

        # Generate charts from history if caller did not provide them
        if not chart_paths:
            chart_paths = self._generate_charts(history, stats, quarter, year)

        doc = Document()
        self._setup_section(doc.sections[0])

        self._build_page1(doc, quarter, year, stats, history)
        self._build_page2(doc, stats, chart_paths, quarter, year, history)
        self._build_page3(doc, stats, chart_paths, quarter, year, history)
        self._build_page4(doc, stats, chart_paths, quarter, year, history)
        self._build_page5(doc, stats, chart_paths, quarter, year, history)
        self._build_page6(doc, stats, chart_paths, quarter, year, history)
        self._build_last_page(doc)

        reports_dir = os.path.normpath(
            os.path.join(os.path.dirname(__file__), '..', '..', '..', '..', 'storage', 'reports')
        )
        os.makedirs(reports_dir, exist_ok=True)
        file_name = f'VAP report {quarter} {year}.docx'
        file_path = os.path.join(reports_dir, file_name)
        try:
            doc.save(file_path)
        except PermissionError:
            # File is open in Word — save with timestamp suffix instead
            ts        = datetime.now().strftime('%H%M%S')
            file_name = f'VAP report {quarter} {year}_{ts}.docx'
            file_path = os.path.join(reports_dir, file_name)
            doc.save(file_path)
            logger.warning(f"Original file was locked; saved as: {file_name}")
        logger.info(f"VAP DOCX saved: {file_path}")

        return {'filePath': file_path, 'fileName': file_name}

    def _generate_charts(self, history, stats=None, quarter=None, year=None):
        """Generate all 6 VAP charts from history + current quarter; return {chart_key: BytesIO}."""
        try:
            from app.infection_control.VAP.vap_chart_generator import VAPChartGenerator
            from app.infection_control.VAP.vap_history import VAPHistory
            from app.infection_control.VAP.vap_statistics import STANDARD_FLOORS, FLOOR_TARGETS

            full_history = list(history)

            # Inject the current quarter entry so charts include Q4 (or whatever is current)
            if stats and quarter and year:
                summary     = stats.get('summary', {})
                floor_stats = stats.get('floor_stats', {})
                cur_entry = {
                    'quarter':         str(quarter),
                    'year':            str(year),
                    'total_cases':     summary.get('total_cases', 0),
                    'total_vent_days': summary.get('total_vent_days', 0),
                    'floors': {
                        floor: {
                            'cases':           floor_stats.get(floor, {}).get('cases', 0),
                            'ventilator_days': floor_stats.get(floor, {}).get('ventilator_days', 0),
                            'rate':            floor_stats.get(floor, {}).get('rate', 0.0),
                            'target':          FLOOR_TARGETS.get(floor, 0.0),
                        }
                        for floor in STANDARD_FLOORS
                    },
                    'germs_overall':  stats.get('germs_overall', {}),
                    'germs_by_floor': {
                        floor: data
                        for floor, data in stats.get('germs_by_floor', {}).items()
                        if data.get('total', 0) > 0
                    },
                }
                # Only append if not already present (avoid duplicates)
                already = any(
                    e.get('quarter') == str(quarter) and e.get('year') == str(year)
                    for e in full_history
                )
                if not already:
                    full_history.append(cur_entry)

            vap_hist         = VAPHistory.__new__(VAPHistory)
            vap_hist.history = full_history
            chart_data       = vap_hist.get_chart_data()

            gen  = VAPChartGenerator(output_dir='storage/charts/infection_control/VAP')
            bufs = gen.generate_all_charts(chart_data)
            logger.info(f"Generated {len(bufs)} VAP charts for report")
            return bufs
        except Exception as exc:
            logger.warning(f"Chart generation skipped (charts will be placeholders): {exc}")
            return {}

    # =========================================================================
    # SECTION SETUP
    # =========================================================================

    def _setup_section(self, section):
        section.page_width    = PAGE_WIDTH
        section.page_height   = PAGE_HEIGHT
        section.top_margin    = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin   = MARGIN
        section.right_margin  = MARGIN

        sectPr    = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '12')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

        self._add_header(section)
        self._add_footer(section)

    def _add_header(self, section):
        header = section.header
        for element in list(header._element):
            header._element.remove(element)

        table  = header.add_table(rows=1, cols=3, width=Inches(7))
        self._set_table_cell_margins(table, top=0, bottom=0, left=0, right=0)
        cells  = table.rows[0].cells
        cells[0].width = Inches(1.0)
        cells[1].width = Inches(5.0)
        cells[2].width = Inches(1.0)

        if os.path.exists(self.logo_path):
            para = cells[0].paragraphs[0]
            para.add_run().add_picture(self.logo_path, width=Inches(0.5))
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._zero_spacing(para)

        para      = cells[1].paragraphs[0]
        run       = para.add_run("نموذج تحليل البيانات")
        run.font.name = FONT_ANALYSIS
        run.font.size = Pt(16)
        run.bold  = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(para)
        self._zero_spacing(para)

        para      = cells[2].paragraphs[0]
        run       = para.add_run('QS-36F-103(4)\nApd 15/12/2019')
        run.font.name = FONT_EN
        run.font.size = Pt(7)
        run.bold  = True
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._zero_spacing(para)

        line_para = header.add_paragraph()
        line_para.paragraph_format.space_before = Pt(0)
        line_para.paragraph_format.space_after  = Pt(8)
        line_para.paragraph_format.line_spacing = Pt(1)
        line_para.add_run("").font.size = Pt(1)
        pPr    = line_para._element.get_or_add_pPr()
        pBdr   = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_footer(self, section):
        footer = section.footer
        for element in list(footer._element):
            footer._element.remove(element)

        para  = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run   = para.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._element.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)

        run.font.size = Pt(10)
        run.font.name = FONT_EN

    # =========================================================================
    # PAGE 1: Metadata + Results table + Analysis box + Floor comparison + Summary
    # =========================================================================

    def _build_page1(self, doc, quarter, year, stats, history):
        self._add_spacer(doc, space_before=6)
        self._add_metadata_table(doc, quarter, year)
        doc.add_paragraph()
        self._add_results_table(doc, stats, history)
        self._add_spacer(doc)
        self._add_analysis_box(doc, quarter, year, stats)
        doc.add_paragraph()
        self._add_floor_comparison_section(doc, stats, history, quarter, year)
        self._add_page_break(doc)

    def _add_metadata_table(self, doc, quarter, year):
        """3-row × 4-col metadata table."""
        table = doc.add_table(rows=3, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=0, bottom=0, left=60, right=60)

        col_w = [Inches(1.93), Inches(1.06), Inches(2.71), Inches(1.17)]
        for i, width in enumerate(col_w):
            for row in table.rows:
                row.cells[i].width = width

        # Row 0
        self._style_cell_text(table.cell(0, 3), "موضوع التحليل", FONT, 10, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)
        self._style_cell_text(table.cell(0, 2), " تحليل بيانات مؤشر ال VAP ", FONT_EN, 9, bold=True)
        self._style_cell_text(table.cell(0, 1), "الإدارة", FONT, 10, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(0, 0), "الجودة والسلامة", FONT, 11, rtl=True)

        # Row 1
        self._style_cell_text(table.cell(1, 3), "", FONT, 9, rtl=True)
        self._set_cell_shading(table.cell(1, 3), SHADE_HEADER)
        self._style_cell_text(table.cell(1, 2), "", FONT_EN, 10)
        self._style_cell_text(table.cell(1, 1), "الوحدة الإدارية", FONT, 10, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(1, 0), "ضبط العدوى", FONT, 10, rtl=True)

        # Row 2
        self._style_cell_text(
            table.cell(2, 3),
            "مصادر\nالبيانات\n)تحديد اسماء\nالنماذج(",
            FONT, 9, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(2, 3), SHADE_HEADER)
        self._style_cell_text(
            table.cell(2, 2),
            "Health care associated infection (HCAI) sheet",
            FONT, 9, rtl=True
        )
        self._style_cell_text(
            table.cell(2, 1),
            "الشهر / العام\nأو\nالفصل / العام",
            FONT, 9, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(2, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(2, 0), f"{quarter}  {year}/", FONT, 11, rtl=True)

        # Vertical merges
        table.cell(0, 2).merge(table.cell(1, 2))
        table.cell(0, 3).merge(table.cell(1, 3))
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)

        row0_tcs = table.rows[0]._tr.findall(qn('w:tc'))
        row1_tcs = table.rows[1]._tr.findall(qn('w:tc'))

        self._set_cell_border(table.cell(0, 0), top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THICK, right=BORDER_THIN)
        self._set_cell_border(table.cell(1, 0), top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THICK, right=BORDER_THIN)
        self._set_cell_border(table.cell(2, 0), top=BORDER_THIN,  bottom=BORDER_THICK, left=BORDER_THICK, right=BORDER_THIN)
        self._set_cell_border(table.cell(0, 1), top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THIN,  right=BORDER_THICK)
        self._set_cell_border(table.cell(1, 1), top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THIN,  right=BORDER_THICK)
        self._set_cell_border(table.cell(2, 1), top=BORDER_THIN,  bottom=BORDER_THICK, left=BORDER_THIN,  right=BORDER_THICK)
        self._set_tc_border(row0_tcs[2], top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THICK, right=BORDER_THIN)
        self._set_tc_border(row1_tcs[2], top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THICK, right=BORDER_THIN)
        self._set_cell_border(table.cell(2, 2), top=BORDER_THIN, bottom=BORDER_THICK, left=BORDER_THICK, right=BORDER_THIN)
        self._set_tc_border(row0_tcs[3], top=BORDER_THICK, bottom=BORDER_THIN,  left=BORDER_THIN, right=BORDER_THICK)
        self._set_tc_border(row1_tcs[3], top=BORDER_THIN,  bottom=BORDER_THIN,  left=BORDER_THIN, right=BORDER_THICK)
        self._set_cell_border(table.cell(2, 3), top=BORDER_THIN, bottom=BORDER_THICK, left=BORDER_THIN, right=BORDER_THICK)

    def _add_results_table(self, doc, stats, history):
        """
        4-row × 6-col results table.
        Col 5 (rightmost / first in RTL): current quarter floor stats.
        Cols 0-4: last 5 history quarters floor stats (oldest → newest left → right).
        """
        COLS   = 6
        N_HIST = COLS - 1   # 5 history columns

        summary = stats.get('summary', {})
        cur_quarter = summary.get('quarter', '')
        cur_year    = str(summary.get('year', ''))

        # Last 5 history entries, oldest first
        last5 = history[-N_HIST:] if len(history) >= N_HIST else history
        while len(last5) < N_HIST:
            last5 = [{}] + last5

        table = doc.add_table(rows=4, cols=COLS)
        table.autofit = False
        self._set_table_cell_margins(table, top=40, bottom=40, left=60, right=60)

        w_col = Inches(6.87 / COLS)
        for i in range(COLS):
            for row in table.rows:
                row.cells[i].width = w_col

        # Row 0: merged title
        cell0 = table.cell(0, 0)
        for c in range(1, COLS):
            cell0 = cell0.merge(table.cell(0, c))
        self._style_cell_text(table.cell(0, 0), "النتائج", FONT, 11, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)

        # Row 1: instruction (cols 0 to N_HIST-1 merged) + "النتيجة الحالية" (col N_HIST)
        instr = table.cell(1, 0)
        for c in range(1, N_HIST):
            instr = instr.merge(table.cell(1, c))
        self._style_cell_text(
            table.cell(1, 0),
            "في حال كان هناك نتائج سابقة عن الموضوع الذي يتم تحليله يجب ذكره مع تحديد الفترة:\nالشهر / الفصل / العام",
            FONT, 11, rtl=True
        )
        self._set_cell_shading(table.cell(1, 0), SHADE_HEADER)
        self._style_cell_text(table.cell(1, N_HIST), "النتيجة الحالية", FONT, 11, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, N_HIST), SHADE_TABLE)

        # Row 2: last 5 history quarter labels + current quarter label (highlighted)
        for i in range(N_HIST):
            entry = last5[i]
            label = f"{entry.get('quarter', '')} {entry.get('year', '')}" if entry.get('quarter') else ""
            self._style_cell_text(table.cell(2, i), label, FONT, 9, rtl=True)
            self._set_cell_shading(table.cell(2, i), SHADE_HEADER)
        self._style_cell_text(
            table.cell(2, N_HIST),
            f"{cur_quarter} {cur_year}",
            FONT, 10, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(2, N_HIST), SHADE_TABLE)

        # Row 3: floor rates per quarter
        # Current quarter (rightmost): from stats['floor_stats']
        self._fill_rates_cell(table.cell(3, N_HIST), stats.get('floor_stats', {}))

        # History columns: from entry['floors']
        for i, entry in enumerate(last5):
            if entry and entry.get('floors'):
                self._fill_rates_cell(table.cell(3, i), entry['floors'])
            else:
                table.cell(3, i).text = ''

        # Borders
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, top=BORDER_THIN, bottom=BORDER_THIN,
                                      left=BORDER_THIN, right=BORDER_THIN)
        for c in range(COLS):
            self._set_cell_border(table.cell(0, c), top=BORDER_THICK,
                                  left=BORDER_THIN, right=BORDER_THIN, bottom=BORDER_THIN)
            self._set_cell_border(table.cell(3, c), bottom=BORDER_THICK,
                                  left=BORDER_THIN, right=BORDER_THIN, top=BORDER_THIN)
        for r in range(4):
            self._set_cell_border(table.cell(r, 0),        left=BORDER_THICK,
                                  top=BORDER_THIN, right=BORDER_THIN, bottom=BORDER_THIN)
            self._set_cell_border(table.cell(r, COLS - 1), right=BORDER_THICK,
                                  top=BORDER_THIN, left=BORDER_THIN,  bottom=BORDER_THIN)

    def _fill_rates_cell(self, cell, floors):
        FLOOR_ORDER = ['ICU', 'CCU', 'CSU', 'Ped', 'ICN', 'ITU']
        FLOOR_LABEL = {
            'ICU': 'ICU', 'CCU': 'CCU', 'CSU': 'CSU',
            'Ped': 'Ped', 'ICN': 'Neonatal', 'ITU': 'ITU',
        }
        pairs = []
        for f in FLOOR_ORDER:
            fd = (floors or {}).get(f, {})
            try:
                rate = float(fd.get('rate', 0.0) if fd else 0.0)
            except (TypeError, ValueError):
                rate = 0.0
            rate_str = f'{rate:.2f}'.rstrip('0').rstrip('.') if rate else '0'
            pairs.append(f'{FLOOR_LABEL[f]} = {rate_str}‰')

        cell.paragraphs[0].clear()
        for i in range(0, len(pairs), 2):
            chunk = '  \n  '.join(pairs[i:i + 2])
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(chunk)
            r.font.name = FONT_EN
            r.font.size = Pt(8)

    def _add_analysis_box(self, doc, quarter, year, stats):
        """Box with section header + chart comparison caption."""
        box = doc.add_table(rows=1, cols=1)
        box.autofit = False
        box.columns[0].width = USABLE_WIDTH
        cell = box.cell(0, 0)
        self._zero_spacing(cell.paragraphs[0])
        self._set_table_cell_margins(box, top=0, bottom=10, left=20, right=20)
        # Top border only — no bottom/left/right
        self._set_cell_border(cell, top=BORDER_MEDIUM, left=BORDER_THICK, right=BORDER_THICK)

        for paragraph in list(cell.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)

        # Header row inside box
        title_tbl  = cell.add_table(rows=1, cols=1)
        title_tbl.autofit = False
        title_tbl.columns[0].width = USABLE_WIDTH
        title_cell = title_tbl.cell(0, 0)
        self._set_table_cell_margins(title_tbl, top=0, bottom=0, left=20, right=20)
        self._set_cell_shading(title_cell, SHADE_HEADER)
        self._set_cell_border(title_cell, bottom=BORDER_MEDIUM)
        self._style_cell_text(title_cell, "*التحليل", FONT, 11, bold=True, rtl=True)

        for p in cell.paragraphs:
            self._zero_spacing(p)

    def _describe_quarters(self, quarter_year_list):
        """
        Return Arabic description of a list of (quarter, year) pairs.
        Groups by year and applies Arabic grammar patterns.
        E.g. [(Q3,2024),(Q4,2024),(Q1,2025),(Q2,2025),(Q3,2025)]
          → "الفصلين الثالث والرابع من العام 2024 والفصول الثلاثة الأولى من العام 2025"
        """
        ORDINAL = {
            'الفصل الأول':  'الأول',
            'الفصل الثاني': 'الثاني',
            'الفصل الثالث': 'الثالث',
            'الفصل الرابع': 'الرابع',
        }
        year_quarters = {}
        for q, y in quarter_year_list:
            year_quarters.setdefault(str(y), []).append(q)

        parts = []
        for y in sorted(year_quarters.keys()):
            qs    = year_quarters[y]
            n     = len(qs)
            names = [ORDINAL.get(q, q) for q in qs]
            if n == 1:
                part = f"الفصل {names[0]} من العام {y}"
            elif n == 2:
                part = f"الفصلين {names[0]} و{names[1]} من العام {y}"
            elif n == 3:
                if qs == ['الفصل الأول', 'الفصل الثاني', 'الفصل الثالث']:
                    part = f"الفصول الثلاثة الأولى من العام {y}"
                elif qs == ['الفصل الثاني', 'الفصل الثالث', 'الفصل الرابع']:
                    part = f"الفصول الثلاثة الأخيرة من العام {y}"
                else:
                    part = f"الفصول {' و'.join(names)} من العام {y}"
            elif n == 4:
                part = f"فصول العام {y} الأربعة"
            else:
                part = f"فصول العام {y}"
            parts.append(part)

        return " و".join(parts)

    def _add_floor_comparison_section(self, doc, stats, history, quarter, year):
        """
        Dynamic title + 19-column floor comparison table + placeholder analysis text.

        Columns: الفصل والعام | ICU×3 | CCU×3 | CSU×3 | Ped×3 | ICN×3 | ITU×3
        Each floor triple: النسبة | عدد الحالات | أيام التنفس
        Rows: last 5 history quarters (oldest→newest) + current quarter (highlighted)
        """
        FLOORS_DISP  = ['ICU', 'CCU', 'CSU', 'Ped', 'ICN', 'ITU']
        TARGETS_DISP = {'ICU': 25.0, 'CCU': 15.0, 'CSU': 9.5,
                        'Ped': 5.5,  'ICN': 10.0,  'ITU': 25.0}

        recent_hist = history[-5:] if len(history) >= 5 else list(history)
        cur_quarter = str(quarter)
        cur_year    = str(year)

        # (quarter, year) pairs for title generation
        all_qy = [(e.get('quarter', ''), str(e.get('year', ''))) for e in recent_hist]
        all_qy.append((cur_quarter, cur_year))

        # ── Dynamic title ─────────────────────────────────────────────────────
        desc       = self._describe_quarters(all_qy)
        title_text = f"جدول يظهر مقارنة نتائج ال VAP بين {desc} بحسب الأقسام:"

        tp = doc.add_paragraph()
        self._set_rtl(tp)                        # sets bidi → auto right-justifies
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after  = Pt(3)
        tr = tp.add_run(title_text)
        tr.font.name = FONT
        tr.font.size = Pt(14)
        tr.bold      = True
        tr.underline = True
        self._set_run_rtl(tr)

        # ── Table layout ──────────────────────────────────────────────────────
        N_FL       = len(FLOORS_DISP)
        TOTAL_COLS = 1 + N_FL * 3           # 19
        n_data     = len(recent_hist) + 1   # history + current
        TOTAL_ROWS = 2 + n_data             # 2 header rows + data rows

        table = doc.add_table(rows=TOTAL_ROWS, cols=TOTAL_COLS)
        table.autofit = False          # fixed layout – respect explicit widths
        self._set_table_cell_margins(table, top=20, bottom=20, left=10, right=10)

        # Set explicit table preferred width = full usable page width
        # Without <w:tblW>, Word ignores fixed-layout widths and auto-sizes.
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        for _old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(_old)
        _tblW = OxmlElement('w:tblW')
        _tblW.set(qn('w:w'), str(int(USABLE_WIDTH / 914400 * 1440)))  # EMU → dxa
        _tblW.set(qn('w:type'), 'dxa')
        tblPr.append(_tblW)

        # Column widths – 1.30 + 6×(0.34+0.28+0.34) = 1.30 + 5.76 = 7.06"
        # Word will compress slightly to fit tblW (6.87"), keeping proportions.
        q_w     = Inches(1.30)
        rate_w  = Inches(0.34)   # النسبة
        cases_w = Inches(0.28)   # عدد الحالات
        vent_w  = Inches(0.34)   # أيام التنفس

        for row in table.rows:
            row.cells[0].width = q_w
            for fi in range(N_FL):
                b = 1 + fi * 3
                row.cells[b].width     = rate_w
                row.cells[b + 1].width = cases_w
                row.cells[b + 2].width = vent_w

        # ── Row 0: "الفصل والعام" (vertical merge rows 0-1) + floor headers ──
        table.cell(0, 0).merge(table.cell(1, 0))
        self._style_cell_text(table.cell(0, 0), "الفصل\nوالعام", FONT, 9, bold=True, rtl=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)

        for fi, floor in enumerate(FLOORS_DISP):
            b      = 1 + fi * 3
            target = TARGETS_DISP.get(floor, 0.0)
            t_str  = str(int(target)) if target == int(target) else str(target)
            hdr    = f"{floor}  (Target={t_str}‰)"
            table.cell(0, b).merge(table.cell(0, b + 2))
            self._style_cell_text(table.cell(0, b), hdr, FONT_EN, 8, bold=True)
            self._set_cell_shading(table.cell(0, b), SHADE_HEADER)

        # ── Row 1: sub-column labels ──────────────────────────────────────────
        for fi in range(N_FL):
            b = 1 + fi * 3
            self._style_cell_text(table.cell(1, b),     "النسبة",       FONT, 7, bold=True, rtl=True)
            self._style_cell_text(table.cell(1, b + 1), "عدد\nالحالات", FONT, 7, bold=True, rtl=True)
            self._style_cell_text(table.cell(1, b + 2), "أيام\nالتنفس", FONT, 7, bold=True, rtl=True)
            self._set_cell_shading(table.cell(1, b),     SHADE_HEADER)
            self._set_cell_shading(table.cell(1, b + 1), SHADE_HEADER)
            self._set_cell_shading(table.cell(1, b + 2), SHADE_HEADER)

        # ── History data rows ─────────────────────────────────────────────────
        for r_i, entry in enumerate(recent_hist):
            ri      = 2 + r_i
            q_label = f"{entry.get('quarter', '')} {entry.get('year', '')}"
            self._style_cell_text(table.cell(ri, 0), q_label, FONT, 8, rtl=True,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(table.cell(ri, 0), SHADE_HEADER)

            fl = entry.get('floors', {})
            for fi, floor in enumerate(FLOORS_DISP):
                b    = 1 + fi * 3
                fd   = fl.get(floor, {})
                rate = fd.get('rate', 0.0)
                cas  = fd.get('cases', 0)
                vnt  = fd.get('ventilator_days', 0)
                rs   = f"{rate:.2f}‰" if rate else "0‰"
                self._style_cell_text(table.cell(ri, b),     rs,       FONT_EN, 8)
                self._style_cell_text(table.cell(ri, b + 1), str(cas), FONT_EN, 8)
                self._style_cell_text(table.cell(ri, b + 2), str(vnt), FONT_EN, 8)

        # ── Current quarter row (highlighted) ─────────────────────────────────
        ci        = 2 + len(recent_hist)
        cur_label = f"{cur_quarter} {cur_year}"
        self._style_cell_text(table.cell(ci, 0), cur_label, FONT, 8, bold=True, rtl=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(table.cell(ci, 0), SHADE_TABLE)

        fl_stats = stats.get('floor_stats', {})
        for fi, floor in enumerate(FLOORS_DISP):
            b    = 1 + fi * 3
            fd   = fl_stats.get(floor, {})
            rate = fd.get('rate', 0.0)
            cas  = fd.get('cases', 0)
            vnt  = fd.get('ventilator_days', 0)
            rs   = f"{rate:.2f}‰" if rate else "0‰"
            self._style_cell_text(table.cell(ci, b),     rs,       FONT_EN, 8, bold=True)
            self._style_cell_text(table.cell(ci, b + 1), str(cas), FONT_EN, 8, bold=True)
            self._style_cell_text(table.cell(ci, b + 2), str(vnt), FONT_EN, 8, bold=True)
            self._set_cell_shading(table.cell(ci, b),     SHADE_TABLE)
            self._set_cell_shading(table.cell(ci, b + 1), SHADE_TABLE)
            self._set_cell_shading(table.cell(ci, b + 2), SHADE_TABLE)

        # ── Borders (outer thick, inner fine) ────────────────────────────────
        # Skip (1, 0): that _tc is the vMerge *continuation* cell of the merged
        # quarter header.  Setting borders on it draws a phantom line that splits
        # the merged area visually.  The restart cell at (0, 0) already carries
        # the correct outer borders.
        for r in range(TOTAL_ROWS):
            for c in range(TOTAL_COLS):
                if r == 1 and c == 0:        # continuation of merged quarter cell
                    continue
                top    = BORDER_THICK if r == 0               else BORDER_FINE
                bottom = BORDER_THICK if r == TOTAL_ROWS - 1  else BORDER_FINE
                left   = BORDER_THICK if c == 0               else BORDER_FINE
                right  = BORDER_THICK if c == TOTAL_COLS - 1  else BORDER_FINE
                self._set_cell_border(table.cell(r, c),
                    top=top, bottom=bottom, left=left, right=right)

        # ── Analysis text (TODO: replace with AI-generated analysis) ──────────
        hist_qy   = [(e.get('quarter', ''), str(e.get('year', ''))) for e in recent_hist]
        hist_desc = self._describe_quarters(hist_qy) if hist_qy else ""
        analysis_text = (
            f"إن الجدول رقم 1 يظهر لنا مقارنة عدد أيام المرضى الموضوعين على جهاز التنفس الاصطناعي "
            f"والحالات ونسبة ال VAP في هذه الأقسام"
            + (f" بين {hist_desc}" if hist_desc else "")
            + " بحسب الأقسام."
        )
        self._add_analysis_paragraph(doc, analysis_text)

   
    # =========================================================================
    # PAGE 2 – Section II: ICU results (trend chart, germs chart, cases table)
    # =========================================================================

    def _build_page2(self, doc, stats, chart_paths, quarter='', year='', history=None):
        """
        Page 2 (sections 1-1 and 1-2) + Page 3 (section 1-3):
          • Section title  II-  (dynamic quarter range)
          • 1-1  ICU trend chart + analysis placeholder
          • 1-2  ICU germs chart + analysis placeholder
          • [page break → page 3]
          • 1-3  ICU cases table  + analysis placeholder
          • page break
        Titles and analysis texts use the actual quarter/year passed in.
        """
        history     = list(history or [])
        cases_table = stats.get('icu_cases_table', [])
        year_s      = str(year)

        # ── Build Arabic quarter-range descriptions ───────────────────────
        # all_desc : last-5-history quarters + current  (used in II- and 1-1)
        recent5  = history[-5:] if len(history) >= 5 else list(history)
        all_qy   = [(e.get('quarter', ''), str(e.get('year', ''))) for e in recent5]
        all_qy.append((quarter, year_s))
        all_desc = self._describe_quarters(all_qy) if all_qy else f"{quarter} {year_s}"

        # two_desc : previous quarter + current  (used in 1-2)
        prev_list = history[-1:] if history else []
        two_qy    = [(e.get('quarter', ''), str(e.get('year', ''))) for e in prev_list]
        two_qy.append((quarter, year_s))
        two_desc  = self._describe_quarters(two_qy) if len(two_qy) > 1 else f"{quarter} {year_s}"

        # ── Main section title II ─────────────────────────────────────────
        self._add_p2_main_title(
            doc,
            f"نتيجة ال VAP خلال {all_desc} بحسب الأقسام:"
        )

        # ── 1-1 ──────────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"مقارنة نسبة ال VAP بين {all_desc} قسم ال ICU:"
        )
        self._add_chart_p2(doc, chart_paths.get('chart1_icu_trend'), "1")
        self._add_p2_analysis(
            doc,
            "إن الرسم البياني رقم 1 يظهر لنا أن نسبة ال VAP  في قسم ال ICU قد إنخفضت خلال "
            f"{quarter} من العام  {year_s} "
            "عن الفصل الثاني منه بنسبة 4.14 ‰ وعن الفصل الأول من العام نفسه بنسبة  8.77‰ "
            'وأيضاً" إنخفضت عن الفصل الرابع من العام 2024  بنسبة 0.47‰ وعن الفصل الثالث من العام نفسه  بنسبة ‰ 11.17\n'
            'النتيجة أعلاه جداً" مشجعة  كونها لم تتعدى النسبة المسموح بها وهي 25‰ '
            'وهناك تحسناً" ملحوظاً" مقارنة مع الفصول المذكورة أعلاه.\n'
            "لقد سجلت حالة  VAP واحدة في قسم ال ICU  (راجع جدول رقم 2)"
        )

        # ── 1-2 ──────────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"مقارنة نسبة ال VAP بين {two_desc} في قسم ال ICU مع ال Germs المسببة:"
        )
        self._add_chart_p2(doc, chart_paths.get('chart2_icu_germs'), "2")
        self._add_p2_analysis(
            doc,
            "إن الرسم البياني رقم 2 يظهر لنا نوع ال Germs التي ظهرت في ال DTA (Deep tracheal aspiration) culture  "
            f"في قسم ال ICU ونسبتهم خلال {two_desc} "
            f"حيث يتبين لنا أن ال     Acinetobacter baumanii & Klebsiella CRE لم تظهرا في {quarter} من العام {year_s}.\n"
            "لقد سجلت حالة VAP  واحدة سببها  ال Proteus ESBL أي ما يعادل مانسبته %100 من مجمل الحالات."
        )

        self._add_page_break(doc)

    # ── Page-2 title helpers ──────────────────────────────────────────────────

    def _add_p2_main_title(self, doc, text):
        """Section title – right-aligned, bold, underlined, 14 pt."""
        para = doc.add_paragraph()
        self._set_rtl(para)                        # sets bidi → auto right-justifies
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(3)
        run = para.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(14)
        run.bold      = True
        run.underline = True
        self._set_run_rtl(run)

    def _add_p2_sub_title(self, doc, text):
        """Sub-section title – right-aligned, bold, underlined, 14 pt."""
        para = doc.add_paragraph()
        self._set_rtl(para)                        # sets bidi → auto right-justifies
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(14)
        run.bold      = True
        run.underline = True
        self._set_run_rtl(run)

    def _add_chart_p2(self, doc, chart, label):
        """Embed chart at compact size (5.5") for page 2."""
        if chart is None:
            para = doc.add_paragraph(f"[ الرسم البياني {label} ]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            return

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run()
        try:
            if isinstance(chart, io.BytesIO):
                chart.seek(0)
                run.add_picture(chart, width=Inches(5.5))
            else:
                run.add_picture(str(chart), width=Inches(5.5))
        except Exception as exc:
            logger.warning(f"Could not embed chart {label}: {exc}")
            para.clear()
            para.add_run(f"[ الرسم البياني {label} – خطأ في التحميل ]")

    def _add_p2_analysis(self, doc, text):
        """Analysis paragraph – right-aligned, 10 pt, RTL, for pages 2+."""
        for line in text.split('\n'):
            para = doc.add_paragraph()
            self._set_rtl(para)                        # sets bidi → auto right-justifies
            run = para.add_run(line or '\u200f')
            run.font.name = FONT_ANALYSIS
            run.font.size = Pt(10)
            self._set_run_rtl(run)                     # <w:rtl> + w:cs font
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)

    def _add_ai_analysis(self, doc, _section_label: str, _context_data: dict,
                         fallback_text: str = ''):
        """
        Integration point for AI-generated analysis paragraphs.

        Currently renders ``fallback_text`` (the static Arabic text written for
        each section).  When AI integration is ready, replace the body with a
        call to AIAnalysisService and remove the fallback_text parameter.

        Args:
            doc           : python-docx Document
            section_label : Human-readable label, e.g. "1-1 ICU trend chart"
            context_data  : Dict of stats/context to be sent to the AI model
            fallback_text : Static Arabic text rendered until AI is wired in

        TODO: replace with AI call, e.g.:
            from app.services.ai_service import AIAnalysisService
            ai   = AIAnalysisService()
            text = ai.generate_analysis(_section_label, _context_data)
            self._add_p2_analysis(doc, text)
        """
        _ = (_section_label, _context_data)   # reserved for AI integration
        self._add_p2_analysis(doc, fallback_text)

    def _add_icu_cases_table_p2(self, doc, cases):
        """
        9-column RTL cases table for section 1-3.

        Visual column order (right → left, bidiVisual enabled so col 0 = rightmost):
          Risk factors | Age/Year | Germs | Date of infection | Date of intubation |
          Date of admission | Diagnosis | Case number | Nb of cases
        """
        HEADERS_AR = [
            "عوامل الخطر",       # col 0 → rightmost
            "العمر/السنة",
            "الجرثومة",
            "تاريخ الإصابة",
            "تاريخ التنبيب",
            "تاريخ الدخول",
            "التشخيص",
            "رقم الحالة",
            "عدد الحالات",       # col 8 → leftmost
        ]
        KEYS = [
            "risk_factors",
            "age",
            "germs",
            "date_of_infection",
            "date_of_intubation",
            "date_of_admission",
            "diagnosis",
            "case_number",
            "nb_of_cases",
        ]
        # Widths sum to ≈ 6.87" (usable page width)
        col_widths = [
            Inches(1.20),  # risk_factors  (trimmed to make room for nb_of_cases)
            Inches(0.55),  # age
            Inches(1.00),  # germs
            Inches(0.65),  # date_of_infection
            Inches(0.65),  # date_of_intubation
            Inches(0.65),  # date_of_admission
            Inches(1.00),  # diagnosis
            Inches(0.62),  # case_number
            Inches(0.55),  # nb_of_cases (widened so "عدد الحالات" fits on one line)
        ]

        n_data = len(cases) if cases else 1   # at least one (empty) data row
        n_cols = len(HEADERS_AR)
        table  = doc.add_table(rows=1 + n_data, cols=n_cols)
        table.autofit = False
        self._set_table_cell_margins(table, top=30, bottom=30, left=40, right=40)
        self._set_table_bidi(table)   # column 0 appears on the RIGHT

        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w

        # ── Header row ────────────────────────────────────────────────────
        for i, hdr in enumerate(HEADERS_AR):
            self._style_cell_text(table.cell(0, i), hdr, FONT, 10, bold=True, rtl=True)
            self._set_cell_shading(table.cell(0, i), SHADE_HEADER)
            self._set_cell_border(table.cell(0, i),
                top=BORDER_MEDIUM, bottom=BORDER_MEDIUM,
                left=BORDER_FINE, right=BORDER_FINE)

        # ── Data rows ─────────────────────────────────────────────────────
        if cases:
            for r_i, case in enumerate(cases):
                row_idx = 1 + r_i
                for c_i, key in enumerate(KEYS):
                    val = case.get(key, '\u2014')
                    val = str(val) if val is not None else '\u2014'
                    self._style_cell_text(table.cell(row_idx, c_i), val, FONT_EN, 8)
                    self._set_cell_border(table.cell(row_idx, c_i),
                        top=BORDER_FINE, bottom=BORDER_FINE,
                        left=BORDER_FINE, right=BORDER_FINE)
        else:
            # Empty placeholder row when no data yet
            for c_i in range(n_cols):
                self._style_cell_text(table.cell(1, c_i), '', FONT_EN, 8)
                self._set_cell_border(table.cell(1, c_i),
                    top=BORDER_FINE, bottom=BORDER_FINE,
                    left=BORDER_FINE, right=BORDER_FINE)

    def _build_page3(self, doc, stats, chart_paths, quarter='', year='', history=None):
        history    = list(history or [])
        cases_table = stats.get('icu_cases_table', [])
        year_s     = str(year)

        prev_list = history[-1:] if history else []
        two_qy    = [(e.get('quarter', ''), str(e.get('year', ''))) for e in prev_list]
        two_qy.append((quarter, year_s))
        two_desc  = self._describe_quarters(two_qy) if len(two_qy) > 1 else f"{quarter} {year_s}"

        # ── 1-3 title ────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"جدول مفصل عن حالات ال VAP في قسم ال ICU خلال {quarter} من العام {year_s}:"
        )

        # ── ICU cases table ───────────────────────────────────────────────
        self._add_icu_cases_table_p2(doc, cases_table)

        # ── 1-3 analysis ─────────────────────────────────────────────────
        self._add_p2_analysis(
            doc,
            f"إن الجدول رقم 2 يظهر لنا شرحاً مفصلاً عن حالة ال VAP المكتسبة في قسم ال ICU خلال {quarter} من العام {year_s} "
            "ويظهر لنا التشخيص عند الدخول, تاريخ دخول المريض إلى المستشفى , تاريخ وضع المريض على جهاز التنفس , "
            "تاريخ إلتقاط العدوى مع إسم الجرثومة المسببة بالإضافة إلى العوامل الصحية المساعدة لإلتقاط العدوى  \n"
            "من خلال هذا الجدول يتبين لنا أن :\n"
            "- المريض  عمره 92 سنة ولديه عدة عوامل صحية بالإضافة إلى إقامته الطويل في المستشفى.\n"
            "وهذه جميعه عوامل مساعدة تساهم في تعرض المريض لإلتقاط العدوى المكتسبة بالإضافة إلى عوامل أخرى مذكورة في الجدول أعلاه"
        )

        # ── 2-1 title ────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"مقارنة نسبة ال VAP بين {two_desc} في قسم ال CCU:"
        )

        # ── Chart 3 (CCU trend) ───────────────────────────────────────────
        self._add_chart_p2(doc, chart_paths.get('chart3_ccu_trend'), "3")

        self._add_p2_analysis(
            doc,
            f"إن الرسم البياني رقم 3 يظهر لنا أن نسبة ال VAP في قسم ال CCU قد انخفضت خلال {quarter} من العام {year_s} "
            "عن الفصل الثاني منه بنسبة 6.84‰،\n"
            "في حين ارتفعت عن الفصل الأول من العام نفسه والفصل الرابع من العام 2024 بنسبة 4.36‰،\n"
            "بينما انخفضت عن الفصل الثالث من العام 2024 بنسبة 3.51‰.\n\n"
            "النتيجة أعلاه مشجعة كونها لم تتعدَّ النسبة المسموح بها وهي 15‰."
        )
        self._add_page_break(doc)
  
    def _build_page4(self, doc, stats, chart_paths, quarter='', year='', history=None):
        history     = list(history or [])
        ccu_cases   = stats.get('ccu_cases_table', [])
        year_s      = str(year)

        prev_list = history[-1:] if history else []
        two_qy    = [(e.get('quarter', ''), str(e.get('year', ''))) for e in prev_list]
        two_qy.append((quarter, year_s))
        two_desc  = self._describe_quarters(two_qy) if len(two_qy) > 1 else f"{quarter} {year_s}"

        # ── 3-2 title ────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"مقارنة نسبة ال VAP بين {two_desc} في قسم ال ICN مع ال Germs المسببة:"
        )

        # ── Chart 4 (ICN germs) ───────────────────────────────────────────
        self._add_chart_p2(doc, chart_paths.get('chart4_ccu_germs'), "4")

        # ── Chart 4 analysis ─────────────────────────────────────────────
        self._add_p2_analysis(
            doc,
            "إن الرسم البياني رقم 4 يظهر لنا نوع ال Germs التي ظهرت في ال DTA (Deep tracheal aspiration) culture "
            f"في قسم ال ICN ونسبتهم خلال {two_desc} "
            "حيث يتبين لنا أن ال Xanthomonas maltophilia هي السبب في الفصلين الثاني والثالث من العام 2025.\n"
            "لقد سجلت حالة VAP واحدة سببها ال Xanthomonas maltophilia أي ما يعادل مانسبته %100 من مجمل الحالات."
        )

        # ── 2-3 title ────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"جدول مفصل عن حالات ال VAP في قسم ال CCU خلال {quarter} من العام {year_s}:"
        )

        # ── CCU cases table ───────────────────────────────────────────────
        self._add_icu_cases_table_p2(doc, ccu_cases)

        # ── CCU cases analysis ────────────────────────────────────────────
        self._add_p2_analysis(
            doc,
            f"إن الجدول رقم 3 يظهر لنا شرحاً مفصلاً عن حالة ال VAP المكتسبة في قسم ال CCU خلال {quarter} من العام {year_s} "
            "ويظهر لنا التشخيص عند الدخول, تاريخ دخول المريض إلى المستشفى , تاريخ وضع المريض على جهاز التنفس , "
            "تاريخ إلتقاط العدوى مع إسم الجرثومة المسببة بالإضافة إلى العوامل الصحية المساعدة لإلتقاط العدوى\n"
            "من خلال هذا الجدول يتبين لنا أن :\n"
            f"- المريض عمره 64 سنة ولديه عدة عوامل صحية مذكورة في الجدول أعلاه"
        )

        self._add_page_break(doc)

    def _build_page5(self, doc, stats, chart_paths, quarter='', year='', history=None):
        history   = list(history or [])
        icn_cases = stats.get('icn_cases_table', [])
        year_s    = str(year)

        prev_list = history[-1:] if history else []
        two_qy    = [(e.get('quarter', ''), str(e.get('year', ''))) for e in prev_list]
        two_qy.append((quarter, year_s))
        two_desc  = self._describe_quarters(two_qy) if len(two_qy) > 1 else f"{quarter} {year_s}"

        # ── 3-1 title ────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"مقارنة نسبة ال VAP بين {two_desc} في قسم ال ICN:"
        )

        # ── Chart 5 (ICN trend) ───────────────────────────────────────────
        self._add_chart_p2(doc, chart_paths.get('chart5_icn_trend'), "5")

        # ── Chart 5 analysis ─────────────────────────────────────────────
        self._add_p2_analysis(
            doc,
            f"إن الرسم البياني رقم 5 يظهر لنا أن نسبة ال VAP في قسم ال ICN قد انخفضت خلال {quarter} من العام {year_s} "
            "عن الفصل الثاني منه بنسبة 1.47‰ وعن الفصل الأول من العام نفسه بنسبة 7.52‰ "
            "بينما ارتفعت عن الفصل الرابع من العام 2024 بنسبة 4.78‰ "
            "في حين انخفضت عن الفصل الثالث من العام نفسه بنسبة 0.86‰\n"
            "النتيجة أعلاه مقبولة كونها لم تتعدَّ النسبة المسموح بها وهي 10‰\n"
            f"لقد سجلت حالة VAP في قسم ال ICN خلال هذا الفصل (راجع الجدول رقم 5)"
        )
        # ── 3-2 title ────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"مقارنة نسبة ال VAP بين {two_desc} في قسم ال ICN مع ال Germs المسببة:"
        )

        # ── Chart 6 (ICN germs) ───────────────────────────────────────────
        self._add_chart_p2(doc, chart_paths.get('chart6_icn_germs'), "6")

        # ── Chart 6 analysis ─────────────────────────────────────────────
        self._add_p2_analysis(
            doc,
            "إن الرسم البياني رقم 6 يظهر لنا نوع ال Germs التي ظهرت في ال DTA (Deep tracheal aspiration) culture "
            f"في قسم ال ICN ونسبتهم خلال {two_desc} "
            "حيث يتبين لنا أن ال Xanthomonas maltophilia هي السبب في الفصلين الثاني والثالث من العام 2025.\n"
            "لقد سجلت حالة VAP واحدة سببها ال Xanthomonas maltophilia أي ما يعادل مانسبته %100 من مجمل الحالات."
        )

        self._add_page_break(doc)

    def _build_page6(self, doc, stats, chart_paths, quarter='', year='', history=None):
        history   = list(history or [])
        icn_cases = stats.get('icn_cases_table', [])
        year_s    = str(year)

        # ── 3-3 title ────────────────────────────────────────────────────
        self._add_p2_sub_title(
            doc,
            f"جدول مفصل عن حالات ال VAP في قسم ال ICN خلال {quarter} من العام {year_s}:"
        )

        # ── ICN cases table ───────────────────────────────────────────────
        self._add_icu_cases_table_p2(doc, icn_cases)

        # ── ICN cases analysis ────────────────────────────────────────────
        self._add_p2_analysis(
            doc,
            f"إن الجدول رقم 4 يظهر لنا شرحاً مفصلاً عن حالة ال VAP المكتسبة في قسم ال ICN خلال {quarter} من العام {year_s} "
            "ويظهر لنا التشخيص عند الدخول, تاريخ دخول المريض إلى المستشفى , تاريخ وضع المريض على جهاز التنفس , "
            "تاريخ إلتقاط العدوى مع إسم الجرثومة المسببة بالإضافة إلى العوامل الصحية المساعدة لإلتقاط العدوى:\n"
            "- تشوهات خلقية في القلب\n"
            "- البقاء لفترة طويلة في القسم على جهاز التنفس الاصطناعي\n"
            "هذه كلها عوامل تساهم في التقاط العدوى المكتسبة في المستشفى.\n\n"
            "*النتيجة في أقسام ال Pediatric + CSU + ITU جاءت جداً مشجعة في جميع الفصول المذكورة أعلاه 0‰"
        )


    # =========================================================================
    # LAST PAGE: Final result, action tables, approval
    # =========================================================================

    def _build_last_page(self, doc):
        doc.add_page_break()
        self._add_spacer(doc, space_before=6)

        self._add_titled_row(doc, "النتيجة النهائية", shade=SHADE_HEADER, border=BORDER_MEDIUM)
        self._add_titled_row(doc, "Favorable", shade=None, border=BORDER_FINE)

        self._add_spacer(doc, space_before=6)
        self._add_checkbox_question_box(doc)
        self._add_spacer(doc, space_before=6)
        self._add_previous_actions_table(doc)
        self._add_spacer(doc, space_before=6)
        self._add_current_actions_table(doc)

        p_title = doc.add_paragraph()
        self._set_rtl(p_title)
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after  = Pt(4)
        p_title.paragraph_format.line_spacing = 1
        run_t = p_title.add_run('يجب ان تكون الاجراءات الواجب اتخاذها محددة وقابلة للقياس')
        run_t.font.name  = FONT_ANALYSIS
        run_t.font.size  = Pt(14)
        run_t.bold       = True
        run_t.underline  = True

        self._add_approval_table(doc)

        p_note = doc.add_paragraph()
        self._set_rtl(p_note)
        p_note.paragraph_format.space_before = Pt(4)
        p_note.paragraph_format.space_after  = Pt(0)
        p_note.paragraph_format.line_spacing = 1
        run_note = p_note.add_run(
            'يجب توقيع جميع المعنيين بمتابعة الاجراءات المتخذة او من ينوب عنهم في خانة الموافقة '
            'اضافة الى مدراء الادارات المعنيين'
        )
        run_note.font.name = FONT_ANALYSIS
        run_note.font.size = Pt(11)
        run_note.bold      = True

    # =========================================================================
    # LAST PAGE COMPONENTS (identical to mortality)
    # =========================================================================

    def _add_checkbox_question_box(self, doc):
        box_tbl = doc.add_table(rows=1, cols=1)
        box_tbl.autofit = False
        box_tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(box_tbl, top=30, bottom=30, left=60, right=60)

        cell = box_tbl.cell(0, 0)
        self._set_cell_border(cell, top=BORDER_FINE, bottom=BORDER_FINE,
                              left=BORDER_FINE, right=BORDER_FINE)

        p = cell.paragraphs[0]
        self._set_rtl(p)
        self._zero_spacing(p)
        self._add_run(p, 'هل تم اتخاذ اجراءات تحسينية سابقا حول موضوع التحليل :      ', FONT_ANALYSIS, 11)
        self._add_run(p, 'نعم', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐', None, 11)
        self._add_run(p, '         ', None, 11)
        self._add_run(p, 'كلا  ', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐ ', None, 11)

        p2 = cell.add_paragraph()
        self._set_rtl(p2)
        self._zero_spacing(p2)
        self._add_run(p2,
            'في حال نعم , يجب تعبئة جدول " الاجراءات المتخذة سابقا" و في حال كلا , الانتقال الى جدول "الاجراءات المتخذة الحالية"',
            FONT_ANALYSIS, 11)

    def _add_previous_actions_table(self, doc):
        table = doc.add_table(rows=3, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 4))
        self._style_cell_text(table.cell(0, 0), 'الاجراءات المتخذة  السابقة',
                              FONT, 10, bold=True, rtl=True)

        headers = ['الرقم', 'الاجراءات المقترحة سابقا', 'التاريخ المتوقع لتنفيذ الاجراء',
                   'تم التنفيذ )نعم/كلا(', 'في حال كلا ) الاجراءات التصحيحية الجديدة(']
        for j, h in enumerate(headers):
            self._style_cell_text(table.cell(1, j), h, FONT, 11, bold=True, rtl=True)

        for j in range(4):
            table.cell(2, j).text = ''

        last_cell = table.cell(2, 4)
        last_cell.text = ''
        p1 = last_cell.paragraphs[0]
        self._add_run(p1, ' ☐  الغاء الاجراء', FONT, 11)
        self._zero_spacing(p1)
        self._set_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = last_cell.add_paragraph()
        self._add_run(p2, ' ☐ اضافة الاجراء ضمن الاجراءات المتخذة الحالية.)الجدول ادناه(', FONT, 11)
        self._zero_spacing(p2)
        self._set_rtl(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(2):
            self._shade_row_cells(table.rows[i].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(1.8), Inches(1.37), Inches(1.1), Inches(2.1)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_current_actions_table(self, doc):
        table = doc.add_table(rows=6, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 3))
        self._style_cell_text(table.cell(0, 0), 'الاجراءات المتخذة  الحالية',
                              FONT, 10, bold=True, rtl=True)

        headers = ['الرقم', 'الاجراءات الواجب اتخاذها لحل الموضوع',
                   'الشخص المعني بالمتابعة', 'التاريخ المتوقع للتنفيذ']
        for j, h in enumerate(headers):
            self._style_cell_text(table.cell(1, j), h, FONT, 11, bold=True, rtl=True)

        for i in range(2, 6):
            for j in range(4):
                table.cell(i, j).text = ''

        for i in range(2):
            self._shade_row_cells(table.rows[i].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(3.07), Inches(1.8), Inches(1.5)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_approval_table(self, doc):
        table = doc.add_table(rows=4, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).text = ''
        self._style_cell_text(table.cell(0, 1), 'اعداد', FONT, 11, bold=True, rtl=True)
        table.cell(0, 2).merge(table.cell(0, 4))
        self._style_cell_text(table.cell(0, 2), 'موافقة', FONT, 11, bold=True, rtl=True)

        for i, label in enumerate(['الاسم', 'التوقيع', 'التاريخ']):
            self._style_cell_text(table.cell(i + 1, 0), label, FONT, 11, bold=True, rtl=True)
            for j in range(1, 5):
                table.cell(i + 1, j).text = ''

        self._shade_row_cells(table.rows[0].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.8), Inches(1.5), Inches(1.52), Inches(1.52), Inches(1.53)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w


    # =========================================================================
    # ANALYSIS HELPERS
    # =========================================================================

    def _add_analysis_paragraph(self, doc, text):
        """Analysis paragraph – right-aligned, 11 pt, RTL, for page 1."""
        for line in text.split('\n'):
            para = doc.add_paragraph()
            self._set_rtl(para)                        # sets bidi → auto right-justifies
            run = para.add_run(line or '\u200f')
            run.font.name = FONT_ANALYSIS
            run.font.size = Pt(11)
            self._set_run_rtl(run)                     # <w:rtl> + w:cs font
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)

    # =========================================================================
    # UTILITY HELPERS
    # =========================================================================

    def _add_spacer(self, doc, space_before=4):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(0)

    def _add_titled_row(self, doc, text, shade=None, border=BORDER_FINE):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(tbl,
            top=0 if shade else 20, bottom=0 if shade else 20, left=20, right=20)
        cell = tbl.cell(0, 0)
        if shade:
            self._set_cell_shading(cell, shade)
        self._set_cell_border(cell, top=border, bottom=border, left=border, right=border)
        self._style_cell_text(cell, text, FONT, 11, bold=True, rtl=True)

    def _add_run(self, para, text, font_name, size, bold=False):
        run = para.add_run(text)
        if font_name:
            run.font.name = font_name
        run.font.size = Pt(size)
        run.bold      = bold
        return run
    
    def _shade_row_cells(self, cells, color):
        for cell in cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            shading.set(qn('w:val'),  'clear')
            cell._element.get_or_add_tcPr().append(shading)

    def _apply_borders_to_table(self, table, border_style):
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, top=border_style, bottom=border_style,
                                      left=border_style, right=border_style)

    def _set_table_rtl(self, table):
        bidi = OxmlElement('w:bidiVisual')
        bidi.set(qn('w:val'), '1')
        table._tbl.tblPr.append(bidi)

    def _add_page_break(self, doc):
        para = doc.add_paragraph()
        run  = para.add_run()
        br   = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._element.append(br)

    def _zero_spacing(self, para):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)

    def _set_rtl(self, para):
        # Use python-docx abstraction – properly places <w:bidi/> in schema order
        para.paragraph_format.right_to_left = True
        pPr = para._element.get_or_add_pPr()
        if pPr.find(qn('w:bidi')) is None:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

    def _add_rtl_run(self, para, text, bold=False, size=10):
        run = para.add_run(text)
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)

        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        return run
    def _style_cell_text(self, cell, text, font_name, font_size,
                         bold=False, italic=False, rtl=False,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER):
        para = cell.paragraphs[0]
        para.clear()
        run  = para.add_run(text)
        run.font.name  = font_name
        run.font.size  = Pt(font_size)
        run.bold       = bold
        run.italic     = italic
        para.alignment = alignment
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        if rtl:
            self._set_rtl(para)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _set_cell_shading(self, cell, fill_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_color)
        tcPr.append(shd)

    def _set_cell_border(self, cell, top=None, bottom=None, left=None, right=None):
        tc        = cell._tc
        tcPr      = tc.get_or_add_tcPr()
        # Remove every existing tcBorders element so we never accumulate duplicates
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = OxmlElement('w:tcBorders')
        for side, attrs in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            if attrs:
                el = OxmlElement(f'w:{side}')
                for key, val in attrs.items():
                    el.set(qn(f'w:{key}'), val)
                tcBorders.append(el)
        tcPr.append(tcBorders)

    def _set_tc_border(self, tc, top=None, bottom=None, left=None, right=None):
        tcPr      = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = OxmlElement('w:tcBorders')
        for side, attrs in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            if attrs:
                el = OxmlElement(f'w:{side}')
                for key, val in attrs.items():
                    el.set(qn(f'w:{key}'), val)
                tcBorders.append(el)
        tcPr.append(tcBorders)

    def _set_run_rtl(self, run):
        """Configure a run for correct Arabic RTL rendering:
          • <w:rtl val="1"/>   – run direction
          • w:rFonts w:cs      – CS font (Arabic chars use CS slot, not ascii/hAnsi)
          • <w:bCs/>           – CS bold  (run.bold sets <w:b/> for Latin; Arabic
                                 needs <w:bCs/> to actually appear bold)
        """
        rPr = run._element.get_or_add_rPr()
        # 1. RTL direction
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        # 2. Mirror regular font → CS font
        try:
            cs_font = run.font.name
            if cs_font:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:cs'), cs_font)
        except Exception:
            pass
        # 3. If the run is bold, also set CS bold so Arabic glyphs are bold
        try:
            if run.bold:
                bCs = OxmlElement('w:bCs')
                rPr.append(bCs)
        except Exception:
            pass

    def _set_table_bidi(self, table):
        """Add <w:bidiVisual/> so column 0 appears on the RIGHT (RTL table)."""
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        bidi = OxmlElement('w:bidiVisual')
        tblPr.append(bidi)

    def _set_table_cell_margins(self, table, top=0, bottom=0, left=0, right=0):
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblCellMar = OxmlElement('w:tblCellMar')
        for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'),    str(val))
            el.set(qn('w:type'), 'dxa')
            tblCellMar.append(el)
        tblPr.append(tblCellMar)