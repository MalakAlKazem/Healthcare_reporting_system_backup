"""
DOCX Report Generator for Healthcare Mortality System
=====================================================
Generates a multi-page Arabic RTL DOCX report with embedded Matplotlib charts.
Uses python-docx with raw XML manipulation for fine-grained control over
borders, shading, RTL layout, and cell merging.

Pages:
  1. Metadata table, results table, analysis box with chart1
  2. Building distribution (chart2), admission sources (chart3), age distribution (chart5)
  3. Age vs quarter (chart6), department pie (chart7) + data table, department analysis
  4. Department comparison (chart8), WHO age category (chart9)
  5. WHO diagnosis (chart10), doctor specialty table
  6. Final result, previous/current action tables, approval table
"""

import re
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from loguru import logger
from app.mortality.chart_generator import matplotlib_chart_generator
from app.mortality.ai_service import ai_service
# =============================================================================
# BIDI CONTROL CHARACTERS (defined as escaped strings — never embed literally)
# LRE/PDF wrap parentheses to force correct direction in RTL Word paragraphs
# =============================================================================
LRE = '\u202a'   # Left-to-Right Embedding  — opens  LTR scope
PDF = '\u202c'   # Pop Directional Formatting — closes LTR scope
RLM = '\u200f'   # Right-to-Left Mark        — anchors RTL runs



# =============================================================================
# FONTS
# =============================================================================
FONT          = 'Traditional Arabic'
FONT_EN       = 'Calibri'
FONT_ANALYSIS = 'Frutiger LT Arabic 45 Light'

# =============================================================================
# LAYOUT
# =============================================================================
PAGE_WIDTH   = Inches(8.27)
PAGE_HEIGHT  = Inches(11.69)
MARGIN       = Inches(0.7)
USABLE_WIDTH = Inches(6.87)

# =============================================================================
# BORDER PRESETS
# =============================================================================
BORDER_THICK  = {'val': 'single', 'sz': '18', 'color': '000000', 'space': '0'}
BORDER_MEDIUM = {'val': 'single', 'sz': '12', 'color': '000000', 'space': '0'}
BORDER_THIN   = {'val': 'single', 'sz': '8',  'color': '000000', 'space': '0'}
BORDER_FINE   = {'val': 'single', 'sz': '4',  'color': '000000', 'space': '0'}
BORDER_NONE   = {'val': 'none',   'sz': '0',  'color': 'FFFFFF', 'space': '0'}

# =============================================================================
# SHADING
# =============================================================================
SHADE_HEADER = 'EDEBE3'
SHADE_TABLE  = 'D9E2F3'
SHADE_TOTAL  = 'F2F2F2'


class MatplotlibDocxGenerator:

    def __init__(self):
        self.logo_path = 'assets/LOGO.png'

    # =========================================================================
    # PUBLIC API
    # =========================================================================

    async def generate_report(self, data, history=None, options=None):
        options  = options or {}
        quarter  = options.get('quarter', 'الفصل الثالث')
        year     = options.get('year', '2025')
        history  = history or []
        logger.info(f"Generating DOCX report for {quarter} {year}...")

        stats             = data.get('statistics', {})
        mortality_metrics = stats.get('mortality_metrics', {})
        all_records  = data.get('records', [])
        non_kpi_records = [
            r for r in all_records
            if str(r.get('include_kpi', 'YES')).strip().upper() != 'YES'
        ]

        current_stats = {
            'quarter':            quarter,
            'year':               year,
            'mortality_rate':     mortality_metrics.get('rate', 0),
            'mortality_metrics':  mortality_metrics,
            'rate':               mortality_metrics.get('rate', 0),
            'total_deaths':       stats.get('total_deaths', 0),
            'deaths':             stats.get('total_deaths', 0),
            'kpi_deaths':         stats.get('kpi_deaths', stats.get('total_deaths', 0)),
            'total_patients':     stats.get('total_patients', 0),
            'buildings':          stats.get('buildings', {}),
            'demographics':       stats.get('demographics', {}),
            'clinical':           stats.get('clinical', {}),
            'departments':        stats.get('departments', []),
            'specialties':        stats.get('specialties', []),
            'who_categories':     data.get('who_categories', []),
            'who_categories_kpi': stats.get('who_categories_kpi', {}),
            'non_kpi_deaths':     non_kpi_records,
        }

        logger.info("Generating Matplotlib charts...")
        charts = matplotlib_chart_generator.generate_all_charts(current_stats, history)
        logger.info(f"Generated {len(charts)} charts")

        doc     = Document()
        section = doc.sections[0]
        self._setup_section(section)

        # FIX 1+2+3: history now flows into every page that needs it
        self._build_page1(doc, quarter, year, current_stats, charts, history)
        self._build_page2(doc, charts, current_stats, history)
        self._build_page3(doc, charts, current_stats, history)
        self._build_page4(doc, charts, current_stats, history)
        self._build_page5(doc, charts, current_stats)
        self._build_non_kpi_page(doc, current_stats)
        self._build_last_page(doc, current_stats)

        reports_dir = os.path.normpath(
            os.path.join(os.path.dirname(__file__), '..', '..', 'storage', 'reports')
        )
        os.makedirs(reports_dir, exist_ok=True)
        file_name = f'mortality-rate report {quarter} {year}.docx'
        file_path = os.path.join(reports_dir, file_name)
        doc.save(file_path)
        logger.success(f"DOCX report saved: {file_path}")
        return {'filePath': file_path, 'fileName': file_name}

    # =========================================================================
    # SECTION SETUP
    # =========================================================================

    def _setup_section(self, section):
        section.page_width    = PAGE_WIDTH
        section.page_height   = PAGE_HEIGHT
        section.top_margin    = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin   = MARGIN
        section.right_margin  = MARGIN

        sectPr    = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   'single')
            border.set(qn('w:sz'),    '12')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        sectPr.append(pgBorders)

        self._add_header(section)
        self._add_footer(section)

    def _add_header(self, section):
        header = section.header
        for element in list(header._element):
            header._element.remove(element)

        table          = header.add_table(rows=1, cols=3, width=Inches(7))
        self._set_table_cell_margins(table, top=0, bottom=0, left=0, right=0)
        cells          = table.rows[0].cells
        cells[0].width = Inches(1.0)
        cells[1].width = Inches(5.0)
        cells[2].width = Inches(1.0)

        if os.path.exists(self.logo_path):
            para = cells[0].paragraphs[0]
            para.add_run().add_picture(self.logo_path, width=Inches(0.5))
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._zero_spacing(para)

        para     = cells[1].paragraphs[0]
        run      = para.add_run("نموذج تحليل البيانات")
        run.font.name = FONT_ANALYSIS
        run.font.size = Pt(16)
        run.bold      = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(para)
        self._zero_spacing(para)

        para     = cells[2].paragraphs[0]
        run      = para.add_run(f'QS-36F-103(4)\n{datetime.now().strftime("%d/%m/%Y")}')
        run.font.name = FONT_EN
        run.font.size = Pt(7)
        run.bold      = True
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._zero_spacing(para)

        line_para = header.add_paragraph()
        line_para.paragraph_format.space_before = Pt(0)
        line_para.paragraph_format.space_after  = Pt(8)
        line_para.paragraph_format.line_spacing = Pt(1)
        line_para.add_run("").font.size = Pt(1)
        pPr    = line_para._element.get_or_add_pPr()
        pBdr   = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_footer(self, section):
        footer = section.footer
        for element in list(footer._element):
            footer._element.remove(element)

        para = footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._element.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)
        run.font.size = Pt(10)
        run.font.name = FONT_EN

    # =========================================================================
    # PAGE 1
    # =========================================================================

    def _build_page1(self, doc, quarter, year, stats, charts, history):
        self._add_spacer(doc, space_before=6)
        self._add_metadata_table(doc, quarter, year)
        doc.add_paragraph()
        self._add_results_table(doc, stats, history)
        self._add_spacer(doc)
        # FIX 3: pass history so AI can compare with previous quarter
        self._add_analysis_box(doc, quarter, year, stats, charts, history)

    # =========================================================================
    # PAGE 2  — FIX 1: added history=None
    # =========================================================================

    def _build_page2(self, doc, charts, stats, history=None):
        history = history or []
        doc.add_page_break()

        if 'chart2' in charts:
            logger.info("Embedding Chart 2: Building Distribution...")
            self._add_rtl_para(doc, '- توزيع الوفيات بحسب المبنى', 9, bold=True, font=FONT_ANALYSIS)
            self._add_centered_chart(doc, charts['chart2'], width=4)

            buildings  = stats.get('buildings', {})
            bci        = buildings.get('bci', {})
            rah        = buildings.get('rah', {})
            bci_pct    = bci.get('percentage', 0)
            rah_pct    = rah.get('percentage', 0)
            bci_rate   = bci.get('rate', 0)
            rah_rate   = rah.get('rate', 0)
            rah_deaths = rah.get('deaths', 0)

            bci_deaths = bci.get('deaths', 0)

            # FIX 3: pass bci_deaths; fallback is now a complete static template
            building_text = ai_service.analyze_building_distribution(
                bci_pct, rah_pct, bci_rate, rah_rate, rah_deaths,
                bci_deaths=bci_deaths
            )
            if not building_text:
                # Static template — always includes both counts AND patient rates
                bci_rate_part = (
                    f"، وشكّلت نسبة الوفيات بين مرضاه {bci_rate:.1f}%"
                    if bci_rate > 0 else ""
                )
                rah_rate_part = (
                    f"، وشكّلت نسبة الوفيات بين مرضاه {rah_rate:.1f}%"
                    if rah_rate > 0 else ""
                )
                building_text = (
                    f"بالرغم من أن نسبة الوفيات في مركز القلب بلغت {bci_pct:.0f}% من إجمالي الوفيات "
                    f"{LRE}({bci_deaths} وفاة){PDF}{bci_rate_part}، إلا أن المبنى العام سجّل "
                    f"النسبة الأعلى بـ {rah_pct:.0f}% من الإجمالي "
                    f"{LRE}({rah_deaths} وفاة){PDF}{rah_rate_part}."
                )
            self._add_analysis_text(doc, building_text)
            logger.info("Chart 2 embedded")

        if 'chart3' in charts:
            logger.info("Embedding Chart 3: Admission Sources...")
            self._add_rtl_para(doc, 'توزيع الوفيات بحسب وجهة الدخول', 9, bold=True, font=FONT_ANALYSIS)
            self._add_centered_chart(doc, charts['chart3'], width=5.5)
            logger.info("Chart 3 embedded")

        if 'chart5' in charts:
            logger.info("Embedding Chart 5: Age Distribution...")
            self._add_rtl_para(doc, 'توزيع الوفيات بحسب عمر المتوفي', 9, bold=True, font=FONT_ANALYSIS)
            self._add_centered_chart(doc, charts['chart5'], width=5)
            logger.info("Chart 5 embedded")

    # =========================================================================
    # PAGE 3  — FIX 1: added history=None  |  FIX 4: added missing add_analysis_text
    # =========================================================================

    def _build_page3(self, doc, charts, stats, history=None):
        history = history or []
        doc.add_page_break()

        if 'chart6' in charts:
            logger.info("Embedding Chart 6: Age vs Quarter...")
            self._add_rtl_para(doc, 'مقارنة الوفيات بحسب العمر و الفصل', 11, bold=True, font=FONT_ANALYSIS)
            self._add_centered_chart(doc, charts['chart6'], width=6.5)

            age_groups_ar = [
                'اقل من 5 سنوات', 'من 5 الى 15 سنة', 'من 16 الى 30 سنة',
                'من 31 الى 50 سنة', 'من 51 الى 60 سنة', 'من 61 الى 70 سنة',
                'من 71 الى 80 سنة', 'اكثر من 81 سنة'
            ]
            curr_age_cats = stats.get('demographics', {}).get('age_categories', [])
            curr_ages     = [a.get('count', 0) for a in curr_age_cats[:8]] if curr_age_cats else [0] * 8
            prev_hist     = history[-1] if history else {}
            prev_ages     = prev_hist.get('age_groups', [0] * 8)
            if len(prev_ages) < 8:
                prev_ages = list(prev_ages) + [0] * (8 - len(prev_ages))

            age_changes = [
                {'group': g, 'prev': prev_ages[i], 'current': curr_ages[i] if i < len(curr_ages) else 0}
                for i, g in enumerate(age_groups_ar)
            ]
            prev_q_label = f"{prev_hist.get('quarter', '')} {prev_hist.get('year', '')}"
            curr_q_label = f"{stats.get('quarter', '')} {stats.get('year', '')}"

            age_text = ai_service.analyze_age_by_quarter(curr_q_label, prev_q_label, age_changes)
            if not age_text:
                age_text = "يُظهر الرسم البياني مقارنة توزيع الوفيات بحسب الفئات العمرية خلال الفصول."
            self._add_analysis_text(doc, age_text)
            logger.info("Chart 6 embedded")

        departments = stats.get('departments', [])
        if departments:
            logger.info("Embedding Department section...")
            self._add_rtl_para(doc, '- توزيع الوفيات بحسب الاقسام', 11, bold=True, font=FONT_ANALYSIS)
            chart7_img = charts.get('chart7')
            if chart7_img:
                self._add_dept_chart_with_table(doc, chart7_img, departments)
            else:
                self._add_dept_table_only(doc, departments)

            total_deaths = stats.get('total_deaths', 0)
            icu_keywords = ['icu', 'ccu', 'csu', 'itcu', 'icn', 'picu', 'tcu', 'icvu']
            er_keywords  = ['er', 'emergency', 'طوارئ']

            def _is_icu(n): return any(k in n.lower() for k in icu_keywords)
            def _is_er(n):  return any(k in n.lower() for k in er_keywords) and not _is_icu(n)

            dept_total = sum(d.get('count', 0) for d in departments)
            icu_d  = sum(d['count'] for d in departments if _is_icu(d.get('name', '')))
            er_d   = sum(d['count'] for d in departments if _is_er(d.get('name', '')))
            ward_d = dept_total - icu_d - er_d

            clinical      = stats.get('clinical', {})
            er_details    = None
            ward_details  = None

            er_patients = clinical.get('er_inpatient_deaths', [])
            if er_patients:
                los_values = [p.get('los', 0) for p in er_patients if p.get('los')]
                er_details = {
                    'count':       len(er_patients),
                    'los_min':     min(los_values) if los_values else None,
                    'los_max':     max(los_values) if los_values else None,
                    'specialties': list({p.get('specialty', '') for p in er_patients if p.get('specialty')})
                }

            ward_patients = clinical.get('ward_deaths', [])
            if ward_patients:
                ward_details = [
                    {'dept': p.get('department', ''), 'los': p.get('los', ''), 'specialty': p.get('specialty', '')}
                    for p in ward_patients
                ]

            dept_text = ai_service.analyze_departments(
                total_deaths, icu_d, er_d, ward_d, departments[:5],
                er_details=er_details,
                ward_details=ward_details
            )
            if not dept_text:
                dept_text = (
                    f"بلغ عدد الوفيات في أقسام العناية المركزة {icu_d} حالة، "
                    f"وفي قسم الطوارئ {er_d} حالة "
                    f"{LRE}(مرضى مقيمون){PDF}، "
                    f"وفي الأقسام الاستشفائية {ward_d} حالة."
                )
            self._add_analysis_text(doc, dept_text)
            logger.info("Department section embedded")

    # =========================================================================
    # PAGE 4  — FIX 1: added history=None
    # =========================================================================

    def _build_page4(self, doc, charts, stats, history=None):
        history = history or []
        doc.add_page_break()

        if 'chart8' in charts:
            logger.info("Embedding Chart 8: Department Comparison...")
            self._add_rtl_para(doc, 'مقارنة عدد الوفيات بحسب الاقسام', 11, bold=True, font=FONT_ANALYSIS)
            self._add_centered_chart(doc, charts['chart8'], width=6.5)

            curr_depts     = {d['name']: d['count'] for d in stats.get('departments', [])}
            prev_h         = history[-1] if history else {}
            prev_depts     = prev_h.get('departments', {})
            all_dept_names = set(curr_depts.keys()) | set(prev_depts.keys())
            dept_changes   = sorted(
                [
                    {'name': n, 'prev': prev_depts.get(n, 0), 'current': curr_depts.get(n, 0)}
                    for n in all_dept_names
                ],
                key=lambda x: abs(x['current'] - x['prev']),
                reverse=True
            )
            curr_q = f"{stats.get('quarter', '')} {stats.get('year', '')}"
            prev_q = f"{prev_h.get('quarter', '')} {prev_h.get('year', '')}"

            chart8_text = ai_service.analyze_dept_comparison(curr_q, prev_q, dept_changes)
            if not chart8_text:
                chart8_text = "يُظهر الرسم البياني مقارنة عدد الوفيات بحسب الأقسام خلال الفصول المختلفة."
            self._add_analysis_text(doc, chart8_text)
            logger.info("Chart 8 embedded")

        if 'chart9' in charts:
            logger.info("Embedding Chart 9: WHO Category Comparison...")
            self._add_rtl_para(
                doc,
                'توزيع الوفيات بحسب التشخيص للسبب الذي نجم عنه الوفاة بحسب الفصل',
                11, bold=True, font=FONT_ANALYSIS
            )
            self._add_centered_chart(doc, charts['chart9'], width=6.5)

            who_kpi     = stats.get('who_categories_kpi', {})
            top_who     = sorted(
                [{'category': k, 'count': v} for k, v in who_kpi.items()],
                key=lambda x: x['count'], reverse=True
            )[:4]
            curr_q_who  = f"{stats.get('quarter', '')} {stats.get('year', '')}"
            # FIX 5: static WHO text — AI hallucinated wrong Arabic disease names
            chart9_text = self._build_who_comparison_text(curr_q_who, top_who)
            self._add_analysis_text(doc, chart9_text)
            logger.info("Chart 9 embedded")

    # =========================================================================
    # PAGE 5
    # =========================================================================

    def _build_page5(self, doc, charts, stats):
        doc.add_page_break()
        quarter     = stats.get('quarter', '')
        specialties = stats.get('specialties', [])

        if 'chart10' in charts:
            logger.info("Embedding Chart 10: WHO Diagnosis...")
            self._add_rtl_para(
                doc,
                'توزيع الوفيات بحسب التشخيص للسبب الذي نجم عنه الوفاة',
                11, bold=True, font=FONT_ANALYSIS
            )
            who_kpi_p5 = stats.get('who_categories_kpi', {})
            if who_kpi_p5:
                top_cat = max(who_kpi_p5, key=who_kpi_p5.get)
                top_cnt = who_kpi_p5[top_cat]
            else:
                top_cat, top_cnt = 'cardiovascular disease', 0

            # FIX 5b: static WHO text with correct Arabic names
            who_text = self._build_who_diagnosis_text(quarter, stats.get('year',''), top_cat, top_cnt)
            self._add_analysis_text(doc, who_text)
            self._add_centered_chart(doc, charts['chart10'], width=6.5)
            logger.info("Chart 10 embedded")

        self._add_rtl_para(doc, 'توزيع الوفيات بحسب اختصاص الطبيب المعالج', 11, bold=True, font=FONT_ANALYSIS)
        self._add_doctor_specialty_table(doc, specialties)

    # =========================================================================
    # NON-KPI PAGE  (deaths within 24 hours — excluded from indicator)
    # =========================================================================

    def _build_non_kpi_page(self, doc, stats):
        non_kpi = stats.get('non_kpi_deaths', [])
        doc.add_page_break()
        self._add_spacer(doc, space_before=6)

        # Title row
        self._add_titled_row(
            doc,
            'وفيات لا تدخل ضمن هذا المؤشر لانها حصلت خلال 24 ساعة',
            shade=SHADE_HEADER,
            border=BORDER_MEDIUM,
        )
        self._add_spacer(doc, space_before=6)

        if not non_kpi:
            p = doc.add_paragraph()
            r = p.add_run('لا توجد وفيات خلال 24 ساعة')
            r.font.name = FONT
            r.font.size = Pt(11)
            self._set_rtl(p)
            self._zero_spacing(p)
            return

        # 11 columns — Arabic headers matching Excel column names
        # الإقامة uses length_of_stay_original to preserve the raw Excel value
        COLS = [
            ('رقم الملف',                                  'file_number'),
            ('إسم المريض',                                 'patient_name'),
            ('العمر',                                      'age'),
            ('وجهة الدخول',                               'admission_source'),
            ('تاريخ الدخول',                               'admission_date'),
            ('تاريخ الوفاة',                               'death_date'),
            ('التشخيص للسبب الذي نجم عنه الوفاة',         'underlying_cause_of_death'),
            ('القسم التمريضي',                             'nursing_department'),
            ('الإختصاص',                                   'specialty'),
            ('الطبيب المعالج',                             'treating_doctor'),
            ('الإقامة',                                    'length_of_stay_original'),
        ]
        n_cols = len(COLS)
        n_rows = len(non_kpi) + 1  # header + data

        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.autofit = False
        self._set_table_cell_margins(table, top=10, bottom=10, left=12, right=12)
        self._set_table_rtl(table)

        # Column widths — total exactly 6.87 in to fit inside page border
        col_widths = [
            Inches(0.52),  # رقم الملف
            Inches(0.92),  # إسم المريض
            Inches(0.33),  # العمر
            Inches(0.55),  # وجهة الدخول
            Inches(0.60),  # تاريخ الدخول
            Inches(0.60),  # تاريخ الوفاة
            Inches(1.25),  # التشخيص
            Inches(0.55),  # القسم التمريضي
            Inches(0.52),  # الإختصاص
            Inches(0.60),  # الطبيب المعالج
            Inches(0.43),  # الإقامة
        ]

        # Header row — light gold (SHADE_HEADER) matching document style
        self._shade_row_cells(table.rows[0].cells, SHADE_HEADER)
        for j, (hdr, _) in enumerate(COLS):
            self._style_cell_text(table.cell(0, j), hdr, FONT, 8, bold=True, rtl=True, align='center')
            self._set_cell_border(table.cell(0, j),
                                  top=BORDER_FINE, bottom=BORDER_FINE,
                                  left=BORDER_FINE, right=BORDER_FINE)

        # Data rows
        DATE_FIELDS = {'admission_date', 'death_date'}

        def _fmt_date(s):
            """Convert ISO date string to DD/MM/YYYY without time component."""
            import re as _re
            m = _re.match(r'(\d{4})-(\d{2})-(\d{2})', s)
            return f"{m.group(3)}/{m.group(2)}/{m.group(1)}" if m else s

        for i, rec in enumerate(non_kpi):
            row_idx = i + 1
            if i % 2 == 1:
                self._shade_row_cells(table.rows[row_idx].cells, SHADE_TOTAL)

            def _val(field, _rec=rec):
                if field == 'underlying_cause_of_death':
                    # First mapped occurrence, then any pandas-suffixed duplicates (.1, .2 …)
                    v = _rec.get('underlying_cause_of_death')
                    if not v or str(v).strip().lower() in ('nan', 'none', 'nat', ''):
                        for key in sorted(_rec.keys()):
                            if 'التشخيص للسبب الذي نجم عنه الوفاة' in str(key):
                                candidate = _rec[key]
                                if candidate and str(candidate).strip().lower() not in ('nan', 'none', 'nat', ''):
                                    v = candidate
                                    break
                else:
                    v = _rec.get(field)
                if v is None or (isinstance(v, float) and v != v):
                    return '—'
                s = str(v).strip()
                if not s or s.lower() in ('nan', 'none', 'nat'):
                    return '—'
                if field in DATE_FIELDS:
                    s = _fmt_date(s)
                return s

            values = [_val(field) for _, field in COLS]

            for j, val in enumerate(values):
                is_ar = any('؀' <= c <= 'ۿ' for c in val)
                font  = FONT if is_ar else FONT_EN
                self._style_cell_text(table.cell(row_idx, j), val, font, 8,
                                      rtl=is_ar, align='center')
                self._set_cell_border(table.cell(row_idx, j),
                                      top=BORDER_FINE, bottom=BORDER_FINE,
                                      left=BORDER_FINE, right=BORDER_FINE)

        # Apply widths
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w

    # =========================================================================
    # LAST PAGE
    # =========================================================================

    def _build_last_page(self, doc, stats=None):
        stats = stats or {}
        from app.config import load_targets as _load_targets
        rate   = float(stats.get('mortality_rate', 0) or 0)
        target = float(_load_targets().get('mortality', {}).get('rate', 2.0) or 2.0)
        result_text = "مشجعة" if rate < target else "غير مشجعة"

        doc.add_page_break()
        self._add_spacer(doc, space_before=6)
        self._add_titled_row(doc, "النتيجة النهائية", shade=SHADE_HEADER, border=BORDER_MEDIUM)
        self._add_titled_row(doc, result_text, shade=None, border=BORDER_FINE)
        self._add_spacer(doc, space_before=6)
        self._add_checkbox_question_box(doc)
        self._add_spacer(doc, space_before=6)
        self._add_previous_actions_table(doc)
        self._add_spacer(doc, space_before=6)
        self._add_current_actions_table(doc)

        p_title = doc.add_paragraph()
        self._set_rtl(p_title)
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after  = Pt(4)
        p_title.paragraph_format.line_spacing = 1
        run_t           = p_title.add_run('يجب ان تكون الاجراءات الواجب اتخاذها محددة وقابلة للقياس')
        run_t.font.name = FONT_ANALYSIS
        run_t.font.size = Pt(11)
        run_t.bold      = True
        run_t.underline = True

        self._add_approval_table(doc)

        p_note = doc.add_paragraph()
        self._set_rtl(p_note)
        p_note.paragraph_format.space_before = Pt(4)
        p_note.paragraph_format.space_after  = Pt(0)
        p_note.paragraph_format.line_spacing = 1
        run_note           = p_note.add_run(
            'يجب توقيع جميع المعنيين بمتابعة الاجراءات المتخذة او من ينوب عنهم في خانة الموافقة '
            'اضافة الى مدراء الادارات المعنيين'
        )
        run_note.font.name = FONT_ANALYSIS
        run_note.font.size = Pt(11)
        run_note.bold      = True

    # =========================================================================
    # PAGE 1 COMPONENTS
    # =========================================================================

    def _add_metadata_table(self, doc, quarter, year):
        table = doc.add_table(rows=3, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=0, bottom=0, left=30, right=30)

        col_w = [Inches(1.93), Inches(1.06), Inches(2.71), Inches(1.17)]
        for i, width in enumerate(col_w):
            for row in table.rows:
                row.cells[i].width = width

        # Row heights — generous minimum so content is never cut
        for ri in range(3):
            table.rows[ri].height      = Inches(0.35)
            table.rows[ri].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        # Row 0
        self._style_cell_text(table.cell(0, 3), "موضوع التحليل", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)
        self._style_cell_text(table.cell(0, 2), "Inpatient Mortality rate", FONT_EN, 8, bold=True)
        self._style_cell_text(table.cell(0, 1), "الإدارة", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(0, 0), "الطبية", FONT, 10, rtl=True)

        # Row 1
        self._style_cell_text(table.cell(1, 3), "", FONT, 9, rtl=True)
        self._set_cell_shading(table.cell(1, 3), SHADE_HEADER)
        self._style_cell_text(table.cell(1, 2), "", FONT_EN, 9)
        self._style_cell_text(table.cell(1, 1), "الوحدة الإدارية", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(1, 0), "---------------- ", FONT, 9, rtl=True)

        # Row 2
        self._style_cell_text(table.cell(2, 3), "مصادر البيانات) تحديد اسماء النماذج(", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(2, 3), SHADE_HEADER)
        self._style_cell_text(table.cell(2, 2), "سجل الوفيات", FONT, 9, rtl=True)
        self._style_cell_text(table.cell(2, 1), "الشهر / الفصل / العام", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(2, 1), SHADE_HEADER)
        self._style_cell_text(table.cell(2, 0), f"{quarter} {year} /", FONT, 10, rtl=True)

        # Vertical merges for cols 2 and 3 rows 0-1
        table.cell(0, 2).merge(table.cell(1, 2))
        table.cell(0, 3).merge(table.cell(1, 3))
        self._set_cell_shading(table.cell(0, 3), SHADE_HEADER)

        # Borders — all sides thick
        for r in range(3):
            for c in range(4):
                self._set_cell_border(table.cell(r, c),
                                      top=BORDER_THICK, bottom=BORDER_THICK,
                                      left=BORDER_THICK, right=BORDER_THICK)

        # Cols 2-3 in row 1 are vMerge continuation cells — set their borders directly
        row1_tcs = table.rows[1]._tr.findall(qn('w:tc'))
        for tc in row1_tcs[2:4]:
            self._set_tc_border(tc,
                                top=BORDER_THICK, bottom=BORDER_THICK,
                                left=BORDER_THICK, right=BORDER_THICK)

    def _add_results_table(self, doc, stats, history):
        COLS   = 6
        N_HIST = COLS - 1

        table = doc.add_table(rows=4, cols=COLS)
        table.autofit = False
        self._set_table_cell_margins(table, top=10, bottom=10, left=30, right=30)

        w_col = Inches(6.87 / COLS)
        for i in range(COLS):
            for row in table.rows:
                row.cells[i].width = w_col

        # Compact row heights
        table.rows[0].height      = Inches(0.20)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[1].height      = Inches(0.28)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[2].height      = Inches(0.20)
        table.rows[2].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[3].height      = Inches(0.35)
        table.rows[3].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        # Row 0: merged title
        cell0 = table.cell(0, 0)
        for c in range(1, COLS):
            cell0 = cell0.merge(table.cell(0, c))
        self._style_cell_text(table.cell(0, 0), "النتائج", FONT, 11, bold=True, rtl=True)
        self._set_cell_shading(table.cell(0, 0), SHADE_HEADER)

        # Row 1: instruction (cols 0-4 merged) + current label (col 5)
        instr = table.cell(1, 0)
        for c in range(1, N_HIST):
            instr = instr.merge(table.cell(1, c))
        self._style_cell_text(
            table.cell(1, 0),
            "في حال كان هناك نتائج سابقة عن الموضوع الذي يتم تحليله يجب ذكره مع تحديد الفترة: الشهر / الفصل / العام",
            FONT, 8, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(1, 0), SHADE_HEADER)
        self._style_cell_text(table.cell(1, N_HIST), "النتيجة الحالية", FONT, 9, bold=True, rtl=True)
        self._set_cell_shading(table.cell(1, N_HIST), SHADE_TABLE)

        # Row 2: quarter labels
        last5 = history[-N_HIST:] if len(history) >= N_HIST else list(history)
        while len(last5) < N_HIST:
            last5 = [{}] + last5

        for i in range(N_HIST):
            entry = last5[i]
            label = f"{entry['quarter']} {entry['year']}" if entry.get('quarter') else ""
            self._style_cell_text(table.cell(2, i), label, FONT, 9, bold=True, rtl=True)
            self._set_cell_shading(table.cell(2, i), SHADE_HEADER)
        self._style_cell_text(
            table.cell(2, N_HIST),
            f"{stats.get('quarter', '')} {stats.get('year', '')}",
            FONT, 10, bold=True, rtl=True
        )
        self._set_cell_shading(table.cell(2, N_HIST), SHADE_TABLE)

        # Row 3: rate values
        for i in range(N_HIST):
            entry = last5[i]
            val   = f"{entry['rate']:.2f}%" if entry.get('rate') else ""
            self._style_cell_text(table.cell(3, i), val, FONT_EN, 11)

        current_rate = stats.get('mortality_rate', 0)
        try:
            curr = float(current_rate)
        except Exception:
            curr = 0.0
        self._style_cell_text(table.cell(3, N_HIST), f"%{curr:.2f}", FONT, 11, bold=True, rtl=True)

        # Borders — all sides thick
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell,
                                      top=BORDER_THICK, bottom=BORDER_THICK,
                                      left=BORDER_THICK, right=BORDER_THICK)

    def _add_analysis_box(self, doc, quarter, year, stats, charts, history=None):
        """
        Page 1 bordered analysis box with proper RTL rendering.
        """
        history = history or []
        mortality_rate = stats.get('mortality_rate', 0)

        box = doc.add_table(rows=1, cols=1)
        box.autofit = False
        para_before = doc.paragraphs[-2]
        para_before.paragraph_format.space_before = Pt(0)
        para_before.paragraph_format.space_after = Pt(0)
        box.columns[0].width = USABLE_WIDTH
        cell = box.cell(0, 0)
        self._zero_spacing(cell.paragraphs[0])
        self._set_table_cell_margins(box, top=0, bottom=10, left=20, right=20)
        self._set_cell_border(cell, top=BORDER_MEDIUM, bottom=BORDER_MEDIUM, 
                            left=BORDER_MEDIUM, right=BORDER_MEDIUM)

        # Clear existing content
        cell.text = ""
        for paragraph in cell.paragraphs:
            paragraph._element.getparent().remove(paragraph._element)

        # Title row
        title_tbl = cell.add_table(rows=1, cols=1)
        title_tbl.autofit = False
        title_tbl.columns[0].width = Inches(9)
        title_cell = title_tbl.cell(0, 0)
        self._set_table_cell_margins(title_tbl, top=0, bottom=0, left=20, right=20)
        self._set_cell_shading(title_cell, SHADE_HEADER)
        self._set_cell_border(title_cell, bottom=BORDER_MEDIUM)
        self._style_cell_text(title_cell, "*التحليل", FONT, 11, bold=True, rtl=True)

        # Clear any extra paragraphs
        for p_after in cell.paragraphs:
            p_after._element.getparent().remove(p_after._element)

        # First paragraph - Introduction
        p2 = cell.add_paragraph()
        self._set_rtl(p2)
        self._zero_spacing(p2)
        run1 = p2.add_run(
            f"هذا التحليل يتضمن عدد الوفيات في {quarter} من العام {year}، وتم احتساب نسبة "
            f"الوفيات للمرضى الذين مكثوا في المستشفى أكثر من 24 ساعة.\n"
        )
        run1.font.name = FONT_ANALYSIS
        run1.font.size = Pt(11)
        rPr1 = run1._element.get_or_add_rPr()
        if not rPr1.findall(qn('w:rtl')):
            rPr1.append(OxmlElement('w:rtl'))

        # Subtitle — Arabic run + separate LTR run for "Target", matching IC bidi pattern
        p_subtitle = cell.add_paragraph()
        self._set_rtl(p_subtitle)
        self._zero_spacing(p_subtitle)
        # Arabic part: add <w:rtl/> on the run so Word's bidi engine treats it as RTL
        r_ar = p_subtitle.add_run("- مقارنة نسبة الوفيات مع الـ ")
        r_ar.font.name = FONT_ANALYSIS
        r_ar.font.size = Pt(11)
        r_ar.bold = True
        r_ar.underline = True
        rPr_ar = r_ar._element.get_or_add_rPr()
        if not rPr_ar.findall(qn('w:rtl')):
            rPr_ar.append(OxmlElement('w:rtl'))
        # LTR part: no <w:rtl/> so Word renders "Target" left-to-right within RTL paragraph
        r_en = p_subtitle.add_run("Target")
        r_en.font.name = FONT_ANALYSIS
        r_en.font.size = Pt(11)
        r_en.bold = True
        r_en.underline = True

        # Get data for analysis — read target from config (mortality_metrics has no 'target' key)
        from app.config import load_targets as _load_targets
        target_rate = float(_load_targets().get('mortality', {}).get('rate', 2.0) or 2.0)
        prev = history[-1] if history else {}
        prev_rate = prev.get('rate', 0)
        prev_deaths = prev.get('deaths', 0)
        prev_q = f"{prev.get('quarter', '')} {prev.get('year', '')}"
        curr_deaths = stats.get('kpi_deaths', stats.get('total_deaths', 0))

        # Try to get AI text, use fallback if needed
        ai_text = ai_service.analyze_mortality_trend(
            quarter, year, mortality_rate,
            curr_deaths, prev_rate, prev_deaths, prev_q,
            target_rate=target_rate
        )
        
        if not ai_text:
            vs_target = "أقل من" if mortality_rate < target_rate else "أعلى من"
            trend_word = "ارتفاع" if mortality_rate > prev_rate else "انخفاض"
            diff = abs(mortality_rate - prev_rate)
            
            # FIX: Proper Arabic sentence structure - all text flows RTL
            # The parentheses naturally appear in the correct position in RTL text
            ai_text = (
                f"بلغت نسبة الوفيات في {quarter} من العام {year} مستوى {mortality_rate:.2f}% "
                f"مع تسجيل {curr_deaths} حالة، مقارنة بالفصل السابق ({prev_q}) الذي سجل نسبة "
                f"{prev_rate:.2f}% و{prev_deaths} حالة، مما يشير إلى {trend_word} بمقدار {diff:.2f}%. "
                f"ومع ذلك، فإن النسبة الحالية لا تزال {vs_target} الـ Target البالغ {target_rate:.1f}%."
            )

        # Analysis text paragraph - with proper RTL
        p_ai = cell.add_paragraph()
        self._set_rtl(p_ai)
        self._zero_spacing(p_ai)
        
        # Add the text as a single run - this works better for RTL in Word
        run3 = p_ai.add_run(ai_text)
        run3.font.name = FONT_ANALYSIS
        run3.font.size = Pt(11)
        rPr3 = run3._element.get_or_add_rPr()
        if not rPr3.findall(qn('w:rtl')):
            rPr3.append(OxmlElement('w:rtl'))

        # Chart section
        if 'chart1' in charts:
            p3 = cell.add_paragraph()
            run = p3.add_run(f"Inpatient mortality rate\nT<{target_rate:.1f}%")
            run.font.name = FONT_EN
            run.font.size = Pt(11)
            run.bold = True
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(0)
            p3.paragraph_format.space_after = Pt(4)

            pimg = cell.add_paragraph()
            pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pimg.add_run().add_picture(charts['chart1'], width=Inches(6.8))
            pimg.paragraph_format.space_before = Pt(0)
            pimg.paragraph_format.space_after = Pt(2)

        # Ensure zero spacing on all paragraphs
        for p in cell.paragraphs:
            self._zero_spacing(p)

    # =========================================================================
    # PAGE 3 COMPONENTS
    # =========================================================================

    def _add_dept_chart_with_table(self, doc, chart7, departments):
        layout_tbl = doc.add_table(rows=1, cols=2)
        layout_tbl.autofit = False
        for cell in layout_tbl.rows[0].cells:
            self._set_cell_border(cell, top=BORDER_NONE, bottom=BORDER_NONE, left=BORDER_NONE, right=BORDER_NONE)

        layout_tbl.rows[0].cells[0].width = Inches(2.67)
        layout_tbl.rows[0].cells[1].width = Inches(4.2)

        chart_cell = layout_tbl.rows[0].cells[1]
        chart_p    = chart_cell.paragraphs[0]
        chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._zero_spacing(chart_p)
        chart_p.add_run().add_picture(chart7, width=Inches(4))

        data_cell = layout_tbl.rows[0].cells[0]
        tcPr      = data_cell._element.get_or_add_tcPr()
        tcMar     = OxmlElement('w:tcMar')
        for edge, val in [('w:right', '80'), ('w:left', '80')]:
            m = OxmlElement(edge)
            m.set(qn('w:w'), val)
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)

        for p in data_cell.paragraphs:
            p._element.getparent().remove(p._element)

        dept_display_names = {'c2': 'Cardiac 2'}
        total_count = sum(d.get('count', 0) for d in departments)
        dept_data   = [('النسبة', 'عدد الوفيات', 'القسم')]
        for dept in departments:
            cnt  = dept.get('count', 0)
            pct  = f"{(cnt / total_count * 100):.0f}%" if total_count > 0 else "0%"
            name = dept_display_names.get(dept.get('name', '').lower().strip(), dept.get('name', ''))
            dept_data.append((pct, str(cnt), name))
        dept_data.append(('100%', str(total_count), 'المجموع'))

        num_rows   = len(dept_data)
        dept_table = data_cell.add_table(rows=num_rows, cols=3)
        dept_table.autofit = False
        self._set_table_cell_margins(dept_table, top=20, bottom=20, left=40, right=40)

        last_row = num_rows - 1
        for i, (col1, col2, col3) in enumerate(dept_data):
            row_cells = dept_table.rows[i].cells
            for j, val in enumerate([col1, col2, col3]):
                is_arabic = any('\u0600' <= c <= '\u06FF' for c in val)
                font = FONT if is_arabic else FONT_EN
                self._style_cell_text(row_cells[j], val, font, 11,
                                      bold=(i == 0 or i == last_row), rtl=is_arabic, align='center')
                self._set_cell_border(row_cells[j], top=BORDER_FINE, bottom=BORDER_FINE,
                                      left=BORDER_FINE, right=BORDER_FINE)
            if i == 0:
                self._shade_row_cells(row_cells, SHADE_TABLE)
            elif i == last_row:
                self._shade_row_cells(row_cells, SHADE_TOTAL)

        for row in dept_table.rows:
            for j, w in enumerate([Inches(0.7), Inches(0.87), Inches(1.0)]):
                row.cells[j].width = w

    def _add_dept_table_only(self, doc, departments):
        """Render a standalone department table (no chart) when chart7 is unavailable."""
        dept_display_names = {'c2': 'Cardiac 2'}
        total_count = sum(d.get('count', 0) for d in departments)
        dept_data = [('النسبة', 'عدد الوفيات', 'القسم')]
        for dept in departments:
            cnt  = dept.get('count', 0)
            pct  = f"{(cnt / total_count * 100):.0f}%" if total_count > 0 else "0%"
            name = dept_display_names.get(dept.get('name', '').lower().strip(), dept.get('name', ''))
            dept_data.append((pct, str(cnt), name))
        dept_data.append(('100%', str(total_count), 'المجموع'))

        num_rows  = len(dept_data)
        table     = doc.add_table(rows=num_rows, cols=3)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        last_row = num_rows - 1
        for i, (col1, col2, col3) in enumerate(dept_data):
            row_cells = table.rows[i].cells
            for j, val in enumerate([col1, col2, col3]):
                is_arabic = any('\u0600' <= c <= '\u06FF' for c in val)
                font = FONT if is_arabic else FONT_EN
                self._style_cell_text(row_cells[j], val, font, 11,
                                      bold=(i == 0 or i == last_row), rtl=is_arabic, align='center')
                self._set_cell_border(row_cells[j], top=BORDER_FINE, bottom=BORDER_FINE,
                                      left=BORDER_FINE, right=BORDER_FINE)
            if i == 0:
                self._shade_row_cells(row_cells, SHADE_TABLE)
            elif i == last_row:
                self._shade_row_cells(row_cells, SHADE_TOTAL)

        for row in table.rows:
            row.cells[0].width = Inches(1.0)
            row.cells[1].width = Inches(1.2)
            row.cells[2].width = Inches(4.67)

    # =========================================================================
    # PAGE 5 COMPONENTS
    # =========================================================================

    def _add_doctor_specialty_table(self, doc, specialties):
        total_count = sum(s.get('count', 0) for s in specialties)
        doctor_data = [('اختصاص الطبيب', 'العدد', 'النسبة')]
        for spec in specialties:
            count = spec.get('count', 0)
            pct = (count / total_count * 100) if total_count > 0 else 0
            doctor_data.append((
                spec.get('name', ''),
                str(count),
                f"{pct:.1f}%"
            ))
        doctor_data.append(('المجموع', str(total_count), '100%'))

        num_rows = len(doctor_data)
        table    = doc.add_table(rows=num_rows, cols=3)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        last_row = num_rows - 1
        for i, (col1, col2, col3) in enumerate(doctor_data):
            row_cells = table.rows[i].cells
            for j, val in enumerate([col1, col2, col3]):
                is_arabic = any('\u0600' <= c <= '\u06FF' for c in val)
                font = FONT if is_arabic else FONT_EN
                self._style_cell_text(row_cells[j], val, font, 11,
                                      bold=(i == 0 or i == last_row), rtl=True, align='center')
                self._set_cell_border(row_cells[j], top=BORDER_FINE, bottom=BORDER_FINE,
                                      left=BORDER_FINE, right=BORDER_FINE)
            if i == 0:
                self._shade_row_cells(row_cells, SHADE_TABLE)
            elif i == last_row:
                self._shade_row_cells(row_cells, SHADE_TOTAL)

        for row in table.rows:
            row.cells[0].width = Inches(3.5)
            row.cells[1].width = Inches(1.0)
            row.cells[2].width = Inches(1.0)

    # =========================================================================
    # LAST PAGE COMPONENTS
    # =========================================================================

    def _add_checkbox_question_box(self, doc):
        box_tbl = doc.add_table(rows=1, cols=1)
        box_tbl.autofit = False
        box_tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(box_tbl, top=30, bottom=30, left=60, right=60)
        cell = box_tbl.cell(0, 0)
        self._set_cell_border(cell, top=BORDER_FINE, bottom=BORDER_FINE, left=BORDER_FINE, right=BORDER_FINE)

        p = cell.paragraphs[0]
        self._set_rtl(p)
        self._zero_spacing(p)
        self._add_run(p, 'هل تم اتخاذ اجراءات تحسينية سابقا حول موضوع التحليل :      ', FONT_ANALYSIS, 11)
        self._add_run(p, 'نعم', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐', None, 11)
        self._add_run(p, '         ', None, 11)
        self._add_run(p, 'كلا  ', FONT_ANALYSIS, 11)
        self._add_run(p, ' ☐ ', None, 11)

        p2 = cell.add_paragraph()
        self._set_rtl(p2)
        self._zero_spacing(p2)
        self._add_run(p2,
            'في حال نعم , يجب تعبئة جدول " الاجراءات المتخذة سابقا" '
            'و في حال كلا , الانتقال الى جدول "الاجراءات المتخذة الحالية"',
            FONT_ANALYSIS, 11
        )

    def _add_previous_actions_table(self, doc):
        table = doc.add_table(rows=3, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 4))
        self._style_cell_text(table.cell(0, 0), 'الاجراءات المتخذة  السابقة', FONT, 10, bold=True, rtl=True, align='center')

        headers = [
            'الرقم', 'الاجراءات المقترحة سابقا',
            'التاريخ المتوقع لتنفيذ الاجراء',
            'تم التنفيذ )نعم/كلا(',
            'في حال كلا ) الاجراءات التصحيحية الجديدة('
        ]
        for j, h in enumerate(headers):
            self._style_cell_text(table.cell(1, j), h, FONT, 11, bold=True, rtl=True, align='center')

        for j in range(4):
            table.cell(2, j).text = ''

        last_cell = table.cell(2, 4)
        last_cell.text = ''
        p1 = last_cell.paragraphs[0]
        self._add_run(p1, ' ☐  الغاء الاجراء', FONT, 11)
        self._zero_spacing(p1)
        self._set_rtl(p1)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = last_cell.add_paragraph()
        self._add_run(p2, ' ☐ اضافة الاجراء ضمن الاجراءات المتخذة الحالية.)الجدول ادناه(', FONT, 11)
        self._zero_spacing(p2)
        self._set_rtl(p2)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(2):
            self._shade_row_cells(table.rows[i].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(1.8), Inches(1.37), Inches(1.1), Inches(2.1)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_current_actions_table(self, doc):
        table = doc.add_table(rows=6, cols=4)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).merge(table.cell(0, 3))
        self._style_cell_text(table.cell(0, 0), 'الاجراءات المتخذة  الحالية', FONT, 10, bold=True, rtl=True, align='center')

        headers = ['الرقم', 'الاجراءات الواجب اتخاذها لحل الموضوع', 'الشخص المعني بالمتابعة', 'التاريخ المتوقع للتنفيذ']
        for j, h in enumerate(headers):
            self._style_cell_text(table.cell(1, j), h, FONT, 11, bold=True, rtl=True, align='center')

        for i in range(2, 6):
            for j in range(4):
                table.cell(i, j).text = ''

        for i in range(2):
            self._shade_row_cells(table.rows[i].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.5), Inches(3.07), Inches(1.8), Inches(1.5)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w

    def _add_approval_table(self, doc):
        table = doc.add_table(rows=4, cols=5)
        table.autofit = False
        self._set_table_cell_margins(table, top=20, bottom=20, left=40, right=40)
        self._set_table_rtl(table)

        table.cell(0, 0).text = ''
        self._style_cell_text(table.cell(0, 1), 'اعداد', FONT, 11, bold=True, rtl=True, align='center')
        table.cell(0, 2).merge(table.cell(0, 4))
        self._style_cell_text(table.cell(0, 2), 'موافقة', FONT, 11, bold=True, rtl=True, align='center')

        for i, label in enumerate(['الاسم', 'التوقيع', 'التاريخ']):
            self._style_cell_text(table.cell(i + 1, 0), label, FONT, 11, bold=True, rtl=True, align='center')
            for j in range(1, 5):
                table.cell(i + 1, j).text = ''

        self._shade_row_cells(table.rows[0].cells, SHADE_HEADER)
        self._apply_borders_to_table(table, BORDER_FINE)

        widths = [Inches(0.8), Inches(1.5), Inches(1.52), Inches(1.52), Inches(1.53)]
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w


    # =========================================================================
    # WHO STATIC ANALYSIS — Arabic lookup table + templates
    # Replaces AI calls to prevent hallucinated disease names
    # =========================================================================

    # Correct Arabic translations for all WHO ICD-10 categories
    WHO_AR = {
        'cardiovascular disease':              'أمراض القلب والأوعية الدموية',
        'cardiovascular diseases':             'أمراض القلب والأوعية الدموية',
        'genitourinary diseases':              'أمراض الجهاز البولي التناسلي',
        'injuries':                            'الإصابات والحوادث',
        'malignant neoplasms':                 'الأورام الخبيثة',
        'congenital anomalies':                'التشوهات الخلقية',
        'defined disease':                     'أمراض محددة',
        'digestive diseases':                  'أمراض الجهاز الهضمي',
        'infectious and parasitic diseases':   'الأمراض المعدية والطفيلية',
        'respiratory diseases':                'أمراض الجهاز التنفسي',
        'neuropsychiatric conditions':         'الاضطرابات العصبية والنفسية',
        'perinatal conditions':                'حالات ما حول الولادة',
        'diabetes mellitus and endocrine disorders': 'داء السكري واضطرابات الغدد الصماء',
        'musculoskeletal diseases':            'أمراض الجهاز العضلي الهيكلي',
        'nutritional deficiencies':            'نقص التغذية',
        'skin diseases':                       'أمراض الجلد',
        'sense organ diseases':                'أمراض الحواس',
        'other neoplasms':                     'الأورام الأخرى',
        'infection following a procedure, not elsewhere classified':
                                               'عدوى ما بعد الإجراء الطبي',
    }

    def _who_ar(self, category):
        """Return correct Arabic name; falls back to original if not in table."""
        return self.WHO_AR.get(category.lower().strip(), category)

    def _build_who_comparison_text(self, quarter, top_who):
        """Static WHO comparison text for chart9."""
        if not top_who:
            return (
                f"بناءً لتصنيف منظمة الصحة العالمية، يُظهر الرسم البياني توزيع "
                f"الوفيات بحسب التشخيص في {quarter}."
            )
        parts = [
            f"{self._who_ar(item.get('category',''))} "
            f"{LRE}({item.get('count',0)} حالة){PDF}"
            for item in top_who
        ]
        # First item is the highest — highlight it
        first = parts[0]
        rest  = "، ".join(parts[1:])
        return (
            f"في {quarter}، تصدّرت {first} قائمةَ أسباب الوفاة، "
            f"تلتها كلٌّ من: {rest}."
        )

    def _build_who_diagnosis_text(self, quarter, year, top_cat, top_cnt):
        """Static WHO diagnosis intro text for chart10."""
        ar_top = self._who_ar(top_cat)
        return (
            f"بناءً لتصنيف منظمة الصحة العالمية {LRE}WHO{PDF} تم تصنيف "
            f"{LRE}ICD10{PDF} التابعة للوفيات في {quarter} من العام {year} "
            f"وكانت النتيجة على الشكل التالي، حيث احتلّت {ar_top} "
            f"المرتبة الأولى بعدد {top_cnt} حالة:"
        )

    # =========================================================================
    # HELPERS
    # =========================================================================

    def _add_spacer(self, doc, space_before=0):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(space_before)
        sp.paragraph_format.space_after  = Pt(0)
        sp.paragraph_format.line_spacing = Pt(1)
        sp.add_run("").font.size = Pt(1)

    def _add_centered_chart(self, doc, chart_bytes, width=6.5):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._zero_spacing(para)
        para.add_run().add_picture(chart_bytes, width=Inches(width))

    def _add_bidi_para(self, container, text, font_name, size, bold=False):
        """
        Render mixed Arabic/English text with correct BiDi direction.
        Splits text at Arabic/LTR boundaries and assigns explicit run-level
        direction markers so parentheses, %, and English tokens render correctly.
        """
        # Must start with a Latin LETTER so standalone numbers stay in the RTL flow
        # (e.g. "34 حالة" renders correctly: number flows RTL, appears before the noun)
        LTR_PAT = re.compile(r'[A-Za-z][A-Za-z0-9 ().,;:%&/_\-+<>=*]*')
        parts, cursor = [], 0
        for m in LTR_PAT.finditer(text):
            if m.start() > cursor:
                parts.append((text[cursor:m.start()], False))
            parts.append((m.group(), True))
            cursor = m.end()
        if cursor < len(text):
            parts.append((text[cursor:], False))
        if not parts:
            parts = [(text, False)]


        para = container.add_paragraph()
        self._set_rtl(para)
        self._zero_spacing(para)

        for chunk, is_ltr in parts:
            if not chunk:
                continue
            display = (' ' + chunk.strip() + ' ') if is_ltr else chunk
            run = para.add_run(display)
            run.font.size = Pt(size)
            run.bold = bold
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            if is_ltr:
                run.font.name = FONT_EN
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), FONT_EN)
                for el in rPr.findall(qn('w:rtl')):
                    rPr.remove(el)
            else:
                run.font.name = font_name
                for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), font_name)
                if not rPr.findall(qn('w:rtl')):
                    rPr.append(OxmlElement('w:rtl'))

    def _add_analysis_text(self, doc, text):
        # Strip Unicode bidi control characters the AI may embed in output
        for ch in ('\u202a', '\u202b', '\u202c', '\u202d', '\u202e',
                   '\u200e', '\u200f', '\u2066', '\u2067', '\u2068', '\u2069'):
            text = text.replace(ch, '')
        # Remove parentheses
        text = text.replace('(', '').replace(')', '')
        # Insert Arabic connector before standalone percentages that follow a dept acronym
        # e.g. "ICU 40%" → "ICU بنسبة 40%"
        text = re.sub(r'([A-Z]{2,6})\s+(\d+%)', r'\1 بنسبة \2', text)
        self._add_bidi_para(doc, text, FONT_ANALYSIS, 11)

    def _add_titled_row(self, doc, text, shade=None, border=BORDER_FINE):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.autofit = False
        tbl.columns[0].width = USABLE_WIDTH
        self._set_table_cell_margins(tbl, top=0 if shade else 20, bottom=0 if shade else 20, left=20, right=20)
        cell = tbl.cell(0, 0)
        if shade:
            self._set_cell_shading(cell, shade)
        self._set_cell_border(cell, top=border, bottom=border, left=border, right=border)
        self._style_cell_text(cell, text, FONT, 11, bold=True, rtl=True, align='center')

    def _add_rtl_para(self, doc, text, size, bold=False, font=None):
        para = doc.add_paragraph()
        run  = para.add_run(text)
        run.font.name = font or FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.underline = True
        self._set_rtl(para)
        self._zero_spacing(para)

    def _add_run(self, para, text, font_name, size, bold=False):
        run = para.add_run(text)
        if font_name:
            run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        return run

    def _zero_spacing(self, para):
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing = 1

    def _set_rtl(self, para):
        para.paragraph_format.right_to_left = True
        pPr = para._element.get_or_add_pPr()
        if pPr.find(qn('w:bidi')) is None:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

    def _set_table_rtl(self, table):
        bidi = OxmlElement('w:bidiVisual')
        bidi.set(qn('w:val'), '1')
        table._tbl.tblPr.append(bidi)

    def _shade_row_cells(self, cells, color):
        for cell in cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            shading.set(qn('w:val'), 'clear')
            cell._element.get_or_add_tcPr().append(shading)

    def _apply_borders_to_table(self, table, border_style):
        for row in table.rows:
            for cell in row.cells:
                self._set_cell_border(cell, top=border_style, bottom=border_style,
                                      left=border_style, right=border_style)

    def _set_table_cell_margins(self, table, top=50, bottom=50, left=80, right=80):
        tblCellMar = OxmlElement('w:tblCellMar')
        for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)
        table._tbl.tblPr.append(tblCellMar)

    def _set_cell_shading(self, cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill_hex)
        tcPr.append(shd)

    def _set_cell_border(self, cell, **kwargs):
        tc        = cell._element
        tcPr      = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                tag     = f'w:{edge}'
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                for k, v in kwargs[edge].items():
                    element.set(qn('w:' + k), str(v))

    def _set_tc_border(self, tc, **kwargs):
        tcPr      = tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for edge in ('top', 'left', 'bottom', 'right'):
            if edge in kwargs:
                tag     = f'w:{edge}'
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                for k, v in kwargs[edge].items():
                    element.set(qn('w:' + k), str(v))

    def _style_cell_text(self, cell, text, font_name, font_size, bold=False, rtl=False, align='center', underline=False):
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = font_name
        r.font.size = Pt(font_size)
        r.bold = bold
        r.underline = underline
        self._zero_spacing(p)
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if rtl:
            self._set_rtl(p)


# =============================================================================
# SINGLETON
# =============================================================================
matplotlib_docx_generator = MatplotlibDocxGenerator()